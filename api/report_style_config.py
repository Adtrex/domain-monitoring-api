"""Central style and copy configuration for PDF and DOCX report generation."""

from datetime import datetime, timezone
from io import BytesIO
import logging
import re

from .cve_checker import version_compare, fetch_eol_status, _EOL_PRODUCT_MAP

logger = logging.getLogger(__name__)


def _format_generated_at(value):
    """Render the report timestamp in plain language, e.g. '22 June 2026 at 13:36 UTC'.

    Accepts a datetime or an ISO string; falls back to the original value if it
    cannot be parsed so the header is never left blank.
    """
    dt = value
    if isinstance(value, str):
        try:
            dt = datetime.fromisoformat(value)
        except ValueError:
            return value
    try:
        # %d keeps a leading zero on single-digit days; strip it for natural reading.
        return dt.strftime('%d %B %Y at %H:%M UTC').lstrip('0')
    except (AttributeError, ValueError):
        return str(value)

REPORT_TEXT = {
    'title': 'CERRT Vulnerability Remediation Report',
    'how_to_use_title': 'HOW TO USE THIS REPORT',
    'findings_title': 'FINDINGS AND RECOMMENDED ACTIONS',
    'checks_title': 'ADDITIONAL SCAN CHECKS',
    'no_findings': 'No findings recorded for this scan.',
}

HOW_TO_USE_STEPS = [
    'Fix Critical and High findings first.',
    'Validate each fix in staging before production.',
    'Re-run scan to confirm issues are closed.',
]

PDF_STYLE = {
    'page': {
        'left_margin_mm': 14,
        'right_margin_mm': 14,
        'top_margin_mm': 14,
        'bottom_margin_mm': 14,
    },
    'fonts': {
        'title_size': 18,
        'section_size': 8,
        'body_size': 9,
        'small_size': 8,
        'meta_label_size': 7,
        'meta_value_size': 13,
    },
    'colors': {
        # Base palette
        'white': '#FFFFFF',
        'page_bg': '#F5F5F3',
        'surface': '#FFFFFF',
        'surface_secondary': '#F5F5F3',

        # Text
        'text_primary': '#1A1A18',
        'text_secondary': '#5F5E5A',
        'text_muted': '#888780',
        'text_hint': '#B4B2A9',

        # Borders
        'border': '#E0DFD8',
        'border_light': '#EEECEA',

        # Severity badge backgrounds and text (pill style)
        'severity': {
            'Critical': {
                'text': '#A32D2D',
                'bg': '#FCEBEB',
                'border': '#F7C1C1',
                'dot': '#E24B4A',
            },
            'High': {
                'text': '#854F0B',
                'bg': '#FAEEDA',
                'border': '#FAC775',
                'dot': '#EF9F27',
            },
            'Medium': {
                'text': '#185FA5',
                'bg': '#E6F1FB',
                'border': '#B5D4F4',
                'dot': '#378ADD',
            },
            'Low': {
                'text': '#3B6D11',
                'bg': '#EAF3DE',
                'border': '#C0DD97',
                'dot': '#639922',
            },
            'default': {
                'text': '#444441',
                'bg': '#F1EFE8',
                'border': '#D3D1C7',
                'dot': '#888780',
            },
        },

        # Stats grid
        'stat_label': '#888780',
        'stat_value': '#1A1A18',
        'stat_bg': '#F5F5F3',
        'stat_border': '#E0DFD8',

        # Finding card
        'card_bg': '#FFFFFF',
        'card_border': '#E0DFD8',
        'card_header_bg': '#F5F5F3',
        'action_box_bg': '#F5F5F3',

        # Section label
        'section_label': '#888780',

        # Status
        'status_completed': '#3B6D11',
        'link': '#185FA5',
    },
    'spacing': {
        'card_padding_mm': 4,
        'card_gap_mm': 3,
        'section_gap_mm': 5,
    },
}

DOCX_STYLE = {
    'stats_table_style': 'Table Grid',
    'finding_table_style': 'Light List Accent 1',
}

RISK_ORDER = {
    'Critical': 0,
    'High': 1,
    'Medium': 2,
    'Low': 3,
}

RISK_OWNER = {
    'Critical': 'Security / Platform team',
    'High': 'Security / Platform team',
    'Medium': 'Application team',
    'Low': 'Application team',
}


def risk_order(value):
    return RISK_ORDER.get(value or 'Low', 4)


def parse_remediation(recommendation):
    text = (recommendation or '').strip()
    if not text:
        return 'No remediation provided.', []

    links = re.findall(r'https?://[^\s)]+', text)
    cleaned = re.sub(r'https?://[^\s)]+', '', text).strip()
    if not cleaned:
        cleaned = 'Review the remediation resources listed below.'
    return cleaned, links


_GENERIC_PHRASES = (
    'investigate this finding',
    'no remediation provided',
    'assess its exploitability',
    'update the affected library or plugin',
    'validate the finding and apply remediation',
    'apply the header with a secure value',
    'review vendor guidance',
    'configure spf, dkim, and dmarc records',
    'review and remediate this finding',
    'detected vulnerable library version',
    'detected outdated library version',
    'technology category:',
    'header status is missing',
    'cms: wordpress',
    'an spf txt record was detected',
    'outdated wordpress core, plugins, or themes',
)

_TITLE_REMEDIATION = [
    ('content-security-policy', 'Deploy a Content-Security-Policy (CSP) response header on all HTTP responses. Start with default-src \'self\' and whitelist only trusted sources. Enable report-uri to collect violations before enforcing.'),
    ('strict-transport-security', 'Add Strict-Transport-Security: max-age=31536000; includeSubDomains to all HTTPS responses. Validate, then submit the domain to the HSTS preload list. Ensure HTTP redirects to HTTPS before enabling preload.'),
    ('x-content-type-options', 'Add X-Content-Type-Options: nosniff to all HTTP responses in the web server or application configuration. This prevents MIME-sniffing attacks that allow browsers to execute attacker-supplied content as scripts.'),
    ('x-frame-options', 'Set X-Frame-Options: DENY on all HTTP responses. If same-origin embedding is needed use SAMEORIGIN. Complement with Content-Security-Policy: frame-ancestors for fine-grained control.'),
    ('permissions-policy', 'Add a Permissions-Policy header restricting unused browser APIs. Example: Permissions-Policy: geolocation=(), microphone=(), camera=(). Review and update this policy whenever new browser features are adopted.'),
    ('referrer-policy', 'Set Referrer-Policy: no-referrer or strict-origin-when-cross-origin on all responses. This limits referrer data sent to third-party sites, preventing leakage of internal URLs and session paths.'),
    ('x-xss-protection', 'Add X-XSS-Protection: 1; mode=block to HTTP responses. Pair it with a strong Content-Security-Policy — modern browsers rely on CSP for XSS defence.'),
    ('swiper', 'Update Swiper to the latest stable release. Review the Swiper GitHub security advisories for patched CVEs. If loaded via CDN, update the URL and add a Subresource Integrity (SRI) hash. Test all carousel components after upgrading.'),
    ('chart', 'Update Chart.js to the latest stable release. Check the Chart.js GitHub advisories for patched CVEs. If loaded via CDN, update the URL and add a Subresource Integrity (SRI) hash. Re-test all charts after the upgrade.'),
    ('jquery', 'Update jQuery to the latest stable release (3.7.x or newer). If loaded from a CDN, add a Subresource Integrity (SRI) hash. Audit all .html(), .append(), and similar sink methods to ensure no untrusted data flows into them.'),
    ('bootstrap', 'Upgrade Bootstrap to the latest stable release (5.x). Review the Bootstrap security changelog, update JavaScript bundles, and regression-test the UI before deploying to production.'),
    ('animate', 'Update Animate.css to the latest stable release. If loaded via CDN, pin the version and add a Subresource Integrity (SRI) hash. Remove unused animation classes to reduce the dependency surface.'),
    ('font awesome', 'Update Font Awesome to the latest stable release. Update the CDN link or kit URL and add a Subresource Integrity (SRI) hash. Audit for unused icon sets to minimise the dependency footprint.'),
    ('moment', 'Moment.js is in maintenance-only mode with no active security fixes. Migrate to Day.js, date-fns, or Luxon. If migration is not immediate, pin the current version and monitor advisories closely.'),
    ('revslider', 'Update Revolution Slider to the latest version via the theme vendor dashboard. RevSlider has had critical vulnerabilities including arbitrary file upload. Verify its AJAX handler is not accessible to unauthenticated users.'),
    ('contact form', 'Update Contact Form 7 and all addons to the latest version. Restrict file uploads to safe types, confirm CSRF protection is active, and remove unused addons.'),
    ('lodash', 'Update Lodash to version 4.17.21 or later to fix prototype-pollution vulnerabilities. Switch to lodash-es or individual function imports to reduce the dependency surface.'),
    ('axios', 'Update Axios to the latest stable release to address known CSRF and SSRF vulnerabilities. After upgrading, verify request interceptors and base URL configuration remain correct.'),
    ('vue', 'Update Vue.js to the latest stable release (Vue 3.x). If migrating from Vue 2 (end-of-life Dec 2023), follow the official migration build process to avoid breaking changes.'),
    ('angular', 'Update Angular to the latest LTS release. Run ng update at update.angular.io and execute the full test suite before deploying to production.'),
    ('react', 'Update React and ReactDOM to the latest stable release. Run npm/yarn update, resolve peer-dependency conflicts, and execute the full test suite after upgrading.'),
    ('caa record', 'Publish a CAA DNS record to restrict which CAs may issue certificates for the domain, e.g. 0 issue "letsencrypt.org". Verify with a DNS lookup tool. This prevents fraudulent certificate issuance.'),
    ('subdomain_takeover', 'Audit all CNAME and A records pointing to decommissioned or unclaimed services. Remove or update stale DNS entries immediately. Add a DNS cleanup step to all cloud resource deprovisioning checklists.'),
    ('deprecated tls', 'Disable TLS 1.0 and 1.1 at the web server and load-balancer level. Accept only TLS 1.2 and 1.3. Use the Mozilla SSL Configuration Generator as a reference, then verify with SSLLabs.'),
    ('weak cipher', 'Remove weak cipher suites (RC4, DES, 3DES, NULL, EXPORT) from the TLS configuration. Allow only ECDHE and DHE key exchange with AES-GCM or ChaCha20-Poly1305. Re-test with an SSL scanner after changes.'),
    ('wildcard tls', 'Limit wildcard certificates to specific subdomain groups. Store the private key in an HSM or secrets manager. Rotate the certificate and key at least annually.'),
    ('ssl dns names', 'Review the Subject Alternative Names (SANs) in the TLS certificate. Remove unexpected or sensitive hostnames and reissue with only the required SANs.'),
    ('hsts', 'Enable HSTS: Strict-Transport-Security: max-age=31536000; includeSubDomains; preload. Confirm every subdomain is reachable over HTTPS, then submit to hstspreload.org.'),
    ('tls version', 'Enforce a minimum of TLS 1.2 and prefer TLS 1.3. Disable SSLv3, TLS 1.0, and TLS 1.1 at the application server and upstream load balancers. Confirm the change with an SSLLabs Full Report.'),
    ('ssl certificate issuer', 'Replace self-signed or private-CA certificates used on public-facing services with publicly trusted CA certificates. Automate renewal using ACME (e.g. Let\'s Encrypt with certbot).'),
    ('dkim', 'Publish a DKIM public key in DNS and configure your MTA to sign all outgoing messages with at least 2048-bit keys. Rotate DKIM keys annually and monitor DMARC aggregate reports for failures.'),
    ('dmarc', 'Publish a DMARC record starting with p=none to monitor. Analyse aggregate reports (rua), fix SPF/DKIM alignment issues, then progress to p=quarantine and finally p=reject.'),
    ('spf', 'Publish an SPF TXT record listing all authorised mail servers. End with -all (hard fail) to reject mail from unlisted senders. Validate and update whenever mail infrastructure changes.'),
    ('php', 'Upgrade PHP to the latest supported release published at php.net/supported-versions. End-of-life PHP versions no longer receive security patches and must be moved onto the current stable branch. Test in staging before deploying to production.'),
    ('wordpress', 'Keep WordPress core, active themes, and every plugin updated. Remove inactive plugins and themes. Restrict wp-admin by IP if possible, disable XML-RPC if not required, and deploy a WordPress WAF rule set.'),
    ('nginx', 'Update nginx to the latest stable release. Set server_tokens off to suppress version disclosure. Review nginx security advisories for the installed version and disable unused modules.'),
    ('clickjacking', 'Add X-Frame-Options: DENY and Content-Security-Policy: frame-ancestors \'none\' to all HTTP responses. These prevent the page from being loaded inside a third-party iframe.'),
    ('mx record', 'Verify all MX records point to authorised operational mail servers. Remove stale or legacy MX entries and review them quarterly.'),
    ('ns record', 'Confirm NS records point only to authorised authoritative name servers. Restrict zone transfers (AXFR) to specific secondary server IPs via ACL. Remove any rogue NS entries immediately.'),
    ('txt record', 'Audit all TXT records and remove entries that are no longer required (expired SPF includes, old verification tokens, decommissioned vendor records).'),
    ('elementor', 'Update Elementor (and Elementor Pro) to the latest version. Review the security changelog and apply recommended hardening such as restricting widget access by role.'),
    ('csp unknown', 'Investigate why the CSP status is unknown. Confirm the header is present on all responses and is valid. Parse with a CSP evaluator tool and fix overly permissive directives such as unsafe-inline or unsafe-eval.'),
]

_CATEGORY_REMEDIATION = {
    'cve': 'Apply the vendor-supplied security patch for this CVE. If no patch is available implement the vendor-recommended mitigation and isolate the affected component. Validate the fix by re-scanning after the patch is applied.',
    'ssl': 'Review the SSL/TLS configuration against current best practices. Disable deprecated protocols and weak cipher suites, enforce TLS 1.2 or higher, and re-test using an SSL scanner after applying changes.',
    'dns': 'Audit all DNS records for accuracy and remove any stale or unauthorised entries. Enable DNSSEC where supported to protect against DNS spoofing and cache poisoning attacks.',
    'email': 'Verify SPF, DKIM, and DMARC records are correctly configured and aligned with your mail sending infrastructure. Progress DMARC policy toward p=reject to prevent spoofing.',
    'misconfiguration': 'Correct the misconfiguration by applying the vendor-recommended secure default. Test the change in staging, document it, and verify through a configuration audit.',
    'header': 'Add or correct the missing security header in the web server or application configuration. Ensure the header is applied to all HTTP responses including error and redirect pages.',
    'library': 'Update the identified library to the latest stable supported version. Review the changelog for breaking changes, run the test suite, and deploy to staging before pushing to production.',
}

_RISK_SUFFIX = {
    'critical': ' Escalate immediately — patch, verify, and confirm closure within 24 hours.',
    'high': ' Prioritise this fix — target remediation within 7 days.',
    'medium': ' Schedule remediation within the next sprint or 30-day window.',
}

_TITLE_SUMMARY = [
    ('content-security-policy', 'A Content Security Policy (CSP) header is missing, leaving the site open to Cross-Site Scripting (XSS) where attackers inject malicious scripts into pages viewed by other users.'),
    ('enforce https', 'The site accepts unencrypted HTTP connections. Data in transit — credentials, tokens, form submissions — can be intercepted or tampered with by an attacker with network access.'),
    ('strict-transport-security', 'HTTP Strict Transport Security (HSTS) is not set. Browsers may fall back to plain HTTP, enabling downgrade attacks and cookie hijacking on shared or hostile networks.'),
    ('x-content-type-options', 'The X-Content-Type-Options header is absent. Browsers may misinterpret response content and execute attacker-supplied data as scripts, enabling MIME-confusion and XSS attacks.'),
    ('x-frame-options', 'X-Frame-Options protection is missing. The page can be embedded in an attacker-controlled iframe and used for clickjacking, tricking users into performing unintended actions.'),
    ('permissions-policy', 'No Permissions-Policy header restricts browser API access. Embedded scripts may access geolocation, camera, or microphone without the user\'s knowledge.'),
    ('referrer-policy', 'A Referrer-Policy header is not configured. Full referrer URLs are sent to third-party sites, potentially leaking internal paths, user identifiers, and navigation patterns.'),
    ('x-xss-protection', 'The X-XSS-Protection header is absent. This legacy browser filter adds a defence-in-depth layer against reflected XSS and should be set alongside a strong CSP.'),
    ('hsts', 'HSTS is weak or missing. Without a long max-age and includeSubDomains, browsers can be coerced into connecting over plain HTTP, enabling interception and session hijacking.'),
    ('swiper', 'A vulnerable version of the Swiper JavaScript library is in use. Known vulnerabilities in this version may allow attackers to exploit slider and carousel components on the page.'),
    ('chart', 'A vulnerable version of Chart.js is in use. Known vulnerabilities in this version could allow attackers to exploit charting components, potentially leading to script injection or data exposure.'),
    ('jquery', 'An outdated jQuery version is in use. Older releases contain known XSS vulnerabilities in methods like .html() and .append() that attackers can exploit when user-supplied data flows through them.'),
    ('bootstrap', 'An outdated Bootstrap version is detected. Some earlier releases contain XSS and prototype-pollution vulnerabilities in JavaScript components that attackers can exploit to manipulate the DOM.'),
    ('animate', 'A vulnerable version of Animate.css is in use. Outdated CSS libraries may contain vulnerabilities that could be leveraged in injection attacks or unintended behaviour in the browser.'),
    ('font awesome', 'A vulnerable version of Font Awesome is detected. Outdated icon library versions may contain security issues that could be exploited through crafted content or CDN-based attacks.'),
    ('moment', 'A vulnerable version of Moment.js is in use. The library is in maintenance-only mode and known vulnerabilities (including ReDoS) will not receive active patches going forward.'),
    ('revslider', 'A vulnerable version of Revolution Slider (RevSlider) is detected. RevSlider has historically had critical vulnerabilities including arbitrary file upload and remote code execution.'),
    ('lodash', 'A vulnerable version of Lodash is in use. Known prototype-pollution vulnerabilities in older Lodash versions allow attackers to manipulate the JavaScript prototype chain and alter application behaviour.'),
    ('axios', 'A vulnerable version of Axios is detected. Older versions contain CSRF and SSRF vulnerabilities that allow attackers to forge requests and potentially access internal resources.'),
    ('vue', 'A vulnerable version of Vue.js is in use. Known security issues in older Vue releases may enable script injection or data exposure in reactive components.'),
    ('angular', 'A vulnerable version of Angular is detected. Unpatched Angular versions may contain XSS, template injection, or dependency vulnerabilities that expose the application to attack.'),
    ('react', 'A vulnerable version of React is in use. Outdated React versions may contain vulnerabilities in JSX rendering or event handling that could be exploited via crafted inputs.'),
    ('caa record', 'No CAA DNS record is published. Any Certificate Authority can issue TLS certificates for the domain, raising the risk of fraudulent certificate issuance and impersonation.'),
    ('subdomain_takeover', 'A subdomain DNS record points to an unclaimed or decommissioned service. An attacker can register that service and serve malicious content under the organisation\'s trusted domain.'),
    ('deprecated tls', 'Deprecated TLS 1.0 or 1.1 versions are supported. These protocols have known cryptographic weaknesses (POODLE, BEAST) that allow attackers to decrypt or tamper with encrypted traffic.'),
    ('weak cipher', 'The TLS service accepts weak cipher suites. A determined attacker can exploit these to decrypt encrypted sessions or forge authentication material.'),
    ('wildcard tls', 'A wildcard TLS certificate is in use. If the private key is ever compromised every subdomain covered by the wildcard is immediately at risk of impersonation.'),
    ('ssl dns names', 'The TLS certificate lists unexpected or sensitive internal hostnames in its Subject Alternative Names, exposing internal infrastructure to reconnaissance.'),
    ('tls version', 'An outdated TLS version is in use. Older protocols are vulnerable to downgrade attacks and weak cipher negotiation that can break encryption protecting client-server communications.'),
    ('ssl certificate issuer', 'The TLS certificate is self-signed or issued by an untrusted CA. Clients cannot verify the server\'s identity, opening the connection to man-in-the-middle attacks.'),
    ('dkim', 'DKIM signing is not configured. Without a cryptographic signature on outgoing mail, spoofed emails cannot be distinguished from legitimate ones by receiving mail servers.'),
    ('dmarc', 'A DMARC policy is absent or unenforced (p=none). Attackers can send phishing emails appearing to originate from this domain with no technical barrier preventing delivery.'),
    ('spf', 'No SPF record is published or the record is misconfigured. Unauthorised mail servers can send email purporting to be from this domain, enabling spam and phishing in the organisation\'s name.'),
    ('php', 'An end-of-life PHP version is running. No further security patches are released, leaving the application exposed to known exploits including remote code execution and SQL injection.'),
    ('wordpress', 'Outdated WordPress core, plugins, or themes are detected. Unpatched components are among the most commonly exploited vulnerabilities, enabling attackers to gain admin access or inject persistent malicious content.'),
    ('nginx', 'The web server discloses its software version. This aids attackers in identifying known vulnerabilities for the installed version and targeting exploit attempts accordingly.'),
    ('clickjacking', 'No clickjacking protection headers are present. The page can be overlaid invisibly in an iframe on an attacker-controlled site, causing users to click hidden controls.'),
    ('mx record', 'An MX record anomaly was detected. Stale or misconfigured mail exchange records can redirect inbound email to unintended servers or cause delivery failures.'),
    ('ns record', 'A nameserver record anomaly was detected. Incorrect or stale NS records can cause DNS resolution failures and create exposure to DNS hijacking or cache-poisoning attacks.'),
    ('txt record', 'Stale or unnecessary DNS TXT records are present. These can reveal details about past and current service providers, aiding attackers in mapping the technology stack.'),
    ('elementor', 'An outdated Elementor page-builder version is detected. Known vulnerabilities in older releases include stored XSS and privilege escalation exploitable by authenticated or unauthenticated users.'),
    ('cve', 'A publicly disclosed CVE was identified in software running on this asset. Published CVEs often have available exploit code, making them high-priority targets for automated attacks.'),
]

_CATEGORY_SUMMARY = {
    'ssl': 'An SSL/TLS configuration weakness was detected. Misconfigured transport security can allow attackers to intercept, decrypt, or tamper with encrypted communications.',
    'dns': 'A DNS configuration issue was identified. DNS misconfigurations can disrupt service availability, enable domain hijacking, or expose internal infrastructure to reconnaissance.',
    'email': 'An email authentication weakness was found. Gaps in SPF, DKIM, or DMARC records allow attackers to send spoofed emails that appear to originate from the organisation\'s domain.',
    'header': 'A browser security response header is missing or misconfigured. Security headers instruct browsers on content trust policies; their absence removes key defences against XSS and clickjacking.',
    'library': 'An outdated or vulnerable third-party library is in use. Unpatched libraries expose the application to known vulnerabilities that may allow script execution or data exfiltration.',
    'misconfiguration': 'A security misconfiguration was detected. Misconfigurations are easily discovered by automated scanners and are a leading initial access vector in real-world attacks.',
    'cve': 'A publicly disclosed CVE was identified in software running on this asset. Published CVEs often have available exploit code, making them high-priority targets for automated attacks.',
}


def _keyword_lookup(title, mapping):
    for keyword, text in mapping:
        if keyword in title:
            return text
    return None


def _get_groq_enhancement(finding):
    """Optionally enhance finding with AI-generated content using Groq."""
    try:
        from .remediation_enhancer import enhance_finding as groq_enhance
        enhanced = groq_enhance(finding)
        return {
            'issue_summary': enhanced.get('issue_summary'),
            'remediation_steps': enhanced.get('remediation_steps'),
            'risk_context': enhanced.get('risk_context'),
        }
    except Exception as e:
        logger.debug(f"Groq enhancement skipped: {e}")
        return None


_UNKNOWN_VERSIONS = {'', 'unknown', 'n/a', 'none', 'latest', '-'}


def _normalize_lib_key(name):
    return re.sub(r'[^a-z0-9.+-]', '', (name or '').strip().lower())


def _build_version_index(payload):
    """Map normalized component name -> {display, detected, latest}.

    Uses the *real* latest versions already fetched from the package registries
    (npm/PyPI/Packagist/WordPress) at scan time and stored on the library and
    technology checks. This is the trustworthy source for "what is the latest
    version" — never a hardcoded string or an LLM guess.
    """
    index = {}
    checks = (payload or {}).get('checks', {}) or {}
    for item in checks.get('libraries', []) or []:
        key = _normalize_lib_key(item.get('name'))
        if key:
            index[key] = {
                'display': (item.get('name') or '').strip(),
                'detected': (item.get('detected_version') or '').strip(),
                'latest': (item.get('latest_version') or '').strip(),
            }
    for item in checks.get('technologies', []) or []:
        key = _normalize_lib_key(item.get('name'))
        if key and key not in index:
            index[key] = {
                'display': (item.get('name') or '').strip(),
                'detected': (item.get('version') or '').strip(),
                'latest': (item.get('latest_version') or '').strip(),
            }
    return index


def _match_component_version(title, version_index):
    """Find the component whose name appears in the finding title (longest match wins)."""
    if not version_index:
        return None
    tl = (title or '').lower()
    best = None
    for info in version_index.values():
        disp = (info.get('display') or '').lower()
        if disp and disp in tl and (best is None or len(disp) > len(best[0])):
            best = (disp, info)
    return best[1] if best else None


def _component_from_title(title):
    """Best-effort component info parsed from a finding title like 'PHP 8.4.22'.

    Lets runtimes/servers/frameworks covered by the authoritative endoflife.date
    lookup (``_EOL_PRODUCT_MAP``) be verified against that live source even when
    they are not present in the scan's component inventory — instead of falling
    back to a static, possibly-stale claim about their support status.
    """
    if not title:
        return None
    tl = title.lower()
    # Longest product name first so 'node.js' wins over 'node'.
    for prod in sorted(_EOL_PRODUCT_MAP, key=len, reverse=True):
        m = re.search(
            r'(?<![a-z0-9])' + re.escape(prod) + r'[\s/v:_-]*([0-9]+(?:\.[0-9]+)*)',
            tl,
        )
        if not m:
            continue
        display = title[m.start():m.start() + len(prod)]
        return {'display': display, 'detected': m.group(1), 'latest': ''}
    return None


def _dynamic_component_remediation(info):
    """Build remediation text from the real detected/latest versions — no hardcoding."""
    display = info.get('display') or 'the component'
    detected = (info.get('detected') or '').strip()
    latest = (info.get('latest') or '').strip()
    has_detected = detected.lower() not in _UNKNOWN_VERSIONS
    has_latest = latest.lower() not in _UNKNOWN_VERSIONS

    if has_detected and has_latest:
        try:
            cmp = version_compare(detected, latest)
        except Exception:
            cmp = None
        if cmp is not None and cmp >= 0:
            # Installed version is current (or ahead) — the "vulnerable" flag is
            # almost certainly a signature/false-positive, so say so plainly.
            return (
                f"{display} {detected} is already the latest version published on its package "
                f"registry ({latest}), so this is most likely a false positive from "
                f"signature-based detection. Verify against the official {display} security "
                f"advisories; no upgrade is required while the installed version is current."
            )
        return (
            f"Update {display} from {detected} to {latest} — the current release published on "
            f"its official package registry. If loaded via a CDN, update the URL, pin the new "
            f"version, and add a Subresource Integrity (SRI) hash, then re-test the affected "
            f"components after upgrading."
        )
    if has_detected:
        return (
            f"{display} {detected} is in use. Check the latest release on its official package "
            f"registry (npm/PyPI/vendor) and upgrade if a newer version exists. If loaded via a "
            f"CDN, pin the version and add a Subresource Integrity (SRI) hash, then re-test."
        )
    return (
        f"Update {display} to the latest stable release published on its official package "
        f"registry. If loaded via a CDN, pin the version and add a Subresource Integrity (SRI) "
        f"hash, then re-test after upgrading."
    )


def _support_window_phrase(eol):
    """Describe a component's support window from an endoflife.date record.

    endoflife.date exposes two distinct fields that are easy to conflate:
      * ``support`` — end of *active* support (bug fixes / full support).
      * ``eol``     — end of *security* support (only security patches are issued
                      between the active-support end and this date).

    Crucially the schema is not uniform across products: each field may be an ISO
    date string, a boolean (``true`` = still supported / no announced end, ``false``
    = already ended), or absent entirely. This handles every shape so a runtime,
    framework, or JS library all get a correct, non-contradictory phrase rather than
    a bug-fix date mislabelled as "security support". Returns '' when nothing usable
    is present.
    """
    if not isinstance(eol, dict):
        return ''
    sec_until = eol.get('eol')          # security-support end
    active = eol.get('support')         # active / bug-fix support end
    today = datetime.now(timezone.utc).date().isoformat()

    parts = []
    # Security support (the is_eol==True case is handled by the caller, so here a
    # date means "still patched until then").
    if isinstance(sec_until, str):
        parts.append(f"security support until {sec_until}")
    elif sec_until is False:
        parts.append("no announced security end-of-life date")

    # Active / bug-fix support — may be a date or a boolean.
    if isinstance(active, str):
        if active < today:
            parts.append(f"active bug-fix support ended {active}")
        else:
            parts.append(f"active bug-fix support until {active}")
    elif active is True:
        parts.append("still in active bug-fix support")
    elif active is False:
        parts.append("in the security-fix-only phase (active bug-fix support has ended)")

    return '; '.join(parts)


def _component_intel(info, flagged_vulnerable=False):
    """Return (summary, remediation) for a component, grounded on real data.

    Uses the registry-fetched latest version AND a live endoflife.date lookup so
    end-of-life / supported claims are verified facts, never hardcoded guesses.
    Falls back to neutral version-only wording if the online lookup is unavailable.

    ``flagged_vulnerable`` is True when the scan's own vulnerability databases
    (CVE / Retire.js) flagged this exact version. Support-window status must never
    override that verdict — a version can be fully supported and still carry a known
    CVE — so when it is set we keep the language consistent with a real vulnerability
    rather than calling it "informational" or "not urgent".
    """
    display = info.get('display') or 'the component'
    detected = (info.get('detected') or '').strip()
    latest = (info.get('latest') or '').strip()
    has_detected = detected.lower() not in _UNKNOWN_VERSIONS
    has_latest = latest.lower() not in _UNKNOWN_VERSIONS

    eol = None
    if has_detected:
        try:
            eol = fetch_eol_status(display, detected)
        except Exception:
            eol = None

    # Prefer the package-registry 'latest'; fall back to the EOL feed's latest.
    if not has_latest and eol and eol.get('latest'):
        latest = str(eol['latest']).strip()
        has_latest = latest.lower() not in _UNKNOWN_VERSIONS

    newer_available = False
    if has_latest and has_detected:
        try:
            newer_available = version_compare(detected, latest) < 0
        except Exception:
            newer_available = False

    summary = None
    remediation = None

    if eol is not None and eol.get('is_eol') is True:
        since = f" (end-of-life since {eol['eol']})" if isinstance(eol.get('eol'), str) else " (end-of-life)"
        latest_note = f" The latest release is {latest}." if has_latest else ""
        summary = (
            f"{display} {detected}{since} no longer receives security patches, leaving the "
            f"application exposed to unpatched vulnerabilities.{latest_note}"
        )
        target = f" to {latest}" if has_latest else " to a currently supported release"
        remediation = (
            f"{display} {detected} is end-of-life — upgrade{target} and test in staging before "
            f"deploying to production."
        )
    elif flagged_vulnerable:
        # The scan's CVE / Retire.js matching flagged this version as vulnerable.
        # Report it as such regardless of support-window status, but still surface
        # the (factual) support context so the reader has the full picture.
        support_note = ''
        if eol is not None and eol.get('is_eol') is False:
            phrase = _support_window_phrase(eol)
            window = f" ({phrase})" if phrase else ""
            support_note = (
                f" The {detected} release line is still supported{window}, but that does "
                f"not clear the reported vulnerability."
            )
        latest_note = f" The latest release is {latest}." if has_latest else ""
        summary = (
            f"{display} {detected} was flagged with a known vulnerability by the scan's "
            f"vulnerability databases (CVE / Retire.js).{support_note}{latest_note}"
        )
        target = f" to {latest}" if (has_latest and newer_available) else " to a patched release"
        remediation = (
            f"Upgrade {display} from {detected}{target} to remediate the reported "
            f"vulnerability. Review the {display} security advisories for the affected "
            f"version; if loaded via a CDN, pin the new version and add a Subresource "
            f"Integrity (SRI) hash, then re-test the affected components."
        )
    elif eol is not None and eol.get('is_eol') is False:
        phrase = _support_window_phrase(eol)
        until_txt = f" ({phrase})" if phrase else ""
        # If active bug-fix support has ended the version is in the security-only
        # phase — say "security-supported" rather than implying full support.
        active = eol.get('support')
        today = datetime.now(timezone.utc).date().isoformat()
        security_only = active is False or (isinstance(active, str) and active < today)
        window_word = "security-support window" if security_only else "support window"

        if newer_available:
            summary = (
                f"{display} {detected} is still within its {window_word}{until_txt} and "
                f"receives security patches, but a newer release ({latest}) is available."
            )
            remediation = (
                f"{display} {detected} is supported, so this is not an urgent security gap, but a "
                f"newer release ({latest}) is available — upgrade to stay current and pick up "
                f"non-security fixes. If loaded via a CDN, pin the new version and add a "
                f"Subresource Integrity (SRI) hash, then re-test."
            )
        else:
            summary = (
                f"{display} {detected} is the current release and within its "
                f"{window_word}{until_txt}, so it still receives security patches. This detection "
                f"is informational — confirm against the vendor advisory before treating it as a "
                f"vulnerability."
            )
            remediation = (
                f"{display} {detected} is a currently supported release, so no upgrade is required. "
                f"Keep it patched within its release line"
                + (f" (latest {latest})." if has_latest else ".")
            )
    else:
        # No EOL data and not flagged vulnerable — neutral, version-grounded wording.
        remediation = _dynamic_component_remediation(
            {'display': display, 'detected': detected, 'latest': latest}
        )
        if has_detected and has_latest and not newer_available:
            summary = (
                f"{display} {detected} is the latest version published on its package "
                f"registry, so this detection is informational — no known vulnerability or "
                f"end-of-life status was confirmed for it."
            )
        elif has_detected and newer_available:
            summary = (
                f"{display} {detected} is in use and a newer release ({latest}) is available. "
                f"Automated matching did not confirm a known vulnerability for this version — "
                f"verify against the vendor advisory and keep it on a regular update cadence."
            )
        elif has_detected:
            summary = (
                f"{display} {detected} is in use. Automated matching did not confirm a known "
                f"vulnerability or end-of-life status for this version — verify against the "
                f"vendor advisory and keep it updated."
            )
        else:
            summary = (
                f"{display} was detected on this asset. No version was identified and automated "
                f"matching did not flag a known vulnerability — this is an informational "
                f"inventory entry, not a confirmed weakness."
            )

    return summary, remediation


def _smart_remediation(finding, version_index=None):
    """Return (recommendation_text, issue_summary, links) using rich template advice,
    optionally augmented by Groq AI when its output is substantively better.

    Always computes the keyword/category template result first. The Groq result is
    only used when it is non-trivial (long, distinct from the raw recommendation,
    not URL-only) — otherwise the template wins.
    """
    raw_title = finding.get('title') or ''
    title = raw_title.lower()
    category = (finding.get('category') or '').lower()
    risk = (finding.get('risk_rating') or 'Low').lower()
    raw = (finding.get('recommendation') or '').strip()
    raw_lower = raw.lower()

    is_generic = not raw or any(p in raw_lower for p in _GENERIC_PHRASES)
    links = re.findall(r'https?://[^\s)]+', raw)
    # also collect from evidence and explicit links field
    for src in (finding.get('evidence') or '', ' '.join(finding.get('links') or [])):
        links.extend(re.findall(r'https?://[^\s)]+', src))
    links = list(dict.fromkeys(links))  # dedupe, preserve order
    cleaned_raw = re.sub(r'https?://[^\s)]+', '', raw).strip() or None

    # --- Template-based summary + remediation (always computed) ---
    # For known libraries/technologies, prefer remediation built from the REAL
    # registry-fetched versions over any hardcoded template or raw recommendation.
    component = _match_component_version(raw_title, version_index)
    if not component:
        # Not in the inventory — but if the title names a runtime/server covered by
        # the live endoflife.date lookup, verify it against that authoritative source
        # rather than a static template.
        component = _component_from_title(raw_title)
    component_summary = None
    if component:
        # The scan's own vulnerability detection (CVE / Retire.js) is the source of
        # truth for "is this version vulnerable" — an elevated rating or an explicit
        # vulnerable/CVE signal must not be overridden by support-window status.
        flagged_vulnerable = (
            risk in ('critical', 'high', 'medium')
            or 'vulnerable' in title
            or 'retire' in title
            or 'cve' in title
            or category == 'cve'
        )
        component_summary, rec_text = _component_intel(component, flagged_vulnerable)
    elif is_generic:
        specific = _keyword_lookup(title, _TITLE_REMEDIATION)
        if not specific:
            specific = _CATEGORY_REMEDIATION.get(category)
        rec_text = specific or 'Investigate this finding, assess its exploitability in your environment, and apply the appropriate corrective control or vendor patch.'
    else:
        rec_text = cleaned_raw

    summary = component_summary or _keyword_lookup(title, _TITLE_SUMMARY)
    if not summary:
        summary = _CATEGORY_SUMMARY.get(
            category,
            'A security weakness was identified during the scan. This finding may '
            'reduce the overall security posture and should be assessed for '
            'exploitability in the context of the target environment.',
        )

    # --- Optional Groq augmentation (grounded on the real registry versions) ---
    groq_finding = finding
    if component:
        groq_finding = {
            **finding,
            'detected_version': component.get('detected') or '',
            'latest_version': component.get('latest') or '',
        }
    groq_result = _get_groq_enhancement(groq_finding)
    evidence_lower = (finding.get('evidence') or '').strip().lower()
    if groq_result:
        g_rem = (groq_result.get('remediation_steps') or '').strip()
        g_sum = (groq_result.get('issue_summary') or '').strip()
        # strip URLs Groq may have echoed back
        g_rem_clean = re.sub(r'https?://[^\s)]+', '', g_rem).strip()
        g_sum_clean = re.sub(r'https?://[^\s)]+', '', g_sum).strip()

        def _is_stub(text):
            t = text.lower()
            return (
                len(text) < 80
                or t == raw_lower
                or t == evidence_lower
                or any(p in t for p in _GENERIC_PHRASES)
            )

        # Keep our registry-grounded version remediation for known components;
        # only let Groq replace the remediation when we have no component versions.
        if g_rem_clean and not _is_stub(g_rem_clean) and not component:
            rec_text = g_rem_clean
        # Keep our verified EOL/support summary for components; let Groq fill the
        # rest only when we don't have an authoritative component fact.
        if g_sum_clean and not _is_stub(g_sum_clean) and not component_summary:
            summary = g_sum_clean

    suffix = _RISK_SUFFIX.get(risk, '')
    rec_text = rec_text.rstrip().rstrip('.') + '.' + suffix

    return rec_text, summary, links


def findings_preview(findings, preview_size=None):
    ordered = sorted(
        findings,
        key=lambda item: (risk_order(item.get('risk_rating')), item.get('title', '')),
    )
    if preview_size is None:
        preview = ordered
    else:
        preview = ordered[:preview_size]
    remaining = max(len(ordered) - len(preview), 0)
    return ordered, preview, remaining


def _root_domain(asset):
    """Best-effort apex domain extraction from an asset string or URL."""
    if not asset:
        return 'unknown'
    s = str(asset).strip().lower()
    if '://' in s:
        s = s.split('://', 1)[1]
    s = s.split('/', 1)[0].split(':', 1)[0]
    if s.startswith('www.'):
        s = s[4:]
    parts = s.split('.')
    if len(parts) >= 3 and len(parts[-1]) == 2 and parts[-2] in {
        'edu', 'co', 'gov', 'ac', 'org', 'net', 'com'
    }:
        return '.'.join(parts[-3:])
    if len(parts) >= 2:
        return '.'.join(parts[-2:])
    return s


_CVE_RE = re.compile(r'CVE-\d{4}-\d{4,}', re.IGNORECASE)
# A title that ends in a version number, e.g. "Swiper 8.4.5", "Chart.js 4.0.1",
# "simple-download-monitor 8". The part before the version is the component name.
_VERSION_TITLE_RE = re.compile(r'^(?P<name>.*\S)\s+v?\d+(?:\.\d+)*$')


def _component_base_title(title):
    """Return the clean ``Component X.Y.Z`` base for a component+version finding.

    The scanner emits a separate finding per advisory, so the same component and
    version appears many times with different titles — a generic ``(vulnerable)``
    tag plus one row per CVE / Retire.js hit. Stripping the trailing advisory
    suffix and parenthetical tag gives a single base that merges them all.

    Returns None for non-component titles (headers, DNS, SSL, email, etc.) so they
    keep their own identity and are never merged together.
    """
    if not title:
        return None
    # Drop a trailing advisory suffix such as " - CVE-2026-27212" or
    # " - Retire.js Vulnerability:".
    base = title.split(' - ', 1)[0].strip()
    # Drop a trailing parenthetical tag such as "(vulnerable)" / "(outdated)".
    base = re.sub(r'\s*\([^)]*\)\s*$', '', base).strip()
    if _VERSION_TITLE_RE.match(base):
        return base
    return None


def _aggregate_findings(findings):
    """Collapse duplicate findings into a single entry whose 'assets' list contains
    every affected asset and whose 'cve_ids' list contains every advisory.

    Component+version findings (any ``Name X.Y.Z``) are merged across all of their
    individual CVE / advisory rows into one card; everything else is merged only
    when it shares the same (title, category, risk)."""
    bucket = {}
    order = []
    for f in findings:
        title = (f.get('title') or 'Untitled finding').strip()
        category = (f.get('category') or 'Other').strip()
        risk = (f.get('risk_rating') or 'Low').strip()
        base = _component_base_title(title)
        if base:
            # One card per component+version, regardless of which advisory each
            # individual finding named.
            key = ('component', base.lower())
        else:
            key = ('title', title, category, risk)
        if key not in bucket:
            agg = dict(f)
            agg['assets'] = []
            agg['cve_ids'] = list(f.get('cve_ids') or [])
            agg['_base'] = base
            agg['_vuln'] = False
            agg['_outdated'] = False
            bucket[key] = agg
            order.append(key)
        agg = bucket[key]
        # Keep the highest severity seen across the merged advisories.
        if risk_order(risk) < risk_order(agg.get('risk_rating')):
            agg['risk_rating'] = risk
        # Track whether any merged member is an actual vulnerability / outdated tag,
        # so a benign inventory entry (e.g. a current PHP) is never relabelled.
        tl = title.lower()
        if (risk_order(risk) <= risk_order('Medium')
                or 'vulnerable' in tl or 'cve' in tl or 'retire' in tl
                or (f.get('cve_ids') or [])):
            agg['_vuln'] = True
        if 'outdated' in tl:
            agg['_outdated'] = True
        asset = (f.get('asset') or '').strip()
        if asset and asset not in agg['assets']:
            agg['assets'].append(asset)
        # Collect CVE IDs from the structured field and from the title text.
        cves = list(f.get('cve_ids') or []) + _CVE_RE.findall(title)
        for cve in cves:
            cve_u = cve.upper()
            if cve_u not in agg['cve_ids']:
                agg['cve_ids'].append(cve_u)

    # Finalise the display title for component-grouped cards: keep it clean, and
    # only add a "(vulnerable)"/"(outdated)" tag when the group actually warrants it.
    result = []
    for k in order:
        agg = bucket[k]
        base = agg.pop('_base', None)
        is_vuln = agg.pop('_vuln', False)
        is_outdated = agg.pop('_outdated', False)
        if base:
            if is_vuln:
                agg['title'] = f"{base} (vulnerable)"
            elif is_outdated:
                agg['title'] = f"{base} (outdated)"
            else:
                agg['title'] = base
        result.append(agg)
    return result


def _affects_text(assets):
    if not assets:
        return 'Affects: N/A'
    if len(assets) == 1:
        return f"Affects: {assets[0]}"
    return f"Affects {len(assets)} subdomains: " + ', '.join(assets)


def _group_by_domain(aggregated):
    """Group aggregated findings by root domain of their first asset.
    Returns list of (domain, [findings]) preserving discovery order."""
    groups = {}
    order = []
    for f in aggregated:
        assets = f.get('assets') or []
        domain = _root_domain(assets[0]) if assets else 'unknown'
        if domain not in groups:
            groups[domain] = []
            order.append(domain)
        groups[domain].append(f)
    return [(d, groups[d]) for d in order]


def _split_by_severity(aggregated):
    """Return (top, rest): Critical+High first, Medium+Low second.
    Each list is already in the input order (caller should sort first)."""
    top, rest = [], []
    for f in aggregated:
        risk = (f.get('risk_rating') or 'Low').strip().lower()
        if risk in ('critical', 'high'):
            top.append(f)
        else:
            rest.append(f)
    return top, rest


def _build_check_sections(payload):
    checks = payload.get('checks', {})
    return [
        ('Library Checks', checks.get('libraries', [])),
        ('Technology Checks', checks.get('technologies', [])),
        ('SSL/TLS Checks', checks.get('ssl', [])),
        ('Email Checks', checks.get('email', [])),
        ('Header Checks', checks.get('headers', [])),
        ('DNS Checks', checks.get('dns', [])),
    ]


def _check_section_summary(section_name, grouped_items):
    """Return 1-3 short HTML paragraphs summarising what was found in a check
    section, what it means, and what to fix. Tailored per control family."""
    if not grouped_items:
        return []

    name = (section_name or '').lower()
    total = len(grouped_items)

    def has(*keywords):
        kws = [k.lower() for k in keywords]
        return [g for g in grouped_items if any(k in (g.get('name') or '').lower() for k in kws)
                or any(k in (g.get('status') or '').lower() for k in kws)]

    def issues():
        out = []
        for g in grouped_items:
            status = (g.get('status') or '').lower()
            risk = (g.get('risk_rating') or '').lower()
            if any(t in status for t in ('missing', 'vulnerable', 'outdated', 'expired',
                                          'mismatched', 'unknown', 'fail', 'not set')):
                out.append(g)
            elif risk in ('critical', 'high', 'medium'):
                out.append(g)
        return out

    issue_items = issues()
    issue_count = len(issue_items)

    if 'ssl' in name or 'tls' in name:
        if not issue_count:
            return [
                f"<b>{total} SSL/TLS check(s) recorded.</b> No high-severity protocol "
                "or certificate problems were flagged automatically, but review the "
                "items listed below and remediate any deprecated protocols, weak "
                "ciphers, or untrusted/expired certificates. Re-test with an SSL scanner.",
            ]
        protocol = has('protocol', 'tls version', 'deprecated')
        certs = has('certificate', 'expired', 'mismatched', 'self-signed', 'ssl certificate issuer')
        hsts = has('hsts')
        sans = has('ssl dns names', 'san')
        parts = []
        if protocol:
            parts.append(f"deprecated TLS protocol versions on {len(protocol)} asset(s)")
        if certs:
            parts.append(f"certificate problems (expired, mismatched, or untrusted issuer) on {len(certs)} asset(s)")
        if hsts:
            parts.append(f"weak or missing HSTS on {len(hsts)} asset(s)")
        if sans:
            parts.append(f"exposed internal hostnames in TLS SAN fields on {len(sans)} asset(s)")
        breakdown = '; '.join(parts) if parts else f"{issue_count} configuration issue(s)"
        return [
            f"<b>{issue_count} SSL/TLS issue(s) need attention</b> &mdash; {breakdown}.",
            "<b>What this means:</b> Misconfigured transport security lets attackers "
            "intercept, downgrade, or impersonate encrypted sessions, and exposes "
            "internal infrastructure through certificate metadata.",
            "<b>What to fix:</b> Enforce TLS 1.2 minimum (prefer 1.3) and disable "
            "SSLv3/TLS 1.0/1.1 at every load balancer and origin server. Replace "
            "self-signed, expired, or mismatched certificates with publicly-trusted "
            "ones via ACME automation. Enable HSTS with "
            "<font face='Courier'>max-age=31536000; includeSubDomains; preload</font>. "
            "Remove unused or sensitive hostnames from certificate SANs and re-test "
            "with SSL Labs after every change.",
        ]

    if 'email' in name:
        # SPF / DKIM / DMARC are all required. A row that's missing entirely is
        # itself a gap — not just rows marked failing.
        def _has_record(record_type):
            rt = record_type.lower()
            for g in grouped_items:
                n = (g.get('name') or '').lower()
                s = (g.get('status') or '').lower()
                if rt in n and 'not set' not in s and 'missing' not in s \
                        and 'fail' not in s and 'unknown' not in s:
                    return True
            return False

        spf_ok = _has_record('spf')
        dkim_ok = _has_record('dkim')
        dmarc_ok = _has_record('dmarc')
        gaps = []
        if not spf_ok:
            gaps.append("SPF is missing or misconfigured")
        if not dkim_ok:
            gaps.append("DKIM signing is not configured")
        if not dmarc_ok:
            gaps.append("DMARC is missing or unenforced (p=none)")

        if not gaps:
            return [
                f"<b>All three email-auth controls are configured.</b> SPF, DKIM, "
                "and DMARC are in place. Maintain DMARC aggregate-report monitoring "
                "and progress policy toward <font face='Courier'>p=reject</font>.",
            ]

        configured = []
        if spf_ok:
            configured.append("SPF")
        if dkim_ok:
            configured.append("DKIM")
        if dmarc_ok:
            configured.append("DMARC")
        configured_phrase = (
            f" Currently configured: {', '.join(configured)}." if configured else ""
        )

        return [
            f"<b>{len(gaps)} of 3 email-auth controls need to be set</b> &mdash; "
            f"{'; '.join(gaps)}.{configured_phrase}",
            "<b>What this means:</b> SPF, DKIM, and DMARC must all be configured "
            "and aligned for receiving mail servers to reliably reject spoofed "
            "email. Any one of them missing leaves a path for attackers to send "
            "phishing email that appears to originate from your domain with no "
            "technical barrier preventing inbox delivery &mdash; the leading vector "
            "for business email compromise (BEC).",
            "<b>What to fix:</b> (1) Publish an SPF TXT record listing every "
            "authorised sending IP/host and ending with "
            "<font face='Courier'>-all</font>. "
            "(2) Generate a 2048-bit DKIM key pair, publish the public key in DNS "
            "under a unique selector, and configure the MTA to sign all outgoing "
            "mail. (3) Publish a DMARC record starting at "
            "<font face='Courier'>p=none</font> with an "
            "<font face='Courier'>rua</font> reporting address; analyse the "
            "aggregate reports for 2&ndash;4 weeks, fix any SPF/DKIM alignment "
            "failures, then progress to <font face='Courier'>p=quarantine</font> "
            "and finally <font face='Courier'>p=reject</font>.",
        ]

    if 'dns' in name:
        if not issue_count:
            return [
                f"<b>{total} DNS check(s) recorded.</b> No high-severity DNS issues "
                "were flagged automatically, but review the entries below — including "
                "any subdomain-takeover candidates, missing CAA records, or stale "
                "TXT/CNAME records — and confirm each is expected.",
            ]
        takeover = has('subdomain_takeover', 'takeover')
        caa = has('caa')
        txt = has('txt record')
        parts = []
        if takeover:
            parts.append(f"{len(takeover)} subdomain(s) flagged as takeover candidates")
        if caa:
            parts.append(f"{len(caa)} domain(s) missing a CAA record")
        if txt:
            parts.append(f"stale or excessive TXT records on {len(txt)} asset(s)")
        breakdown = '; '.join(parts) if parts else f"{issue_count} DNS issue(s)"
        return [
            f"<b>{issue_count} DNS issue(s) need attention</b> &mdash; {breakdown}.",
            "<b>What this means:</b> Stale CNAMEs pointing at unclaimed cloud "
            "resources allow attackers to register the target and serve content "
            "from your trusted domain. Missing CAA records let any CA issue "
            "certificates for the domain. Old TXT records leak details about "
            "decommissioned vendors and infrastructure.",
            "<b>What to fix:</b> Audit every CNAME and A record &mdash; remove or "
            "update any pointing at decommissioned services, and add DNS cleanup "
            "to your cloud deprovisioning checklist. Publish a CAA record "
            "(<font face='Courier'>0 issue \"letsencrypt.org\"</font> or your "
            "chosen CA) on the apex domain. Remove TXT records for retired "
            "services and rotate any verification tokens older than 12 months.",
        ]

    if 'header' in name:
        if not issue_count:
            return [
                f"<b>{total} header checks recorded.</b> Core browser security "
                "headers are in place. Continue to monitor for regressions.",
            ]
        return [
            f"<b>{issue_count} missing or misconfigured browser security header(s).</b>",
            "<b>What this means:</b> Security headers are the browser-side defence "
            "layer against XSS, clickjacking, MIME sniffing, and information leakage "
            "via referrer. Without them, even small application bugs become "
            "exploitable, and the site lacks defence-in-depth.",
            "<b>What to fix:</b> Add the following on every HTTP response in your "
            "web-server or reverse-proxy configuration: "
            "<font face='Courier'>Content-Security-Policy</font> (start strict, "
            "iterate via report-only first), "
            "<font face='Courier'>Strict-Transport-Security</font>, "
            "<font face='Courier'>X-Content-Type-Options: nosniff</font>, "
            "<font face='Courier'>X-Frame-Options: DENY</font>, "
            "<font face='Courier'>Referrer-Policy: strict-origin-when-cross-origin</font>, "
            "and a <font face='Courier'>Permissions-Policy</font> that disables "
            "unused browser APIs. Validate with securityheaders.com.",
        ]

    if 'library' in name or 'technolog' in name:
        vuln = [g for g in grouped_items if 'vulnerable' in (g.get('status') or '').lower()]
        outdated = [g for g in grouped_items if 'outdated' in (g.get('status') or '').lower()]
        if not (vuln or outdated):
            return [
                f"<b>{total} component(s) inventoried.</b> None were flagged vulnerable "
                "or outdated by automated version matching, but detection is best-effort "
                "— verify any end-of-life runtimes or frameworks against their official "
                "supported-version lists and keep dependencies on a regular update cadence.",
            ]
        return [
            f"<b>{len(vuln)} vulnerable</b> and <b>{len(outdated)} outdated</b> "
            f"third-party component(s) detected across the technology stack.",
            "<b>What this means:</b> Public vulnerability databases (CVE, GHSA, "
            "Retire.js) flag known issues against the versions running. Attackers "
            "routinely automate exploitation of these because the proof-of-concept "
            "code is published. Outdated components without active CVEs still drift "
            "out of security-patch support.",
            "<b>What to fix:</b> Upgrade each vulnerable component to its current "
            "stable release; review changelogs for breaking changes and run the "
            "test suite before deploying. Where a library is loaded from a CDN, "
            "pin a version and add a Subresource Integrity (SRI) hash. Add a "
            "dependency scanner (npm audit, Dependabot, Renovate) to CI so future "
            "drift is caught at PR time.",
        ]

    # Generic fallback
    if issue_count:
        return [
            f"<b>{issue_count} item(s) flagged in {section_name}.</b> Review the "
            "table below and address each finding according to the recommendations "
            "in the per-finding cards earlier in this report.",
        ]
    return []


def _group_check_items(items):
    """Collapse check rows by (name, status, risk) and collect all unique assets into one entry."""
    grouped = {}
    order = []
    for item in items:
        name = str(item.get('name') or 'Unknown').strip()
        status = str(
            item.get('status') or item.get('detected_version') or
            item.get('version') or item.get('details') or 'Recorded'
        ).strip()
        risk = str(item.get('risk_rating') or 'Low').strip()
        asset = str(item.get('asset') or 'N/A').strip()
        key = (name, status, risk)
        if key not in grouped:
            grouped[key] = {'name': name, 'status': status, 'risk_rating': risk, 'assets': []}
            order.append(key)
        if asset and asset not in grouped[key]['assets']:
            grouped[key]['assets'].append(asset)
    return [grouped[k] for k in order]


def generate_pdf_report(payload):
    """Generate styled PDF report from payload data."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError:
        return None, 'PDF export dependency missing. Install reportlab.'

    C = PDF_STYLE['colors']
    F = PDF_STYLE['fonts']
    SP = PDF_STYLE['spacing']

    summary = payload['summary']
    ordered, preview, _ = findings_preview(payload['findings'])

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=PDF_STYLE['page']['left_margin_mm'] * mm,
        rightMargin=PDF_STYLE['page']['right_margin_mm'] * mm,
        topMargin=PDF_STYLE['page']['top_margin_mm'] * mm,
        bottomMargin=PDF_STYLE['page']['bottom_margin_mm'] * mm,
    )
    styles = getSampleStyleSheet()

    # --- Typography styles ---
    title_style = ParagraphStyle(
        'ReportTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=F['title_size'],
        leading=24,
        textColor=colors.HexColor(C['text_primary']),
        spaceAfter=4,
    )
    meta_style = ParagraphStyle(
        'Meta',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=F['small_size'],
        leading=12,
        textColor=colors.HexColor(C['text_muted']),
        spaceAfter=0,
    )
    section_label_style = ParagraphStyle(
        'SectionLabel',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=F['section_size'],
        leading=11,
        textColor=colors.HexColor(C['section_label']),
        spaceBefore=SP['section_gap_mm'] * mm,
        spaceAfter=3,
    )
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=F['body_size'],
        leading=13,
        textColor=colors.HexColor(C['text_primary']),
    )
    muted_body_style = ParagraphStyle(
        'MutedBody',
        parent=body_style,
        textColor=colors.HexColor(C['text_secondary']),
    )
    hint_style = ParagraphStyle(
        'Hint',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=F['small_size'],
        leading=11,
        textColor=colors.HexColor(C['text_hint']),
    )
    action_label_style = ParagraphStyle(
        'ActionLabel',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=F['meta_label_size'],
        leading=10,
        textColor=colors.HexColor(C['text_muted']),
        spaceAfter=3,
    )
    link_style = ParagraphStyle(
        'Link',
        parent=body_style,
        textColor=colors.HexColor(C['link']),
        fontSize=F['small_size'],
    )
    small_style = ParagraphStyle(
        'Small',
        parent=body_style,
        fontSize=8,
        leading=10,
        textColor=colors.HexColor(C['text_muted']),
    )

    story = []

    # === HEADER ===
    org_name = summary.get('org_name') or summary.get('domain_name') or 'Security'
    report_title = f"{org_name} Vulnerability Remediation Report"
    story.append(Paragraph(report_title, title_style))
    status_color = C['status_completed']
    story.append(
        Paragraph(
            f"Generated {_format_generated_at(payload['generated_at'])}  \u00b7  Scan #{summary['scan_id']}  \u00b7  "
            f"{summary['scan_type']}  \u00b7  "
            f"<font color='{status_color}'>{summary['status']}</font>  \u00b7  "
            f"{summary['total_assets_scanned']} asset checked",
            meta_style,
        )
    )
    story.append(Spacer(1, 8))

    # === STATS GRID ===
    stat_entries = [
        ('TOTAL', str(summary['total_findings']), C['text_primary']),
        ('CRITICAL', str(summary['critical_findings']), C['severity']['Critical']['text']),
        ('HIGH', str(summary['high_findings']), C['severity']['High']['text']),
        ('MEDIUM', str(summary['medium_findings']), C['severity']['Medium']['text']),
        ('LOW', str(summary['low_findings']), C['severity']['Low']['text']),
        ('ASSETS', str(summary['total_assets_scanned']), C['text_primary']),
    ]
    col_w = 29 * mm
    label_row = [
        Paragraph(f"<font color='{C['stat_label']}'><b>{lbl}</b></font>",
                  ParagraphStyle('SL', parent=styles['Normal'], fontName='Helvetica-Bold',
                                 fontSize=F['meta_label_size'], leading=10))
        for lbl, _, _ in stat_entries
    ]
    value_row = [
        Paragraph(f"<font color='{col}'><b>{val}</b></font>",
                  ParagraphStyle('SV', parent=styles['Normal'], fontName='Helvetica-Bold',
                                 fontSize=F['meta_value_size'], leading=18))
        for _, val, col in stat_entries
    ]
    stats_table = Table([label_row, value_row], colWidths=[col_w] * 6)
    stats_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(C['stat_bg'])),
        ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor(C['stat_border'])),
        ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.HexColor(C['border'])),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.HexColor(C['stat_bg'])]),
    ]))
    story.append(stats_table)
    story.append(Spacer(1, 10))

    # === EXECUTIVE SUMMARY ===
    story.append(Paragraph('EXECUTIVE SUMMARY', section_label_style))

    domain_label = summary.get('domain_name') or 'the target domain'
    total = summary['total_findings']
    crit = summary['critical_findings']
    high = summary['high_findings']
    med = summary['medium_findings']
    low = summary['low_findings']
    assets_count = summary['total_assets_scanned']

    # Risk areas derived from findings categories
    categories = []
    for f in payload['findings']:
        cat = (f.get('category') or '').strip()
        if cat and cat not in categories:
            categories.append(cat)
    if len(categories) <= 4:
        risk_areas = ', '.join(categories) if categories else 'multiple control families'
    else:
        risk_areas = ', '.join(categories[:4]) + f", and {len(categories) - 4} more"

    if crit > 0:
        posture, posture_action = 'Critical', (
            'Immediate action is required to address the critical vulnerabilities '
            'identified before they can be exploited.'
        )
    elif high > 0:
        posture, posture_action = 'High', (
            'High-severity issues should be remediated promptly to reduce the '
            'organisation\'s exposure to known exploit techniques.'
        )
    elif med > 0:
        posture, posture_action = 'Medium', (
            'Schedule remediation of the medium-severity findings within the '
            'next sprint and maintain ongoing monitoring.'
        )
    else:
        posture, posture_action = 'Low', (
            'No critical or high findings were identified. Continue routine '
            'monitoring and address low-severity items during regular hardening.'
        )

    exec_lines = [
        f"This report presents the findings of a security scan conducted against "
        f"<b>{domain_label}</b> covering {assets_count} assets. The scan identified "
        f"{total} security findings: {crit} Critical, {high} High, {med} Medium, "
        f"and {low} Low severity.",
        f"Risk areas identified: {risk_areas}.",
        f"The overall risk posture is <b>{posture}</b>. {posture_action}",
        "Remediation priority: address Critical and High findings first, validate "
        "each fix in a staging environment before production deployment, then "
        "re-scan to confirm closure.",
    ]
    exec_rows = [[Paragraph(line, muted_body_style)] for line in exec_lines]
    exec_table = Table(exec_rows, colWidths=[174 * mm])
    exec_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(C['surface_secondary'])),
        ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor(C['border'])),
        ('INNERGRID', (0, 0), (-1, -1), 0, colors.white),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ROUNDEDCORNERS', [6, 6, 6, 6]),
    ]))
    story.append(exec_table)
    story.append(Spacer(1, 10))

    # === FINDINGS ===
    story.append(Paragraph(REPORT_TEXT['findings_title'], section_label_style))

    if not ordered:
        story.append(Paragraph(REPORT_TEXT['no_findings'], body_style))
    else:
        aggregated = _aggregate_findings(preview)
        aggregated.sort(
            key=lambda f: (risk_order(f.get('risk_rating')), f.get('title', ''))
        )
        top_findings, rest_findings = _split_by_severity(aggregated)
        domain_groups = _group_by_domain(top_findings)

        about_label_style = ParagraphStyle(
            'AboutLabel', parent=styles['Normal'], fontName='Helvetica-Bold',
            fontSize=F['meta_label_size'], leading=10,
            textColor=colors.HexColor(C['text_muted']), spaceAfter=3,
        )
        affects_style = ParagraphStyle(
            'Affects', parent=styles['Normal'], fontName='Helvetica',
            fontSize=F['small_size'], leading=12,
            textColor=colors.HexColor(C['text_secondary']),
        )
        domain_header_style = ParagraphStyle(
            'DomainHeader', parent=styles['Normal'], fontName='Helvetica-Bold',
            fontSize=F['body_size'], leading=14,
            textColor=colors.HexColor(C['text_primary']),
        )
        domain_count_style = ParagraphStyle(
            'DomainCount', parent=styles['Normal'], fontName='Helvetica',
            fontSize=F['small_size'], leading=14, alignment=2,
            textColor=colors.HexColor(C['text_muted']),
        )

        version_index = _build_version_index(payload)
        running_index = 0
        for domain, dom_findings in domain_groups:
            # Domain header bar
            dom_header = Table(
                [[
                    Paragraph(domain, domain_header_style),
                    Paragraph(f"{len(dom_findings)} issues", domain_count_style),
                ]],
                colWidths=[120 * mm, 54 * mm],
            )
            dom_header.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(C['card_header_bg'])),
                ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor(C['card_border'])),
                ('LEFTPADDING', (0, 0), (-1, -1), 10),
                ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(dom_header)
            story.append(Spacer(1, 4))

            for finding in dom_findings:
                running_index += 1
                index = running_index
                risk = finding.get('risk_rating') or 'Low'
                sev = C['severity'].get(risk, C['severity']['default'])
                remediation_text, issue_summary, remediation_links = _smart_remediation(finding, version_index)

                badge = Paragraph(
                    f"<font color='{sev['text']}'><b>{risk}</b></font>",
                    ParagraphStyle('Badge', parent=styles['Normal'], fontName='Helvetica-Bold',
                                   fontSize=F['small_size'], leading=11, alignment=1)
                )
                num_cell = Paragraph(
                    f"<font color='{C['text_muted']}'><b>{index}</b></font>",
                    ParagraphStyle('Num', parent=styles['Normal'], fontName='Helvetica-Bold',
                                   fontSize=F['small_size'], leading=11, alignment=1)
                )
                title_cell = Paragraph(
                    f"<b>{finding.get('title') or 'Untitled finding'}</b>",
                    ParagraphStyle('FindingTitle', parent=styles['Normal'], fontName='Helvetica-Bold',
                                   fontSize=F['body_size'], leading=13,
                                   textColor=colors.HexColor(C['text_primary']))
                )

                affects = _affects_text(finding.get('assets') or [])
                affects_cell = Paragraph(affects, affects_style)

                # ABOUT sub-table (left column)
                about_text = (
                    issue_summary or finding.get('evidence')
                    or 'Finding details not available.'
                )
                # List every advisory merged into this component+version card so
                # no CVE/Retire.js detail is lost when the rows are collapsed.
                cve_ids = finding.get('cve_ids') or []
                if cve_ids:
                    shown = ', '.join(cve_ids[:10])
                    if len(cve_ids) > 10:
                        shown += f", +{len(cve_ids) - 10} more"
                    about_text = (
                        f"{about_text}<br/><br/>"
                        f"<b>Associated advisories:</b> {shown}"
                    )
                about_body = Paragraph(about_text, muted_body_style)
                about_sub = Table(
                    [[Paragraph('ABOUT THIS ISSUE', about_label_style)], [about_body]],
                    colWidths=[70 * mm],
                )
                about_sub.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(C['card_bg'])),
                    ('LEFTPADDING', (0, 0), (-1, -1), 8),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ('BOX', (0, 0), (-1, -1), 0, colors.white),
                    ('INNERGRID', (0, 0), (-1, -1), 0, colors.white),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))

                # REMEDIATION sub-table (right column)
                if remediation_text == 'No remediation provided.':
                    remediation_cell = Paragraph(remediation_text, hint_style)
                else:
                    remediation_cell = Paragraph(remediation_text, muted_body_style)
                rem_rows = [
                    [Paragraph('REMEDIATION', action_label_style)],
                    [remediation_cell],
                ]
                # External links intentionally not displayed — they are used only as
                # AI context for richer remediation generation.
                rem_sub = Table(rem_rows, colWidths=[74 * mm])
                rem_sub.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(C['card_bg'])),
                    ('LEFTPADDING', (0, 0), (-1, -1), 8),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ('BOX', (0, 0), (-1, -1), 0, colors.white),
                    ('INNERGRID', (0, 0), (-1, -1), 0, colors.white),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))

                outer_rows = [
                    # Row 0: header (number | title | badge)
                    [num_cell, title_cell, badge],
                    # Row 1: affects line
                    ['', affects_cell, ''],
                    # Row 2: about | remediation two-column body
                    ['', Table([[about_sub, rem_sub]], colWidths=[70 * mm, 74 * mm]), ''],
                ]

                card = Table(outer_rows, colWidths=[9 * mm, 144 * mm, 21 * mm])
                card.setStyle(TableStyle([
                    ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor(C['card_border'])),
                    ('ROUNDEDCORNERS', [6, 6, 6, 6]),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(C['card_header_bg'])),
                    ('BACKGROUND', (0, 1), (-1, 2), colors.HexColor(C['card_bg'])),
                    ('LINEBELOW', (0, 0), (-1, 0), 0.5, colors.HexColor(C['border'])),
                    ('LINEBELOW', (0, 1), (-1, 1), 0.5, colors.HexColor(C['border'])),
                    ('BACKGROUND', (2, 0), (2, 0), colors.HexColor(sev['bg'])),
                    ('BOX', (2, 0), (2, 0), 0.5, colors.HexColor(sev['border'])),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('ALIGN', (0, 0), (0, -1), 'CENTER'),
                    ('ALIGN', (2, 0), (2, 0), 'CENTER'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 8),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                    ('TOPPADDING', (0, 0), (-1, -1), 7),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
                    ('TOPPADDING', (0, 1), (-1, 2), 4),
                    ('BOTTOMPADDING', (0, 1), (-1, 2), 4),
                    ('LEFTPADDING', (1, 1), (1, 2), 8),
                    ('LEFTPADDING', (0, 1), (0, 2), 0),
                    ('LEFTPADDING', (2, 1), (2, 2), 0),
                ]))

                story.append(card)
                story.append(Spacer(1, SP['card_gap_mm'] * mm))

        # === Additional findings — Medium & Low ===
        if rest_findings:
            med_count = sum(1 for f in rest_findings
                            if (f.get('risk_rating') or '').lower() == 'medium')
            low_count = len(rest_findings) - med_count

            # Group by category to summarise themes
            cat_buckets = {}
            for f in rest_findings:
                cat = (f.get('category') or 'Other').strip() or 'Other'
                cat_buckets.setdefault(cat, 0)
                cat_buckets[cat] += 1
            top_cats = sorted(cat_buckets.items(), key=lambda x: -x[1])[:4]
            cat_phrase = ', '.join(f"{name} ({count})" for name, count in top_cats)

            story.append(Spacer(1, 8))
            story.append(Paragraph(
                f"Additional findings &mdash; Medium &amp; Low ({len(rest_findings)})",
                ParagraphStyle('RestHeader', parent=styles['Normal'], fontName='Helvetica-Bold',
                               fontSize=F['section_size'], leading=12,
                               textColor=colors.HexColor(C['section_label']),
                               spaceAfter=4)
            ))

            # Brief intro summary explaining what these findings mean + what to do
            intro_lines = [
                f"This section lists {len(rest_findings)} additional findings "
                f"({med_count} Medium, {low_count} Low) that do not pose immediate "
                f"critical risk but should be addressed during routine hardening. "
                f"The bulk of these relate to {cat_phrase}.",
                "<b>What this means:</b> These are weaknesses such as missing "
                "defence-in-depth headers, outdated libraries with known but lower-impact "
                "issues, stale DNS entries, and TLS configuration drift. Individually each "
                "is low impact, but together they expand the attack surface and make "
                "exploitation of higher-severity issues easier.",
                "<b>Recommended action:</b> Triage the table below by category, batch "
                "fixes by control family (headers, libraries, DNS, TLS), and schedule "
                "remediation across the next 1&ndash;2 sprints. Validate each fix in "
                "staging and re-scan to confirm closure.",
            ]
            intro_rows = [[Paragraph(line, muted_body_style)] for line in intro_lines]
            intro_table = Table(intro_rows, colWidths=[174 * mm])
            intro_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(C['surface_secondary'])),
                ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor(C['border'])),
                ('INNERGRID', (0, 0), (-1, -1), 0, colors.white),
                ('LEFTPADDING', (0, 0), (-1, -1), 10),
                ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 5),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ROUNDEDCORNERS', [6, 6, 6, 6]),
            ]))
            story.append(intro_table)
            story.append(Spacer(1, 6))

            # Compact table — one row per finding with title, short summary, sev, assets
            rest_rows = [[
                Paragraph('<b>Finding</b>', small_style),
                Paragraph('<b>Sev</b>', small_style),
                Paragraph('<b>Affected Assets</b>', small_style),
            ]]
            for f in rest_findings:
                title = f.get('title') or 'Untitled finding'
                _rem, finding_summary, _links = _smart_remediation(f, version_index)
                summary_text = (finding_summary or '').strip()
                finding_html = f"<b>{title}</b>"
                if summary_text:
                    finding_html += (
                        f"<br/><font size='7' color='{C['text_muted']}'>"
                        f"{summary_text}</font>"
                    )
                assets_list = f.get('assets') or []
                if len(assets_list) > 3:
                    assets_str = ', '.join(assets_list[:3]) + f" +{len(assets_list) - 3} more"
                else:
                    assets_str = ', '.join(assets_list)
                rest_rows.append([
                    Paragraph(finding_html, body_style),
                    Paragraph(f.get('risk_rating') or 'Low', muted_body_style),
                    Paragraph(assets_str, muted_body_style),
                ])
            rest_table = Table(rest_rows, colWidths=[90 * mm, 16 * mm, 68 * mm])
            rest_table.setStyle(TableStyle([
                ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor(C['border'])),
                ('INNERGRID', (0, 0), (-1, -1), 0.4, colors.HexColor(C['border_light'])),
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(C['surface_secondary'])),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))
            story.append(rest_table)

        story.append(Spacer(1, 4))
        story.append(
            Paragraph(
                f"Showing {len(aggregated)} unique issue types across "
                f"{len(domain_groups) or 1} domain"
                f"{'s' if (len(domain_groups) or 1) != 1 else ''}.",
                ParagraphStyle('Footer', parent=styles['Normal'], fontName='Helvetica',
                               fontSize=F['small_size'], leading=11,
                               textColor=colors.HexColor(C['text_hint']))
            )
        )

        # === VULNERABILITY SUMMARY ===
        story.append(Spacer(1, 10))
        story.append(Paragraph('VULNERABILITY SUMMARY', section_label_style))
        cat_summary_text = ', '.join(categories[:4]) if categories else 'multiple control families'
        if len(categories) > 4:
            cat_summary_text += ', and others'
        ch_count = summary['critical_findings'] + summary['high_findings']
        story.append(Paragraph(
            f"The scan identified {summary['total_findings']} vulnerability findings, "
            f"with {ch_count} rated Critical or High. These findings span categories "
            f"such as {cat_summary_text}.",
            muted_body_style,
        ))
        story.append(Paragraph(
            "Recommended action: Prioritise remediation of Critical and High findings, "
            "validate fixes in staging, then perform a follow-up Vulnerability "
            "Assessment and Penetration Test (VAPT) and re-scan to confirm closure.",
            muted_body_style,
        ))

    # === ADDITIONAL CHECKS ===
    story.append(Spacer(1, 10))
    story.append(Paragraph(REPORT_TEXT['checks_title'], section_label_style))
    for section_name, section_items in _build_check_sections(payload):
        story.append(Paragraph(section_name, body_style))
        if not section_items:
            story.append(Paragraph('No records for this section.', hint_style))
            story.append(Spacer(1, 4))
            continue

        grouped_items = _group_check_items(section_items)

        # Per-section summary: what was found + what to fix
        summary_lines = _check_section_summary(section_name, grouped_items)
        if summary_lines:
            sum_rows = [[Paragraph(line, muted_body_style)] for line in summary_lines]
            sum_table = Table(sum_rows, colWidths=[174 * mm])
            sum_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(C['surface_secondary'])),
                ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor(C['border'])),
                ('INNERGRID', (0, 0), (-1, -1), 0, colors.white),
                ('LEFTPADDING', (0, 0), (-1, -1), 10),
                ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 5),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ROUNDEDCORNERS', [6, 6, 6, 6]),
            ]))
            story.append(sum_table)
            story.append(Spacer(1, 4))

        check_rows = [[
            Paragraph('<b>Name</b>', small_style),
            Paragraph('<b>Status / Version</b>', small_style),
            Paragraph('<b>Affected Assets</b>', small_style),
        ]]
        for g_item in grouped_items:
            status_risk = g_item['status']
            if g_item['risk_rating'] and g_item['risk_rating'].lower() not in ('low', 'recorded', ''):
                status_risk = f"{g_item['status']}  [{g_item['risk_rating']}]"
            check_rows.append([
                Paragraph(g_item['name'], body_style),
                Paragraph(status_risk, muted_body_style),
                Paragraph(', '.join(g_item['assets']), muted_body_style),
            ])

        check_table = Table(check_rows, colWidths=[55 * mm, 55 * mm, 65 * mm])
        check_table.setStyle(TableStyle([
            ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor(C['border'])),
            ('INNERGRID', (0, 0), (-1, -1), 0.4, colors.HexColor(C['border_light'])),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(C['surface_secondary'])),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(check_table)
        story.append(Spacer(1, 6))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue(), None


def generate_docx_report(payload):
    """Generate styled DOCX report from payload data."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Mm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return None, 'DOCX export dependency missing. Install python-docx.'

    C = PDF_STYLE['colors']

    def hex_to_rgb(hex_color):
        h = hex_color.lstrip('#')
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color.lstrip('#'))
        tcPr.append(shd)

    def add_paragraph_with_color(doc_or_cell, text, font_size=9, bold=False, color_hex=None, italic=False):
        if hasattr(doc_or_cell, 'add_paragraph'):
            p = doc_or_cell.add_paragraph()
        else:
            p = doc_or_cell.paragraphs[0] if doc_or_cell.paragraphs else doc_or_cell.add_paragraph()
            p.clear()
        run = p.add_run(text)
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        if color_hex:
            r, g, b = hex_to_rgb(color_hex)
            run.font.color.rgb = RGBColor(r, g, b)
        return p

    summary = payload['summary']
    ordered, preview, _ = findings_preview(payload['findings'])
    document = Document()

    # Narrow margins
    for section in document.sections:
        section.left_margin = Mm(14)
        section.right_margin = Mm(14)
        section.top_margin = Mm(14)
        section.bottom_margin = Mm(14)

    # === TITLE ===
    title_p = document.add_paragraph()
    title_run = title_p.add_run(REPORT_TEXT['title'])
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    r, g, b = hex_to_rgb(C['text_primary'])
    title_run.font.color.rgb = RGBColor(r, g, b)

    # === META LINE ===
    meta_p = document.add_paragraph()
    meta_text = (
        f"Generated {payload['generated_at']}  \u00b7  Scan #{summary['scan_id']}  \u00b7  "
        f"{summary['scan_type']}  \u00b7  {summary['status']}  \u00b7  "
        f"{summary['total_assets_scanned']} asset checked"
    )
    meta_run = meta_p.add_run(meta_text)
    meta_run.font.size = Pt(8)
    r, g, b = hex_to_rgb(C['text_muted'])
    meta_run.font.color.rgb = RGBColor(r, g, b)

    document.add_paragraph()

    # === STATS TABLE ===
    stats = [
        ('TOTAL', str(summary['total_findings']), C['text_primary']),
        ('CRITICAL', str(summary['critical_findings']), C['severity']['Critical']['text']),
        ('HIGH', str(summary['high_findings']), C['severity']['High']['text']),
        ('MEDIUM', str(summary['medium_findings']), C['severity']['Medium']['text']),
        ('LOW', str(summary['low_findings']), C['severity']['Low']['text']),
        ('ASSETS', str(summary['total_assets_scanned']), C['text_primary']),
    ]
    stats_table = document.add_table(rows=2, cols=6)
    stats_table.style = 'Table Grid'
    for idx, (label, value, color) in enumerate(stats):
        label_cell = stats_table.cell(0, idx)
        set_cell_bg(label_cell, C['stat_bg'].lstrip('#'))
        add_paragraph_with_color(label_cell, label, font_size=7, bold=True, color_hex=C['stat_label'])

        val_cell = stats_table.cell(1, idx)
        set_cell_bg(val_cell, C['stat_bg'].lstrip('#'))
        add_paragraph_with_color(val_cell, value, font_size=13, bold=True, color_hex=color)

    document.add_paragraph()

    # === HOW TO USE ===
    how_p = document.add_paragraph()
    how_run = how_p.add_run(REPORT_TEXT['how_to_use_title'])
    how_run.font.size = Pt(8)
    how_run.font.bold = True
    r, g, b = hex_to_rgb(C['section_label'])
    how_run.font.color.rgb = RGBColor(r, g, b)

    for idx, step in enumerate(HOW_TO_USE_STEPS, start=1):
        p = document.add_paragraph(style='List Number')
        run = p.add_run(step)
        run.font.size = Pt(9)
        r, g, b = hex_to_rgb(C['text_secondary'])
        run.font.color.rgb = RGBColor(r, g, b)

    document.add_paragraph()

    # === FINDINGS ===
    findings_p = document.add_paragraph()
    findings_run = findings_p.add_run(REPORT_TEXT['findings_title'])
    findings_run.font.size = Pt(8)
    findings_run.font.bold = True
    r, g, b = hex_to_rgb(C['section_label'])
    findings_run.font.color.rgb = RGBColor(r, g, b)

    if not ordered:
        document.add_paragraph(REPORT_TEXT['no_findings'])
    else:
        for index, finding in enumerate(preview, start=1):
            risk = finding.get('risk_rating') or 'Low'
            sev = C['severity'].get(risk, C['severity']['default'])
            owner = RISK_OWNER.get(risk, RISK_OWNER['Low'])
            remediation_text, issue_summary, remediation_links = _smart_remediation(finding)

            # Finding card table: 3 cols — number | details | badge
            card_table = document.add_table(rows=4, cols=3)
            card_table.style = 'Table Grid'

            # Row 0: header
            num_cell = card_table.cell(0, 0)
            set_cell_bg(num_cell, C['card_header_bg'].lstrip('#'))
            add_paragraph_with_color(num_cell, str(index), font_size=8, bold=True, color_hex=C['text_muted'])

            title_cell = card_table.cell(0, 1)
            set_cell_bg(title_cell, C['card_header_bg'].lstrip('#'))
            add_paragraph_with_color(title_cell, finding.get('title') or 'Untitled finding',
                                     font_size=9, bold=True, color_hex=C['text_primary'])

            badge_cell = card_table.cell(0, 2)
            set_cell_bg(badge_cell, sev['bg'].lstrip('#'))
            bp = add_paragraph_with_color(badge_cell, risk, font_size=8, bold=True, color_hex=sev['text'])
            bp.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Row 1: meta fields
            meta_fields = [
                ('Category', finding.get('category') or 'Unknown'),
                ('Affected asset', finding.get('asset') or 'N/A'),
                ('Related CVEs', '; '.join(finding.get('cve_ids') or []) or 'Not provided'),
                ('Suggested owner', owner),
            ]
            meta_cell = card_table.cell(1, 1)
            set_cell_bg(meta_cell, C['card_bg'].lstrip('#'))
            meta_cell.paragraphs[0].clear()
            for k, v in meta_fields:
                mp = meta_cell.add_paragraph()
                key_run = mp.add_run(f"{k}: ")
                key_run.font.size = Pt(8)
                key_run.font.bold = False
                r, g, b = hex_to_rgb(C['text_muted'])
                key_run.font.color.rgb = RGBColor(r, g, b)
                val_run = mp.add_run(v)
                val_run.font.size = Pt(8)
                r, g, b = hex_to_rgb(C['text_secondary'])
                val_run.font.color.rgb = RGBColor(r, g, b)

            for col_idx in [0, 2]:
                set_cell_bg(card_table.cell(1, col_idx), C['card_bg'].lstrip('#'))

            # Row 2: about this issue
            about_cell = card_table.cell(2, 1)
            set_cell_bg(about_cell, C['card_bg'].lstrip('#'))
            about_cell.paragraphs[0].clear()

            about_label_p = about_cell.add_paragraph()
            about_label_r = about_label_p.add_run('ABOUT THIS ISSUE')
            about_label_r.font.size = Pt(7)
            about_label_r.font.bold = True
            r, g, b = hex_to_rgb(C['text_primary'])
            about_label_r.font.color.rgb = RGBColor(r, g, b)

            about_p = about_cell.add_paragraph()
            about_r = about_p.add_run(issue_summary if issue_summary else 'Finding details not available.')
            about_r.font.size = Pt(8)
            r, g, b = hex_to_rgb(C['text_secondary'])
            about_r.font.color.rgb = RGBColor(r, g, b)

            for col_idx in [0, 2]:
                set_cell_bg(card_table.cell(2, col_idx), C['card_bg'].lstrip('#'))

            # Row 3: remediation
            rem_cell = card_table.cell(3, 1)
            set_cell_bg(rem_cell, C['action_box_bg'].lstrip('#'))
            rem_cell.paragraphs[0].clear()

            label_run_p = rem_cell.add_paragraph()
            label_r = label_run_p.add_run('REMEDIATION')
            label_r.font.size = Pt(7)
            label_r.font.bold = True
            r, g, b = hex_to_rgb(C['text_muted'])
            label_r.font.color.rgb = RGBColor(r, g, b)

            rem_p = rem_cell.add_paragraph()
            rem_r = rem_p.add_run(remediation_text)
            rem_r.font.size = Pt(8)
            rem_r.font.italic = (remediation_text == 'No remediation provided.')
            r, g, b = hex_to_rgb(C['text_hint'] if remediation_text == 'No remediation provided.' else C['text_secondary'])
            rem_r.font.color.rgb = RGBColor(r, g, b)

            for link in remediation_links:
                lp = rem_cell.add_paragraph()
                lr = lp.add_run(link)
                lr.font.size = Pt(8)
                r, g, b = hex_to_rgb(C['link'])
                lr.font.color.rgb = RGBColor(r, g, b)

            for col_idx in [0, 2]:
                set_cell_bg(card_table.cell(3, col_idx), C['action_box_bg'].lstrip('#'))

            document.add_paragraph()

        footer_p = document.add_paragraph()
        footer_run = footer_p.add_run(f'Showing all {len(ordered)} findings.')
        footer_run.font.size = Pt(8)
        r, g, b = hex_to_rgb(C['text_hint'])
        footer_run.font.color.rgb = RGBColor(r, g, b)

    # === ADDITIONAL CHECKS ===
    checks_heading = document.add_paragraph()
    checks_run = checks_heading.add_run(REPORT_TEXT['checks_title'])
    checks_run.font.size = Pt(8)
    checks_run.font.bold = True
    r, g, b = hex_to_rgb(C['section_label'])
    checks_run.font.color.rgb = RGBColor(r, g, b)

    for section_name, section_items in _build_check_sections(payload):
        section_p = document.add_paragraph()
        section_r = section_p.add_run(section_name)
        section_r.font.size = Pt(9)
        section_r.font.bold = True

        if not section_items:
            note = document.add_paragraph('No records for this section.')
            note.runs[0].italic = True
            continue

        grouped_items = _group_check_items(section_items)
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        headers = ['Name', 'Status / Version', 'Affected Assets']
        for idx, text in enumerate(headers):
            table.cell(0, idx).text = text

        for g_item in grouped_items:
            row = table.add_row().cells
            status_risk = g_item['status']
            if g_item['risk_rating'] and g_item['risk_rating'].lower() not in ('low', 'recorded', ''):
                status_risk = f"{g_item['status']}  [{g_item['risk_rating']}]"
            row[0].text = g_item['name']
            row[1].text = status_risk
            row[2].text = ', '.join(g_item['assets'])

        document.add_paragraph()

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue(), None