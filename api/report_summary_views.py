from io import BytesIO
import logging
import re
from typing import Dict, List, Optional
from urllib.parse import urlparse

logger = logging.getLogger(__name__)

from django.db import transaction
from django.http import FileResponse, HttpResponse
from django.shortcuts import get_object_or_404
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView

from .permissions import get_user_organisation

from .models import (
    Asset,
    DNSSecurityCheck,
    Domain,
    EmailSecurityCheck,
    Finding,
    FrontendLibraryCheck,
    ReportSummary,
    ReportSummaryFinding,
    Scan,
    ScanAsset,
    SSLTLSCheck,
    SecurityHeaderCheck,
    TechnologyCheck,
)
from .report_external_links import get_external_reference


try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable,
        Paragraph as RLParagraph,
        SimpleDocTemplate,
        Spacer,
    )
    from reportlab.pdfgen import canvas
except ImportError:
    A4 = None
    canvas = None
    rl_colors = None
    ParagraphStyle = None
    cm = None
    HRFlowable = None
    RLParagraph = None
    SimpleDocTemplate = None
    Spacer = None

try:
    import docx
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor, Inches
except ImportError:
    docx = None
    OxmlElement = None
    qn = None
    Pt = None
    RGBColor = None
    Inches = None


SEVERITY_ORDER = {
    "Critical": 4,
    "High": 3,
    "Medium": 2,
    "Low": 1,
}


def _normalize_severity(raw: Optional[str]) -> str:
    value = (raw or "Low").strip().title()
    if value not in SEVERITY_ORDER:
        return "Low"
    return value


def _short_text(text: str, max_len: int = 180) -> str:
    cleaned = " ".join((text or "").split())
    if len(cleaned) <= max_len:
        return cleaned
    return cleaned[: max_len - 3].rstrip() + "..."


def _build_issue_and_action(detail: str, recommendation: str, fallback_action: str) -> Dict[str, str]:
    issue = _short_text(detail or "Security issue identified.", 220)
    action_source = fallback_action
    action = _short_text(action_source or "Review the external reference and apply vendor guidance.", 220)
    return {
        "issue_summary": issue,
        "recommendation_summary": action,
    }


def _is_primary_domain_asset(asset_value: str, root_domain: str) -> bool:
    candidate = (asset_value or "").strip().lower()
    if not candidate:
        return False

    if "://" in candidate:
        parsed = urlparse(candidate)
        candidate = parsed.hostname or candidate

    candidate = candidate.rstrip(".")
    if ":" in candidate:
        candidate = candidate.split(":", 1)[0]

    primary = (root_domain or "").strip().lower().rstrip(".")
    return bool(primary) and candidate == primary


def _asset_matches_domain_regex(asset_value: str, root_domain: str) -> bool:
    """Fallback domain matcher for assets when exact root asset is unavailable."""
    candidate = (asset_value or "").strip().lower()
    primary = (root_domain or "").strip().lower().rstrip(".")
    if not candidate or not primary:
        return False

    if "://" in candidate:
        parsed = urlparse(candidate)
        candidate = parsed.hostname or candidate

    candidate = candidate.rstrip(".")
    if ":" in candidate:
        candidate = candidate.split(":", 1)[0]

    # Match primary domain exactly or common host variants such as www.<domain>.
    pattern = rf"^(?:www\.)?{re.escape(primary)}$"
    return bool(re.match(pattern, candidate, re.IGNORECASE))


def _pack_issue_action(issue_summary: str, recommendation_summary: str) -> str:
    issue = " ".join((issue_summary or "Security issue identified.").split())
    action = " ".join((recommendation_summary or "Review the external reference and apply vendor guidance.").split())
    prefix = "Issue: "
    marker = " | Action: "
    max_len = 400

    packed = f"{prefix}{issue}{marker}{action}"
    if len(packed) <= max_len:
        return packed

    # Keep both segments and marker so _unpack_issue_action can always parse the value.
    fixed_len = len(prefix) + len(marker)
    budget = max_len - fixed_len

    # Reserve at least 80 chars for action details when possible.
    action_budget = min(max(len(action), 80), max(40, budget // 2))
    issue_budget = max(40, budget - action_budget)

    issue = _short_text(issue, issue_budget)
    action = _short_text(action, budget - len(issue))
    packed = f"{prefix}{issue}{marker}{action}"

    if len(packed) > max_len:
        action = _short_text(action, max_len - fixed_len - len(issue))
        packed = f"{prefix}{issue}{marker}{action}"

    return packed[:max_len]


def _unpack_issue_action(risk_summary: str) -> Dict[str, str]:
    text = (risk_summary or "").strip()
    marker = " | Action: "
    prefix = "Issue: "

    if marker in text and text.startswith(prefix):
        issue_text, action_text = text.split(marker, 1)
        return {
            "issue_summary": issue_text.replace(prefix, "", 1).strip(),
            "recommendation_summary": action_text.strip(),
        }

    return {
        "issue_summary": text or "Security issue identified.",
        "recommendation_summary": "Review the external reference and apply vendor guidance.",
    }


def _make_scope_note(domain: str, asset: Optional[str], scan_id: Optional[int]) -> str:
    if asset and scan_id:
        return f"Scoped to domain {domain}, asset {asset}, and scan #{scan_id}."
    if asset:
        return f"Scoped to domain {domain} and asset {asset}."
    if scan_id:
        return f"Scoped to domain {domain} across all scanned assets in scan #{scan_id}."
    return f"Scoped to domain {domain} across all scanned assets."


def _format_asset_list(asset_values: List[str], limit: int = 8) -> str:
    cleaned = []
    seen = set()
    for value in asset_values:
        item = (value or "").strip()
        if not item or item in seen:
            continue
        cleaned.append(item)
        seen.add(item)

    if not cleaned:
        return ""

    visible = cleaned[:limit]
    remaining = len(cleaned) - len(visible)
    text = ", ".join(visible)
    if remaining > 0:
        text += f", and {remaining} more"
    return text


def _build_executive_summary(
    org_name: str,
    domain: str,
    asset: Optional[str],
    scan_id: Optional[int],
    asset_values: Optional[List[str]],
    counts: Dict[str, int],
    total: int,
) -> str:
    org_display = org_name or domain
    domain_url = domain if domain.startswith("http") else f"https://{domain}/"
    severity_parts = []
    for label in ["Critical", "High", "Medium", "Low"]:
        count = counts.get(label, 0)
        if count:
            severity_parts.append(f"{count} {label}")
    severity_text = ", ".join(severity_parts) if severity_parts else "no material"

    if asset:
        scope_text = f"This summary was generated for the specific asset/subdomain {asset} under the {org_display} domain ({domain_url})."
    else:
        asset_list_text = _format_asset_list(asset_values or [])
        if asset_list_text:
            scope_text = (
                f"This summary covers the scanned assets for {org_display} ({domain_url}), including {asset_list_text}."
            )
        else:
            scope_text = f"This summary covers the scanned assets for {org_display} ({domain_url})."

    return (
        f"{scope_text} We identified {total} findings in total, including {severity_text} issues. "
        "These weaknesses require remediation to reduce the risk of unauthorized access, service compromise, data exposure, and abuse of the organization's digital assets."
    )


def _extract_findings(domain_obj: Domain, asset_obj: Optional[Asset], scan_obj: Optional[Scan]) -> List[Dict[str, str]]:
    asset_qs = Asset.objects.filter(domain=domain_obj)
    if asset_obj is not None:
        asset_qs = asset_qs.filter(id=asset_obj.id)
    if scan_obj is not None:
        # Limit report scope to assets that were actually part of this scan.
        asset_qs = asset_qs.filter(asset_scans__scan=scan_obj).distinct()

    items: List[Dict[str, str]] = []

    library_qs = FrontendLibraryCheck.objects.filter(asset__in=asset_qs)
    technology_qs = TechnologyCheck.objects.filter(asset__in=asset_qs)
    ssl_qs = SSLTLSCheck.objects.filter(asset__in=asset_qs)
    dns_qs = DNSSecurityCheck.objects.filter(asset__in=asset_qs)
    header_qs = SecurityHeaderCheck.objects.filter(asset__in=asset_qs)
    email_qs = EmailSecurityCheck.objects.filter(asset__in=asset_qs)
    findings_qs = Finding.objects.filter(asset__in=asset_qs)

    if scan_obj is not None:
        library_qs = library_qs.filter(scan=scan_obj)
        technology_qs = technology_qs.filter(scan=scan_obj)
        ssl_qs = ssl_qs.filter(scan=scan_obj)
        dns_qs = dns_qs.filter(scan=scan_obj)
        header_qs = header_qs.filter(scan=scan_obj)
        email_qs = email_qs.filter(scan=scan_obj)
        findings_qs = findings_qs.filter(finding_scans__scan=scan_obj)

    for row in findings_qs.select_related("asset").order_by("-last_seen", "-id")[:150]:
        detail = row.evidence or f"{row.category} finding detected."
        advice = _build_issue_and_action(
            detail=detail,
            recommendation="",
            fallback_action="Validate the finding and apply remediation guidance.",
        )
        items.append(
            {
                "finding_type": row.category,
                "title": row.title,
                "severity": _normalize_severity(row.risk_rating),
                "risk_summary": advice["issue_summary"],
                "recommendation_summary": advice["recommendation_summary"],
                "external_reference": get_external_reference(row.title, detail),
                "affected_asset": row.asset.value,
            }
        )

    for row in library_qs.order_by("-checked_at")[:100]:
        if row.vulnerability_status == "up-to-date":
            continue
        title = f"{row.library_name} {row.detected_version} ({row.vulnerability_status})"
        detail = f"Detected {row.vulnerability_status} library version on {row.asset.value}."
        advice = _build_issue_and_action(
            detail=detail,
            recommendation="",
            fallback_action="Update the affected library or plugin to the latest patched version.",
        )
        items.append(
            {
                "finding_type": "Technology",
                "title": title,
                "severity": _normalize_severity(row.risk_level),
                "risk_summary": advice["issue_summary"],
                "recommendation_summary": advice["recommendation_summary"],
                "external_reference": get_external_reference(title, detail),
                "affected_asset": row.asset.value,
            }
        )

    for row in technology_qs.order_by("-created_at")[:100]:
        version = (row.version or "Unknown").strip() or "Unknown"
        latest_version = (row.latest_version or "Unknown").strip() or "Unknown"
        is_outdated = (
            version.lower() != "unknown"
            and latest_version.lower() != "unknown"
            and version != latest_version
        )
        if is_outdated:
            title = f"{row.technology_name} {version} (outdated)"
            detail = (
                f"Technology category: {row.category}. Detected version {version} is behind latest {latest_version}."
            )
            recommendation = f"Upgrade {row.technology_name} from {version} to {latest_version}."
        else:
            title = f"{row.technology_name} {version}"
            detail = f"Technology category: {row.category}."
            recommendation = "Verify this component is current and supported; update if outdated."
        advice = _build_issue_and_action(
            detail=detail,
            recommendation=recommendation,
            fallback_action="Review vendor release notes and patch to a maintained version.",
        )
        items.append(
            {
                "finding_type": "Technology",
                "title": title,
                "severity": _normalize_severity(row.risk_level),
                "risk_summary": advice["issue_summary"],
                "recommendation_summary": advice["recommendation_summary"],
                "external_reference": get_external_reference(title, f"{detail} {advice['recommendation_summary']}"),
                "affected_asset": row.asset.value,
            }
        )

    for row in ssl_qs.order_by("-checked_at")[:100]:
        title = f"SSL/TLS {row.check_type}"
        detail = row.finding or row.example or "SSL/TLS weakness identified."
        advice = _build_issue_and_action(
            detail=detail,
            recommendation="",
            fallback_action="Update TLS configuration to modern protocols and secure cipher suites.",
        )
        items.append(
            {
                "finding_type": "SSL",
                "title": title,
                "severity": _normalize_severity(row.risk_rating),
                "risk_summary": advice["issue_summary"],
                "recommendation_summary": advice["recommendation_summary"],
                "external_reference": get_external_reference(title, detail),
                "affected_asset": row.asset.value,
            }
        )

    for row in dns_qs.order_by("-checked_at")[:100]:
        title = f"DNS {row.check_type}"
        detail = row.finding or row.example or "DNS issue identified."
        advice = _build_issue_and_action(
            detail=detail,
            recommendation="",
            fallback_action="Update DNS records and remove stale or misconfigured entries.",
        )
        items.append(
            {
                "finding_type": "DNS",
                "title": title,
                "severity": _normalize_severity(row.risk_rating),
                "risk_summary": advice["issue_summary"],
                "recommendation_summary": advice["recommendation_summary"],
                "external_reference": get_external_reference(title, detail),
                "affected_asset": row.asset.value,
            }
        )

    for row in header_qs.order_by("-checked_at")[:100]:
        if row.status == "present":
            continue
        title = f"Missing/Misconfigured Header: {row.header}"
        detail = f"Header status is {row.status}."
        advice = _build_issue_and_action(
            detail=detail,
            recommendation="",
            fallback_action="Apply the header with a secure value at the web server layer.",
        )
        items.append(
            {
                "finding_type": "Header",
                "title": title,
                "severity": _normalize_severity(row.risk_rating),
                "risk_summary": advice["issue_summary"],
                "recommendation_summary": advice["recommendation_summary"],
                "external_reference": get_external_reference(title, detail),
                "affected_asset": row.asset.value,
            }
        )

    # Email authentication (SPF/DKIM/DMARC) is a DOMAIN-level control published on
    # the organisational/root domain — it does not apply to arbitrary subdomains.
    # Only include email findings when the primary/root domain is part of the scope.
    primary_assets = [a.value for a in asset_qs if _is_primary_domain_asset(a.value, domain_obj.root_domain)]
    used_regex_fallback = False
    if not primary_assets:
        primary_assets = [a.value for a in asset_qs if _asset_matches_domain_regex(a.value, domain_obj.root_domain)]
        used_regex_fallback = bool(primary_assets)

    if primary_assets:
        matcher = _asset_matches_domain_regex if used_regex_fallback else _is_primary_domain_asset
        primary_asset = sorted(primary_assets)[0]

        # Real SPF/DKIM/DMARC checks, restricted to the primary domain asset only.
        email_rows = [
            row for row in email_qs.order_by("-checked_at")[:100]
            if row.asset and matcher(row.asset.value, domain_obj.root_domain)
        ]
        for row in email_rows:
            is_pass = (row.status or "").upper() == "PASS"
            title = f"Email Auth {row.check_type}: {row.status}"
            detail = row.details or "Email authentication misconfiguration detected."
            recommendation = "Maintain this control and continue periodic validation." if is_pass else "Correct SPF/DKIM/DMARC records and verify alignment."
            advice = _build_issue_and_action(
                detail=detail,
                recommendation=recommendation,
                fallback_action="Correct SPF/DKIM/DMARC records and verify alignment.",
            )
            items.append(
                {
                    "finding_type": "Email",
                    "title": title,
                    "severity": "Low" if is_pass else _normalize_severity(row.risk_rating),
                    "risk_summary": advice["issue_summary"],
                    "recommendation_summary": advice["recommendation_summary"],
                    "external_reference": get_external_reference(title, detail),
                    "affected_asset": row.asset.value,
                }
            )

        existing_types_for_primary = {(row.check_type or "").upper() for row in email_rows}
        for check_type in ("SPF", "DKIM", "DMARC"):
            if check_type in existing_types_for_primary:
                continue
            title = f"Email Auth {check_type}: NOT SET"
            detail = f"{check_type} record not detected for primary domain {domain_obj.root_domain}."
            advice = _build_issue_and_action(
                detail=detail,
                recommendation="",
                fallback_action="Configure SPF, DKIM, and DMARC records on the primary domain and validate DNS propagation.",
            )
            items.append(
                {
                    "finding_type": "Email",
                    "title": title,
                    "severity": "High" if check_type in {"SPF", "DMARC"} else "Medium",
                    "risk_summary": advice["issue_summary"],
                    "recommendation_summary": advice["recommendation_summary"],
                    "external_reference": get_external_reference(check_type, detail),
                    "affected_asset": primary_asset,
                }
            )

    deduped = {}
    for item in items:
        key = (item["finding_type"], item["title"], item["affected_asset"])
        if key not in deduped:
            deduped[key] = item

    final_items = sorted(
        deduped.values(),
        key=lambda row: SEVERITY_ORDER.get(row["severity"], 1),
        reverse=True,
    )
    return final_items


def _build_email_security_debug(domain_obj: Domain, asset_obj: Optional[Asset], scan_obj: Optional[Scan]) -> Dict[str, object]:
    """Build debug metadata for SPF/DKIM/DMARC evaluation in summary JSON exports."""
    check_types = ["SPF", "DKIM", "DMARC"]

    asset_qs = Asset.objects.filter(domain=domain_obj)
    if asset_obj is not None:
        asset_qs = asset_qs.filter(id=asset_obj.id)
    if scan_obj is not None:
        asset_qs = asset_qs.filter(asset_scans__scan=scan_obj).distinct()

    scoped_assets = list(asset_qs.order_by("value"))
    scoped_asset_values = [a.value for a in scoped_assets]

    email_qs = EmailSecurityCheck.objects.filter(asset__in=asset_qs)
    if scan_obj is not None:
        email_qs = email_qs.filter(scan=scan_obj)
    email_rows = list(email_qs.select_related("asset").order_by("asset__value", "check_type", "-checked_at"))

    by_asset: Dict[str, Dict[str, str]] = {}
    for row in email_rows:
        asset_value = row.asset.value if row.asset else ""
        if not asset_value:
            continue
        check_type = (row.check_type or "").upper()
        if check_type not in check_types:
            continue
        by_asset.setdefault(asset_value, {})
        # Keep the first (latest) status for each check type per asset.
        if check_type not in by_asset[asset_value]:
            by_asset[asset_value][check_type] = (row.status or "").upper() or "UNKNOWN"

    primary_exact = [a.value for a in scoped_assets if _is_primary_domain_asset(a.value, domain_obj.root_domain)]
    primary_regex = [a.value for a in scoped_assets if _asset_matches_domain_regex(a.value, domain_obj.root_domain)]

    if primary_exact:
        strategy = "exact"
        used_assets = sorted(set(primary_exact))
    elif primary_regex:
        strategy = "regex-fallback"
        used_assets = sorted(set(primary_regex))
    else:
        strategy = "none"
        used_assets = []

    evaluated_assets = used_assets or scoped_asset_values
    per_asset = []
    for value in evaluated_assets:
        found = by_asset.get(value, {})
        missing = [c for c in check_types if c not in found]
        per_asset.append(
            {
                "asset": value,
                "found": found,
                "missing": missing,
            }
        )

    return {
        "checks_evaluated": check_types,
        "scope": {
            "domain": domain_obj.root_domain,
            "asset_scope": asset_obj.value if asset_obj else "all-assets",
            "scan_id": scan_obj.id if scan_obj else None,
        },
        "assets_considered": scoped_asset_values,
        "primary_domain_matching": {
            "strategy": strategy,
            "exact_matches": sorted(set(primary_exact)),
            "regex_matches": sorted(set(primary_regex)),
            "used_assets": used_assets,
        },
        "records_found": len(email_rows),
        "per_asset_results": per_asset,
    }


def _safe_filename_part(value: str) -> str:
    safe = re.sub(r"[^a-zA-Z0-9._-]", "-", value.strip())
    return safe[:80] or "report"


# ─────────────────────────────────────────────────────────────────────────────
# Render helpers
# ─────────────────────────────────────────────────────────────────────────────

_SEVERITY_HEX = {
    "Critical": "#C00000",
    "High":     "#FF0000",
    "Medium":   "#FF6600",
    "Low":      "#4472C4",
}


def _sev_badge_html(sev: str) -> str:
    c = _SEVERITY_HEX.get(sev, "#000000")
    return f'<font color="{c}"><b>[{sev}]</b></font>'


def _severity_rgb(sev: str):
    mapping = {
        "Critical": (0xC0, 0x00, 0x00),
        "High":     (0xFF, 0x00, 0x00),
        "Medium":   (0xFF, 0x66, 0x00),
        "Low":      (0x44, 0x72, 0xC4),
    }
    r, g, b = mapping.get(sev, (0, 0, 0))
    return RGBColor(r, g, b)


def _add_hyperlink(paragraph, link_text: str, url: str):
    """Append a blue underlined hyperlink run to a python-docx paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "0563C1")
    rPr.append(color_el)

    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(u_el)

    run_el.append(rPr)
    t_el = OxmlElement("w:t")
    t_el.text = link_text
    run_el.append(t_el)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)
    return hyperlink


def _get_report_date_text(summary: ReportSummary) -> str:
    if summary.scan:
        if summary.scan.started_at:
            return summary.scan.started_at.strftime("%d %B %Y")
        if summary.scan.created_at:
            return summary.scan.created_at.strftime("%d %B %Y")
    if summary.generated_at:
        return summary.generated_at.strftime("%d %B %Y")
    return ""


def _group_findings(findings: List[ReportSummaryFinding]) -> Dict[str, List[Dict[str, str]]]:
    severities = ["Critical", "High", "Medium", "Low"]

    category_map: Dict[str, Dict[str, object]] = {}
    asset_map: Dict[str, Dict[str, object]] = {}
    issue_map: Dict[str, Dict[str, object]] = {}

    for f in findings:
        category = (f.finding_type or "Other").strip() or "Other"
        asset = (f.affected_asset or "Unknown Asset").strip() or "Unknown Asset"
        title = (f.title or "Untitled finding").strip() or "Untitled finding"
        severity = _normalize_severity(f.severity)
        issue_text = _unpack_issue_action(f.risk_summary)["issue_summary"]

        if category not in category_map:
            category_map[category] = {
                "count": 0,
                "assets": set(),
                "severity": {s: 0 for s in severities},
            }
        category_map[category]["count"] += 1
        category_map[category]["assets"].add(asset)
        category_map[category]["severity"][severity] += 1

        if asset not in asset_map:
            asset_map[asset] = {
                "count": 0,
                "categories": set(),
                "high_risk": 0,
            }
        asset_map[asset]["count"] += 1
        asset_map[asset]["categories"].add(category)
        if severity in {"Critical", "High"}:
            asset_map[asset]["high_risk"] += 1

        issue_key = f"{category}::{title}"
        if issue_key not in issue_map:
            issue_map[issue_key] = {
                "title": title,
                "category": category,
                "severity": severity,
                "count": 0,
                "assets": set(),
                "issue": issue_text,
            }
        issue_map[issue_key]["count"] += 1
        issue_map[issue_key]["assets"].add(asset)

    category_rows = []
    for category, info in category_map.items():
        sev = info["severity"]
        category_rows.append(
            {
                "category": category,
                "count": info["count"],
                "assets": len(info["assets"]),
                "severity_text": ", ".join(
                    [f"{s}: {sev[s]}" for s in severities if sev[s] > 0] or ["Low: 0"]
                ),
            }
        )
    category_rows.sort(key=lambda row: row["count"], reverse=True)

    asset_rows = []
    for asset, info in asset_map.items():
        categories = sorted(info["categories"])
        asset_rows.append(
            {
                "asset": asset,
                "count": info["count"],
                "high_risk": info["high_risk"],
                "categories": ", ".join(categories),
            }
        )
    asset_rows.sort(key=lambda row: (row["high_risk"], row["count"]), reverse=True)

    issue_rows = []
    for info in issue_map.values():
        issue_rows.append(
            {
                "title": info["title"],
                "category": info["category"],
                "severity": info["severity"],
                "count": info["count"],
                "assets": len(info["assets"]),
                "issue": info["issue"],
            }
        )
    issue_rows.sort(
        key=lambda row: (
            SEVERITY_ORDER.get(row["severity"], 1),
            row["count"],
        ),
        reverse=True,
    )

    return {
        "categories": category_rows,
        "assets": asset_rows,
        "issues": issue_rows,
    }


def _extract_email_statuses(findings: List[ReportSummaryFinding]) -> Dict[str, str]:
    states: Dict[str, str] = {}
    priority = {"NOT SET": 4, "FAIL": 3, "INVALID": 2, "PASS": 1}

    for f in findings:
        if (f.finding_type or "").strip().lower() != "email":
            continue
        match = re.search(r"Email Auth\s+([A-Za-z]+)\s*:\s*(.+)$", f.title or "", re.IGNORECASE)
        if not match:
            continue
        check_type = match.group(1).upper().strip()
        status = match.group(2).upper().strip()
        if check_type not in {"SPF", "DKIM", "DMARC"}:
            continue
        current = states.get(check_type)
        if not current or priority.get(status, 0) > priority.get(current, 0):
            states[check_type] = status

    return states


def _build_narrative_sections(summary: ReportSummary, findings: List[ReportSummaryFinding]) -> List[Dict[str, object]]:
    sections: List[Dict[str, object]] = []
    domain_name = summary.domain.root_domain if summary.domain else "the domain"

    tech_rows = [
        f for f in findings
        if (f.finding_type or "").strip().lower() == "technology"
        and any(k in (f.title or "").lower() for k in ["outdated", "vulnerable", "check-failed", "unknown"])
    ]
    if tech_rows:
        tech_titles = [f.title for f in tech_rows[:5]]
        tech_list = ", ".join(tech_titles)
        sections.append(
            {
                "title": "Outdated Software Packages",
                "paragraph": (
                    "We observed the presence of obsolete or unsupported software components on the website, "
                    f"including {tech_list}. These versions may contain known weaknesses that can be exploited "
                    "by attackers to gain unauthorized access or compromise user data."
                ),
                "bullets": [],
            }
        )

    header_rows = [f for f in findings if (f.finding_type or "").strip().lower() == "header"]
    if header_rows:
        headers = []
        for f in header_rows:
            title = f.title or ""
            if ":" in title:
                header_name = title.split(":", 1)[1].strip()
            else:
                header_name = title.strip()
            if header_name and header_name not in headers:
                headers.append(header_name)

        header_text = ", ".join(headers[:8])
        sections.append(
            {
                "title": "Missing Security Headers",
                "paragraph": (
                    f"Our monitoring indicates that {domain_name} is not implementing important security headers, "
                    f"including {header_text}. These controls help protect against attacks such as cross-site "
                    "scripting (XSS), clickjacking, content injection, and referrer information leakage."
                ),
                "bullets": [],
            }
        )


    # --- EMAIL SECTION (debug-aware, matches what is actually present) ---
    # Use the same logic as email_security_debug for accuracy
    from .models import Asset, EmailSecurityCheck, ScanAsset
    domain_obj = summary.domain
    asset_obj = summary.asset
    scan_obj = summary.scan
    check_types = ["SPF", "DKIM", "DMARC"]
    asset_qs = Asset.objects.filter(domain=domain_obj)
    if asset_obj is not None:
        asset_qs = asset_qs.filter(id=asset_obj.id)
    if scan_obj is not None:
        asset_qs = asset_qs.filter(asset_scans__scan=scan_obj).distinct()
    scoped_assets = list(asset_qs.order_by("value"))
    primary_exact = [a.value for a in scoped_assets if _is_primary_domain_asset(a.value, domain_obj.root_domain)]
    primary_regex = [a.value for a in scoped_assets if _asset_matches_domain_regex(a.value, domain_obj.root_domain)]
    if primary_exact:
        matcher = _is_primary_domain_asset
        used_assets = sorted(set(primary_exact))
    elif primary_regex:
        matcher = _asset_matches_domain_regex
        used_assets = sorted(set(primary_regex))
    else:
        matcher = None
        used_assets = []

    # SPF/DKIM/DMARC are DOMAIN-level controls published on the organisational/root
    # domain — they do not apply to subdomains. If the report scope has no primary
    # domain asset (e.g. a subdomain-only export), omit the email section entirely.
    if matcher is None:
        return sections

    # Gather all checks for the used assets
    email_qs = EmailSecurityCheck.objects.filter(asset__in=asset_qs)
    if scan_obj is not None:
        email_qs = email_qs.filter(scan=scan_obj)
    email_rows = list(email_qs.select_related("asset").order_by("asset__value", "check_type", "-checked_at"))
    found: dict = {}
    for row in email_rows:
        if matcher and not matcher(row.asset.value, domain_obj.root_domain):
            continue
        check_type = (row.check_type or "").upper()
        if check_type in check_types and check_type not in found:
            found[check_type] = (row.status or "").upper() or "UNKNOWN"
    bullets = []
    risks = []
    for check in check_types:
        status = found.get(check)
        if status == "PASS":
            bullets.append(f"{check} records found, and all tests passed.")
        elif status:
            bullets.append(f"{check} record found, but validation failed ({status}).")
            if check == "DMARC":
                risks.append("The absence or failure of DMARC leaves email handling rules unclear and increases exposure to spoofing and phishing")
            if check == "DKIM":
                risks.append("without DKIM signing, email integrity cannot be strongly verified, increasing the risk of tampering")
            if check == "SPF":
                risks.append("missing or failing SPF weakens sender validation and allows unauthorized systems to send mail on behalf of the domain")
        else:
            bullets.append(f"No {check} record found.")
            if check == "DMARC":
                risks.append("The absence or failure of DMARC leaves email handling rules unclear and increases exposure to spoofing and phishing")
            if check == "DKIM":
                risks.append("without DKIM signing, email integrity cannot be strongly verified, increasing the risk of tampering")
            if check == "SPF":
                risks.append("missing or failing SPF weakens sender validation and allows unauthorized systems to send mail on behalf of the domain")
    email_paragraph = " ".join(risks) + "." if risks else (
        "Email authentication controls are configured correctly and should be continuously monitored to sustain protection."
    )
    if bullets:
        sections.append(
            {
                "title": "Email Authentication Methods",
                "paragraph": email_paragraph[:1].upper() + email_paragraph[1:],
                "bullets": bullets,
            }
        )

    return sections


def _build_narrative_recommendations(summary: ReportSummary, findings: List[ReportSummaryFinding]) -> List[str]:
    recommendations: List[str] = []
    finding_types = {(f.finding_type or "").strip().lower() for f in findings}
    email_states = _extract_email_statuses(findings)

    if "technology" in finding_types:
        recommendations.append(
            "It is strongly recommended that outdated or vulnerable software packages be upgraded to currently supported versions to reduce known security risks."
        )

    if "header" in finding_types:
        domain_name = summary.domain.root_domain if summary.domain else "the domain"
        recommendations.append(
            f"Implement the missing security headers on {domain_name} and validate them across all public-facing assets."
        )

    if email_states and any(email_states.get(x, "NOT SET") != "PASS" for x in ["SPF", "DKIM", "DMARC"]):
        recommendations.append(
            "Add or correct SPF, DKIM, and DMARC to strengthen email security and reduce spoofing, phishing, and impersonation risks."
        )
        recommendations.append(
            "After monitoring legitimate traffic using DMARC aggregate reports, enforce a reject policy (p=reject) to maximize protection."
        )

    recommendations.append(
        "Conduct periodic Vulnerability Assessment and Penetration Testing (VAPT) to identify and remediate newly introduced weaknesses."
    )
    return recommendations


# ─────────────────────────────────────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────────────────────────────────────

def _render_pdf(summary: ReportSummary, findings: List[ReportSummaryFinding]) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )

    RED = rl_colors.HexColor("#000000")
    uid = id(doc)  # unique per call — prevents reportlab style name collisions

    body    = ParagraphStyle(f"NBody_{uid}",  fontName="Helvetica", fontSize=11, leading=16, spaceAfter=4)
    h1      = ParagraphStyle(f"NH1_{uid}",    fontName="Helvetica-Bold", fontSize=14, leading=20,
                             spaceBefore=14, spaceAfter=4, textColor=RED)
    h2      = ParagraphStyle(f"NH2_{uid}",    fontName="Helvetica-Bold", fontSize=12, leading=18,
                             spaceBefore=10, spaceAfter=2)
    meta    = ParagraphStyle(f"NMeta_{uid}",  fontName="Helvetica", fontSize=10, leading=14,
                             spaceAfter=3, textColor=rl_colors.HexColor("#444444"))
    title_s = ParagraphStyle(f"NTitle_{uid}", fontName="Helvetica-Bold", fontSize=16,
                             leading=22, spaceAfter=4, textColor=RED)

    def hr():
        return HRFlowable(width="100%", thickness=1, color=RED, spaceAfter=6)

    def esc(text):
        return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    story = []

    # Title with resolved organization name and report date
    centered_title = ParagraphStyle(
        f"NTitleCentered_{uid}",
        parent=title_s,
        alignment=1,  # 1 = center
    )
    story.append(RLParagraph(f"Report on Website Security Vulnerabilities for {summary.organization_name}", centered_title))
    report_date = _get_report_date_text(summary)
    if report_date:
        centered_meta = ParagraphStyle(
            f"NMetaCentered_{uid}",
            parent=meta,
            alignment=1,  # 1 = center
        )
        story.append(RLParagraph(report_date, centered_meta))

    story.append(Spacer(1, 0.3 * cm))

    # Executive Summary with all links blue/underlined
    story.append(RLParagraph("Executive Summary", h1))
    story.append(hr())
    exec_sum_style = ParagraphStyle(
        f"NExecSum_{uid}",
        fontName="Helvetica",
        fontSize=11,
        leading=16,
        spaceAfter=4,
        alignment=4,  # 4 = justify
    )
    exec_sum = summary.executive_summary
    # Highlight all links (http/https/asset domains) in blue/underline
    def highlight_links(text):
        # Highlight all URLs
        text = re.sub(r'(https?://[\w\.-]+)', r'<font color="#0563C1"><u>\1</u></font>', text)
        # Highlight asset domains (e.g., acemfs.futminna.edu.ng)
        text = re.sub(r'(?<![\w\.-])([\w\.-]+\.futminna\.edu\.ng)(?![\w\.-])', lambda m: f'<font color="#0563C1"><u>{m.group(1)}</u></font>' if not m.group(1).startswith('http') else m.group(1), text)
        return text
    exec_sum = highlight_links(exec_sum)
    story.append(RLParagraph(esc(exec_sum), exec_sum_style))
    story.append(Spacer(1, 0.3 * cm))

    # Preliminary Findings
    story.append(RLParagraph("Preliminary Findings", h1))
    story.append(hr())

    sections = _build_narrative_sections(summary, findings)
    if not sections:
        story.append(RLParagraph("No significant findings were identified in this scope.", body))
    else:
        for index, section in enumerate(sections, start=1):
            # Headings black
            story.append(RLParagraph(f'<b>{index}. {esc(section["title"])}</b>', h2))
            para = section["paragraph"]
            # First, highlight libraries/versions in blue (without 'vulnerable')
            para = re.sub(r'(Swiper [\d.]+|jQuery [\d.]+)', r'<font color="#0563C1"><b>\1</b></font>', para)
            # Then, highlight (vulnerable) in orange inside blue
            para = re.sub(r'\(<font color="#0563C1"><b>([\d.]+)</b></font> \(<font color="#FF6600"><b>vulnerable</b></font>\)\)', r'(<font color="#0563C1"><b>\1</b></font> (<font color="#FF6600"><b>vulnerable</b></font>))', para)
            # Or, highlight (vulnerable) in orange (standalone)
            para = re.sub(r'\([vV]ulnerable\)', r'(<font color="#FF6600"><b>vulnerable</b></font>)', para)
            # Highlight missing headers and email statuses in blue
            para = re.sub(r'(Content-Security-Policy|Strict-Transport-Security|SPF|DKIM|DMARC)', r'<font color="#0563C1"><b>\1</b></font>', para)
            # Highlight vulnerability keywords in orange bold (non-overlapping)
            para = re.sub(r'\b(vulnerab[a-z]+|finding|weakness|exploit|risk|exposure|threat|attack)\b', r'<font color="#FF6600"><b>\1</b></font>', para, flags=re.IGNORECASE)
            story.append(RLParagraph(esc(para), body))
            for bullet in section.get("bullets", []):
                # Highlight missing/failed/none in blue
                bullet = re.sub(r'(No [A-Z]+ record found|validation failed \([A-Z]+\)|records found, and all tests passed)', r'<font color="#0563C1"><b>\1</b></font>', bullet)
                story.append(RLParagraph(f"• {esc(bullet)}", body))
            story.append(Spacer(1, 0.18 * cm))

    # Recommendations (keep black)
    story.append(RLParagraph("Recommendations", h1))
    story.append(hr())
    for recommendation in _build_narrative_recommendations(summary, findings):
        story.append(RLParagraph(f"• {esc(recommendation)}", body))
    story.append(Spacer(1, 0.3 * cm))

    # Conclusion
    story.append(RLParagraph("Conclusion", h1))
    story.append(hr())
    conclusion = (
        "Addressing the identified vulnerabilities is essential to mitigate potential security "
        "risks and uphold the trust and integrity of your digital assets. We strongly recommend "
        "taking immediate action to implement the recommended measures."
    )
    conclusion = re.sub(r"(vulnerab[a-z]+|finding|weakness|exploit|risk|exposure|threat|attack)", r'<font color="#FF6600"><b>\1</b></font>', conclusion, flags=re.IGNORECASE)
    story.append(RLParagraph(esc(conclusion), body))

    doc.build(story)
    buffer.seek(0)
    return buffer


# ─────────────────────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────────────────────

def _render_docx(summary: ReportSummary, findings: List[ReportSummaryFinding]) -> BytesIO:
    document = docx.Document()

    # Verdana 11pt — matches the NAHCON house style
    normal = document.styles["Normal"]
    normal.font.name = "Verdana"
    normal.font.size = Pt(11)

    def _red_heading(text, size=13):
        p = document.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        r.font.name = "Verdana"
        return p

    def _bold_value_para(label, value, indent=False):
        p = document.add_paragraph()
        lb = p.add_run(f"{label}: ")
        lb.bold = True
        lb.font.name = "Verdana"
        vr = p.add_run(value)
        vr.font.name = "Verdana"
        if indent:
            p.paragraph_format.left_indent = Inches(0.25)
        return p

    def _body_para(text="", indent=False):
        p = document.add_paragraph()
        if text:
            r = p.add_run(text)
            r.font.name = "Verdana"
        if indent:
            p.paragraph_format.left_indent = Inches(0.25)
        return p

    # Centered title
    title_p = document.add_paragraph()
    title_p.alignment = 1  # 1 = center
    title_run = title_p.add_run(f"Report on Website Security Vulnerabilities for {summary.organization_name}")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Verdana"
    report_date = _get_report_date_text(summary)
    if report_date:
        date_p = document.add_paragraph()
        date_p.alignment = 1  # 1 = center
        date_run = date_p.add_run(report_date)
        date_run.font.name = "Verdana"
    document.add_paragraph()

    # Executive Summary with all links blue/underlined
    _red_heading("Executive Summary")
    exec_sum_p = document.add_paragraph()
    exec_sum_p.alignment = 3  # 3 = justify
    exec_sum = summary.executive_summary
    # Highlight all links (http/https/asset domains) in blue/underline
    import re as _re
    def docx_highlight_links(paragraph, text):
        # Highlight all URLs
        url_pattern = _re.compile(r'(https?://[\w\.-]+)')
        asset_pattern = _re.compile(r'([\w\.-]+\.futminna\.edu\.ng)')
        pos = 0
        for m in url_pattern.finditer(text):
            if m.start() > pos:
                paragraph.add_run(text[pos:m.start()])
            run = paragraph.add_run(m.group(1))
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.font.underline = True
            run.font.name = "Verdana"
            pos = m.end()
        text2 = text[pos:]
        pos2 = 0
        for m in asset_pattern.finditer(text2):
            if m.start() > pos2:
                paragraph.add_run(text2[pos2:m.start()])
            if not text2[m.start():m.end()].startswith('http'):
                run = paragraph.add_run(m.group(1))
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                run.font.underline = True
                run.font.name = "Verdana"
            else:
                paragraph.add_run(m.group(1))
            pos2 = m.end()
        if pos2 < len(text2):
            paragraph.add_run(text2[pos2:])
    docx_highlight_links(exec_sum_p, exec_sum)
    for r in exec_sum_p.runs:
        r.font.size = Pt(11)
    document.add_paragraph()

    # Preliminary Findings
    _red_heading("Preliminary Findings")

    sections = _build_narrative_sections(summary, findings)
    if not sections:
        _body_para("No significant findings were identified in this scope.")
    else:
        for index, section in enumerate(sections, start=1):
            sh = document.add_paragraph()
            shr = sh.add_run(f"{index}. {section['title']}")
            shr.bold = True
            shr.font.size = Pt(12)
            shr.font.name = "Verdana"
            # Headings black
            # Highlight vulnerable libraries, missing headers, and email statuses in blue, and 'vulnerable' in orange, non-overlapping
            para = section["paragraph"]
            # Step 1: Tokenize for libraries/versions
            tokens = []
            last = 0
            for m in _re.finditer(r'(Swiper [\d.]+|jQuery [\d.]+|Content-Security-Policy|Strict-Transport-Security|SPF|DKIM|DMARC)', para):
                if m.start() > last:
                    tokens.append(("text", para[last:m.start()]))
                tokens.append(("blue", m.group(0)))
                last = m.end()
            if last < len(para):
                tokens.append(("text", para[last:]))
            # Step 2: For each token, further split for (vulnerable) and keywords
            def split_token(token_type, text):
                result = []
                pos = 0
                vuln_pat = _re.compile(r'\([vV]ulnerable\)')
                key_pat = _re.compile(r'\b(vulnerab[a-z]+|finding|weakness|exploit|risk|exposure|threat|attack)\b', _re.IGNORECASE)
                while pos < len(text):
                    m1 = vuln_pat.search(text, pos)
                    m2 = key_pat.search(text, pos)
                    m = None
                    if m1 and (not m2 or m1.start() < m2.start()):
                        m = m1
                        if m.start() > pos:
                            result.append((token_type, text[pos:m.start()]))
                        # vuln_pat has no capturing group, so use group(0)
                        result.append(("orange", m.group(0)))
                        pos = m.end()
                    elif m2:
                        if m2.start() > pos:
                            result.append((token_type, text[pos:m2.start()]))
                        result.append(("orange", m2.group(1).upper()))
                        pos = m2.end()
                    else:
                        result.append((token_type, text[pos:]))
                        break
                return result
            # Step 3: Render
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            for ttype, tval in tokens:
                for stype, sval in split_token(ttype, tval):
                    run = p.add_run(sval)
                    run.font.name = "Verdana"
                    if stype == "blue":
                        run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                        run.font.bold = True
                    elif stype == "orange":
                        run.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
                        run.font.bold = True
            # Bullets
            for bullet in section.get("bullets", []):
                bp = document.add_paragraph(style="List Bullet")
                bp.paragraph_format.left_indent = Inches(0.25)
                # Highlight missing/failed/none in blue
                btxt = bullet
                for pat in [r'(No [A-Z]+ record found|validation failed \([A-Z]+\)|records found, and all tests passed)']:
                    btxt = _re.sub(pat, lambda m: f"<BLUE>{m.group(1)}</BLUE>", btxt)
                # Render bullet with color
                i = 0
                while i < len(btxt):
                    if btxt.startswith('<BLUE>', i):
                        i += 6
                        j = btxt.find('</BLUE>', i)
                        run = bp.add_run(btxt[i:j])
                        run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                        run.font.bold = True
                        run.font.name = "Verdana"
                        i = j + 7
                    else:
                        j = btxt.find('<', i)
                        if j == -1:
                            bp.add_run(btxt[i:])
                            break
                        else:
                            bp.add_run(btxt[i:j])
                            i = j
            document.add_paragraph()

    # Recommendations (keep black)
    _red_heading("Recommendations")
    for recommendation in _build_narrative_recommendations(summary, findings):
        bp = document.add_paragraph(style="List Bullet")
        br = bp.add_run(recommendation)
        br.font.name = "Verdana"

    document.add_paragraph()

    # Conclusion
    _red_heading("Conclusion")
    conclusion = (
        "Addressing the identified vulnerabilities is essential to mitigate potential security "
        "risks and uphold the trust and integrity of your digital assets. We strongly recommend "
        "taking immediate action to implement the recommended measures."
    )
    _body_para(conclusion)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


# ─────────────────────────────────────────────────────────────────────────────
# View
# ─────────────────────────────────────────────────────────────────────────────

class ReportSummaryExportView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, scan_id=None, report_format=None, asset_id=None):
        # --- Import report_style_config for findings_preview and summary logic ---
        from . import report_style_config

        org = get_user_organisation(request)

        domain_name = (request.query_params.get("domain") or "").strip()
        asset_value = (request.query_params.get("asset") or "").strip()
        scan_id_value = scan_id or request.query_params.get("scan_id")
        export_format = (report_format or request.query_params.get("format") or "pdf").strip().lower()

        if export_format not in {"pdf", "docx", "json"}:
            return Response({"error": "format must be pdf, docx, or json"}, status=400)

        scan_obj = None
        if scan_id_value:
            try:
                scan_id_int = int(scan_id_value)
            except ValueError:
                return Response({"error": "scan_id must be an integer"}, status=400)
            scan_obj = get_object_or_404(Scan, id=scan_id_int, organisation=org)

        scan_assets = []
        if scan_obj is not None:
            scan_assets = list(
                ScanAsset.objects.filter(scan=scan_obj)
                .select_related("asset", "asset__domain")
            )
            if not scan_assets:
                return Response({"error": f"No assets found for scan {scan_obj.id}"}, status=400)

        asset_obj = None
        if asset_id is not None:
            asset_obj = get_object_or_404(Asset, id=asset_id, organisation=org)
            if scan_obj is not None and not ScanAsset.objects.filter(scan=scan_obj, asset=asset_obj).exists():
                return Response(
                    {"error": f"Asset {asset_id} is not part of scan {scan_obj.id}"},
                    status=400,
                )
            domain_obj = asset_obj.domain
        else:
            if domain_name:
                domain_obj = get_object_or_404(Domain, root_domain__iexact=domain_name, organisation=org)
            elif scan_obj is not None:
                domain_ids = {row.asset.domain_id for row in scan_assets}
                if len(domain_ids) != 1:
                    return Response(
                        {"error": "Scan contains assets from multiple domains. Provide domain query parameter."},
                        status=400,
                    )
                domain_obj = scan_assets[0].asset.domain
            else:
                return Response({"error": "domain or scan_id is required"}, status=400)

            if asset_value:
                asset_obj = get_object_or_404(Asset, domain=domain_obj, value=asset_value, organisation=org)
            elif scan_assets:
                unique_asset_ids = {row.asset_id for row in scan_assets if row.asset.domain_id == domain_obj.id}
                if len(unique_asset_ids) == 1:
                    asset_obj = scan_assets[0].asset

        # If exporting by asset, always resolve org from asset's domain
        if asset_obj is not None:
            asset_domain = asset_obj.domain
            resolved_org = asset_domain.owner.strip() if getattr(asset_domain, 'owner', None) else asset_domain.root_domain
            domain_for_report = asset_domain.root_domain
        else:
            resolved_org = domain_obj.owner.strip() if getattr(domain_obj, 'owner', None) else domain_obj.root_domain
            domain_for_report = domain_obj.root_domain

        scoped_asset_values = []
        if asset_obj is not None:
            scoped_asset_values = [asset_obj.value]
        elif scan_assets:
            scoped_asset_values = [
                row.asset.value
                for row in scan_assets
                if row.asset and row.asset.domain_id == domain_obj.id
            ]
        elif domain_obj is not None:
            scoped_asset_values = list(
                Asset.objects.filter(domain=domain_obj).order_by("value").values_list("value", flat=True)
            )


        # --- Use findings extraction and preview logic from report_style_config ---

        # --- Build findings dicts to match main report structure exactly ---
        findings_raw = _extract_findings(domain_obj, asset_obj, scan_obj)
        findings = []
        for f in findings_raw:
            finding_dict = {
                'title': f.get('title'),
                'category': f.get('finding_type'),
                'asset': f.get('affected_asset'),
                'risk_rating': f.get('severity'),
                'status': '',
                'evidence': f.get('risk_summary'),
                'recommendation': f.get('recommendation_summary'),
                'links': [f.get('external_reference')] if f.get('external_reference') else [],
            }
            findings.append(finding_dict)

        # Use findings_preview to dedupe and sort, as main report does
        ordered, preview, remaining = report_style_config.findings_preview(findings)

        # Build counts dict to match main report (Critical, High, Medium, Low)
        summary_counts = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0}
        for item in ordered:
            sev = (item.get('risk_rating') or 'Low').title()
            if sev in summary_counts:
                summary_counts[sev] += 1
        summary_total = len(ordered)

        executive_summary = _build_executive_summary(
            org_name=resolved_org,
            domain=domain_for_report,
            asset=asset_obj.value if asset_obj else None,
            scan_id=scan_obj.id if scan_obj else None,
            asset_values=scoped_asset_values,
            counts=summary_counts,
            total=summary_total,
        )

        if export_format == "json":
            preview_items = [
                {
                    "finding_type": item.get("finding_type"),
                    "title": item.get("title"),
                    "severity": item.get("risk_rating"),
                    "risk_summary": item.get("evidence"),
                    "recommendation_summary": item.get("recommendation"),
                    "external_reference": item.get("links", []),
                    "affected_asset": item.get("asset"),
                }
                for item in preview[:10]
            ]
            email_debug = _build_email_security_debug(domain_obj, asset_obj, scan_obj)
            return Response(
                {
                    "debug_preview": True,
                    "organization_name": resolved_org,
                    "domain": domain_obj.root_domain,
                    "asset_scope": asset_obj.value if asset_obj else "all-assets",
                    "scan_id": scan_obj.id if scan_obj else None,
                    "executive_summary": executive_summary,
                    "counts": summary_counts,
                    "total_findings": summary_total,
                    "preview_count": len(preview_items),
                    "findings_preview": preview_items,
                    "email_security_debug": email_debug,
                    "note": "Preview only. No summary rows were persisted.",
                }
            )

        if export_format == "pdf" and not SimpleDocTemplate:
            return Response({"error": "reportlab is not installed"}, status=500)
        if export_format == "docx" and not docx:
            return Response({"error": "python-docx is not installed"}, status=500)

        with transaction.atomic():
            summary, _ = ReportSummary.objects.update_or_create(
                domain=domain_obj,
                asset=asset_obj,
                scan=scan_obj,
                defaults={
                    "organisation": org,
                    "organization_name": resolved_org,
                    "executive_summary": executive_summary,
                    "scope_note": _make_scope_note(
                        domain_obj.root_domain,
                        asset_obj.value if asset_obj else None,
                        scan_obj.id if scan_obj else None,
                    ),
                    "total_findings": summary_total,
                    "critical_risk": summary_counts["Critical"],
                    "high_risk": summary_counts["High"],
                    "medium_risk": summary_counts["Medium"],
                    "low_risk": summary_counts["Low"],
                },
            )

            # Replace prior findings for this exact scope to avoid duplicate rows on re-export.
            summary.findings.all().delete()

            finding_rows = [
                ReportSummaryFinding(
                    summary=summary,
                    finding_type=item.get("category"),
                    title=item.get("title"),
                    severity=item.get("risk_rating"),
                    risk_summary=item.get("evidence"),
                    external_reference=", ".join(item.get("links", [])),
                    affected_asset=item.get("asset"),
                )
                for item in ordered
            ]
            ReportSummaryFinding.objects.bulk_create(finding_rows)

        persisted_findings = list(summary.findings.all().order_by("created_at"))
        scan_part = f"scan-{scan_obj.id}" if scan_obj else "all-scans"
        base_name = _safe_filename_part(f"{domain_obj.root_domain}-{asset_obj.value if asset_obj else 'all-assets'}-{scan_part}")

        if export_format == "pdf":
            output = _render_pdf(summary, persisted_findings)
            return FileResponse(output, as_attachment=True, filename=f"{base_name}-security-summary.pdf")

        output = _render_docx(summary, persisted_findings)
        response = HttpResponse(
            output.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{base_name}-security-summary.docx"'
        return response