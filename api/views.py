from rest_framework import viewsets, status
from rest_framework.decorators import action, api_view
from rest_framework.response import Response
from rest_framework.pagination import PageNumberPagination
from rest_framework.test import APIRequestFactory
from rest_framework.exceptions import ValidationError
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from django.db.models import Count, Q
from django.db import IntegrityError, connection
from django.conf import settings
from django.utils import timezone
from datetime import timedelta
from concurrent.futures import ThreadPoolExecutor
import os
import csv
import json
import logging
import re  # ← ADDED: Import re module for regex operations
import requests
import warnings
import socket
from typing import Dict, Any, List
from urllib.parse import urlparse
from urllib3.exceptions import InsecureRequestWarning

# Disable SSL verification warnings
warnings.filterwarnings('ignore', message='Unverified HTTPS request')
warnings.simplefilter('ignore', InsecureRequestWarning)
from .models import TechnologyCheck
from .serializers import TechnologyCheckSerializer


from .models import (
    Domain, Asset, Scan, ScanAsset, Finding, ScanFinding, CVE, FindingCVE,
    FrontendLibraryCheck, SSLTLSCheck, EmailSecurityCheck,
    SecurityHeaderCheck, DNSSecurityCheck, ReportSummary, finalize_stale_scans
)
from .permissions import OrganisationScopedMixin, get_user_organisation, is_platform_admin, requested_org_id, IsPlatformAdmin
from .audit import log_action
from .serializers import (
    DomainSerializer, AssetSerializer, ScanListSerializer, ScanDetailSerializer,
    ScanCreateSerializer, FindingSerializer, FrontendLibraryCheckSerializer,
    SSLTLSCheckSerializer, EmailSecurityCheckSerializer,
    SecurityHeaderCheckSerializer, DNSSecurityCheckSerializer,
    ScanResultSerializer, ScanSummarySerializer, ReportSummarySerializer
)

# PDF/DOCX generation imports
from io import BytesIO
from django.views import View
from django.http import FileResponse
from rest_framework.views import APIView
from rest_framework.permissions import AllowAny, IsAuthenticated
import tempfile
try:
    from reportlab.pdfgen import canvas
except ImportError:
    canvas = None
try:
    import docx
except ImportError:
    docx = None


def _domain_to_org_name(root_domain):
    """Derive a display-friendly organisation name from a root domain label."""
    if not root_domain:
        return None
    first_label = root_domain.split('.')[0]
    # Short labels (≤6 chars) are typically acronyms → uppercase (e.g. nitda → NITDA)
    if len(first_label) <= 6:
        return first_label.upper()
    return first_label.replace('-', ' ').replace('_', ' ').title()


# ============================================
# REPORT SUMMARY ENDPOINT
# ============================================
class ReportSummaryDownloadView(APIView):

    def get(self, request, format=None):
        """
        Download a report summary as PDF or DOCX.
        Query params: ?id=<report_id>&format=pdf|docx
        """
        report_id = request.GET.get('id')
        file_format = request.GET.get('format', 'pdf').lower()
        summary = get_object_or_404(ReportSummary, id=report_id)

        if file_format == 'pdf':
            if not canvas:
                return Response({'error': 'reportlab not installed'}, status=500)
            buffer = BytesIO()
            p = canvas.Canvas(buffer)
            p.setFont("Helvetica", 14)
            p.drawString(100, 800, f"Report Summary for {summary.domain}")
            p.setFont("Helvetica", 12)
            p.drawString(100, 780, f"Scan Date: {summary.scan_date}")
            p.drawString(100, 760, f"Total Findings: {summary.total_findings}")
            p.drawString(100, 740, f"High Risk: {summary.high_risk}")
            p.drawString(100, 720, f"Medium Risk: {summary.medium_risk}")
            p.drawString(100, 700, f"Low Risk: {summary.low_risk}")
            if summary.notes:
                p.drawString(100, 680, f"Notes: {summary.notes[:80]}")
            p.showPage()
            p.save()
            buffer.seek(0)
            return FileResponse(buffer, as_attachment=True, filename=f"report_summary_{summary.id}.pdf")

        elif file_format == 'docx':
            if not docx:
                return Response({'error': 'python-docx not installed'}, status=500)
            doc = docx.Document()
            doc.add_heading(f"Report Summary for {summary.domain}", 0)
            doc.add_paragraph(f"Scan Date: {summary.scan_date}")
            doc.add_paragraph(f"Total Findings: {summary.total_findings}")
            doc.add_paragraph(f"High Risk: {summary.high_risk}")
            doc.add_paragraph(f"Medium Risk: {summary.medium_risk}")
            doc.add_paragraph(f"Low Risk: {summary.low_risk}")
            if summary.notes:
                doc.add_paragraph(f"Notes: {summary.notes}")
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(tmp.name)
            tmp.seek(0)
            response = FileResponse(open(tmp.name, 'rb'), as_attachment=True, filename=f"report_summary_{summary.id}.docx")
            return response

        return Response({'error': 'Invalid format'}, status=400)

# Import detection and scanning modules
from .nuclei_runner import run_nuclei_scan, get_template_path
from .cve_checker import (
    check_library_vulnerabilities,
    fetch_latest_library_version,
    get_ecosystem_for_library,
    supports_latest_version_lookup,
    version_compare,
)
from .library_detector import detect_technologies
from .org_extractor import extract_organization_name
from .report_style_config import (
    generate_docx_report,
    generate_pdf_report,
)

logger = logging.getLogger(__name__)


def _augment_library_detections_with_analytics(tech_result, detected_libraries):
    """Promote selected analytics detections into library entries for library checks."""
    libraries = list(detected_libraries or [])
    seen_names = {(item.get('name') or '').strip().lower() for item in libraries}

    for analytics_item in tech_result.get('analytics', []) or []:
        analytics_name = (analytics_item.get('name') or '').strip().lower()
        if analytics_name == 'cloudflare beacon' and 'cloudflare beacon' not in seen_names:
            libraries.append(
                {
                    'name': 'Cloudflare Beacon',
                    'version': analytics_item.get('version', 'Unknown') or 'Unknown',
                    'source': 'analytics_fallback',
                    'confidence': analytics_item.get('confidence', 'medium'),
                }
            )
            seen_names.add('cloudflare beacon')

    return libraries


def _normalize_target_url(target: str) -> str:
    """Ensure target has a scheme for HTTP probing."""
    if target.startswith(('http://', 'https://')):
        return target
    return f'https://{target}'


def _check_target_is_active(target: str) -> Dict[str, Any]:
    """Check if a target is resolvable and reachable over HTTP(S)."""
    target_url = _normalize_target_url(target)
    parsed = urlparse(target_url)
    hostname = parsed.hostname

    if not hostname:
        return {
            'active': False,
            'target': target,
            'target_url': target_url,
            'reason': 'Invalid target format',
        }

    try:
        socket.getaddrinfo(hostname, None)
    except socket.gaierror:
        return {
            'active': False,
            'target': target,
            'target_url': target_url,
            'reason': f'DNS resolution failed for {hostname}',
        }

    try:
        requests.head(target_url, timeout=10, verify=False, allow_redirects=True)
        return {
            'active': True,
            'target': target,
            'target_url': target_url,
            'reason': '',
        }
    except requests.exceptions.RequestException:
        try:
            requests.get(target_url, timeout=10, verify=False, allow_redirects=True)
            return {
                'active': True,
                'target': target,
                'target_url': target_url,
                'reason': '',
            }
        except requests.exceptions.RequestException as exc:
            return {
                'active': False,
                'target': target,
                'target_url': target_url,
                'reason': str(exc),
            }


def _validate_assets_are_active(assets) -> List[Dict[str, Any]]:
    """Return a list of inactive assets with reasons."""
    inactive_assets = []

    for asset in assets:
        check = _check_target_is_active(asset.value)
        if not check['active']:
            inactive_assets.append({
                'asset_id': asset.id,
                'target': asset.value,
                'target_url': check['target_url'],
                'reason': check['reason'],
            })

    return inactive_assets


def _split_active_assets(assets):
    """Split assets into active and inactive buckets based on reachability checks."""
    asset_list = list(assets)
    inactive_assets = _validate_assets_are_active(asset_list)
    inactive_ids = {item['asset_id'] for item in inactive_assets}
    active_assets = [asset for asset in asset_list if asset.id not in inactive_ids]
    return active_assets, inactive_assets


def _asset_value_candidates(raw_value: str) -> List[str]:
    """Build normalized value candidates used to match existing assets."""
    value = (raw_value or '').strip()
    if not value:
        return []

    candidates = set()
    candidates.add(value.lower())
    candidates.add(value.rstrip('/').lower())

    parsed = urlparse(value if value.startswith(('http://', 'https://')) else f'https://{value}')
    host = (parsed.netloc or parsed.path or '').strip().lower().rstrip('/')
    path = (parsed.path or '').strip().rstrip('/')

    if host:
        candidates.add(host)
        # Always include both URL-prefixed and bare-domain forms so the lookup
        # matches regardless of how the value was originally stored.
        candidates.add(f'https://{host}')
        candidates.add(f'http://{host}')
        if host.startswith('www.'):
            bare = host[4:]
            candidates.add(bare)
            candidates.add(f'https://{bare}')
            candidates.add(f'http://{bare}')

        if path and path != '/':
            candidates.add(f"{host}{path}".lower())
            candidates.add(f"https://{host}{path}".lower())
            candidates.add(f"http://{host}{path}".lower())
            if host.startswith('www.'):
                candidates.add(f"{host[4:]}{path}".lower())
                candidates.add(f"https://{host[4:]}{path}".lower())

    return [candidate for candidate in candidates if candidate]


class StandardPagination(PageNumberPagination):
    page_size = 10
    page_size_query_param = 'page_size'
    max_page_size = 100


def _scoped_scan(request, scan_id):
    """Fetch a scan by id, scoped to the caller's organisation.

    Platform admins have full cross-org access, so they can fetch any scan.
    """
    if is_platform_admin(getattr(request, 'user', None)):
        return get_object_or_404(Scan, id=scan_id)
    org = get_user_organisation(request)
    return get_object_or_404(Scan, id=scan_id, organisation=org)


def _scoped_domain(request, domain_id):
    """Fetch a domain by id, scoped to the caller's organisation.

    Platform admins have full cross-org access, so they can fetch any domain.
    """
    if is_platform_admin(getattr(request, 'user', None)):
        return get_object_or_404(Domain, id=domain_id)
    org = get_user_organisation(request)
    return get_object_or_404(Domain, id=domain_id, organisation=org)


# ============================================
# SUBDOMAIN DISCOVERY
# ============================================

_DISCOVERY_HEADERS = {
    'User-Agent': 'Mozilla/5.0 (domainscan subdomain-discovery)',
    'Accept': 'application/json, text/plain, */*',
}


def _clean_host(raw: str, root_lower: str) -> str:
    """Normalise a candidate hostname; return '' if it should be discarded."""
    if not raw:
        return ''
    h = raw.strip().lower().lstrip('*.').rstrip('.')
    if '://' in h:
        h = h.split('://', 1)[1]
    h = h.split('/', 1)[0].split(':', 1)[0]
    if not h or h == root_lower:
        return ''
    if not h.endswith('.' + root_lower):
        return ''
    if any(c in h for c in (' ', '@', '?', '#', ',')):
        return ''
    return h


# --- Individual sources. Each returns a list of hostnames or [] on any error. ---

def _src_sublist3r(root_domain: str) -> List[str]:
    try:
        import sublist3r
        return sublist3r.main(
            domain=root_domain, threads=40, savefile=None, ports=None,
            silent=True, verbose=False, enable_bruteforce=False, engines=None,
        ) or []
    except Exception as exc:
        logging.warning("sublist3r failed for %s: %s", root_domain, exc)
        return []


def _src_crtsh(root_domain: str) -> List[str]:
    try:
        r = requests.get(
            f"https://crt.sh/?q=%25.{root_domain}&output=json",
            headers=_DISCOVERY_HEADERS, timeout=25,
        )
        r.raise_for_status()
        out = []
        for entry in r.json():
            for h in (entry.get('name_value') or '').splitlines():
                out.append(h)
        return out
    except Exception as exc:
        logging.warning("crt.sh failed for %s: %s", root_domain, exc)
        return []


def _src_certspotter(root_domain: str) -> List[str]:
    try:
        r = requests.get(
            f"https://api.certspotter.com/v1/issuances",
            params={
                'domain': root_domain,
                'include_subdomains': 'true',
                'expand': 'dns_names',
            },
            headers=_DISCOVERY_HEADERS, timeout=20,
        )
        r.raise_for_status()
        out = []
        for entry in r.json():
            for h in entry.get('dns_names') or []:
                out.append(h)
        return out
    except Exception as exc:
        logging.warning("certspotter failed for %s: %s", root_domain, exc)
        return []


def _src_hackertarget(root_domain: str) -> List[str]:
    try:
        r = requests.get(
            f"https://api.hackertarget.com/hostsearch/?q={root_domain}",
            headers=_DISCOVERY_HEADERS, timeout=15,
        )
        r.raise_for_status()
        text = r.text or ''
        if 'API count exceeded' in text or 'error' in text.lower()[:30]:
            return []
        # CSV: host,ip
        return [line.split(',', 1)[0] for line in text.splitlines() if line.strip()]
    except Exception as exc:
        logging.warning("hackertarget failed for %s: %s", root_domain, exc)
        return []


def _src_alienvault(root_domain: str) -> List[str]:
    try:
        r = requests.get(
            f"https://otx.alienvault.com/api/v1/indicators/domain/{root_domain}/passive_dns",
            headers=_DISCOVERY_HEADERS, timeout=20,
        )
        r.raise_for_status()
        data = r.json() or {}
        return [rec.get('hostname', '') for rec in (data.get('passive_dns') or [])]
    except Exception as exc:
        logging.warning("alienvault failed for %s: %s", root_domain, exc)
        return []


def _src_anubis(root_domain: str) -> List[str]:
    try:
        r = requests.get(
            f"https://jldc.me/anubis/subdomains/{root_domain}",
            headers=_DISCOVERY_HEADERS, timeout=15,
        )
        r.raise_for_status()
        return r.json() or []
    except Exception as exc:
        logging.warning("anubis failed for %s: %s", root_domain, exc)
        return []


def _src_urlscan(root_domain: str) -> List[str]:
    try:
        r = requests.get(
            f"https://urlscan.io/api/v1/search/",
            params={'q': f'domain:{root_domain}', 'size': 10000},
            headers=_DISCOVERY_HEADERS, timeout=20,
        )
        r.raise_for_status()
        out = []
        for res in (r.json() or {}).get('results') or []:
            host = ((res.get('page') or {}).get('domain') or '')
            if host:
                out.append(host)
            task_url = ((res.get('task') or {}).get('url') or '')
            if task_url:
                out.append(task_url)
        return out
    except Exception as exc:
        logging.warning("urlscan failed for %s: %s", root_domain, exc)
        return []


def _src_wayback(root_domain: str) -> List[str]:
    try:
        r = requests.get(
            "https://web.archive.org/cdx/search/cdx",
            params={
                'url': f'*.{root_domain}/*',
                'output': 'json',
                'fl': 'original',
                'collapse': 'urlkey',
                'limit': 10000,
            },
            headers=_DISCOVERY_HEADERS, timeout=25,
        )
        r.raise_for_status()
        rows = r.json() or []
        # first row is a header
        return [row[0] for row in rows[1:] if row]
    except Exception as exc:
        logging.warning("wayback failed for %s: %s", root_domain, exc)
        return []


_DISCOVERY_SOURCES = [
    ('sublist3r', _src_sublist3r),
    ('crt.sh', _src_crtsh),
    ('certspotter', _src_certspotter),
    ('hackertarget', _src_hackertarget),
    ('alienvault', _src_alienvault),
    ('anubis', _src_anubis),
    ('urlscan', _src_urlscan),
    ('wayback', _src_wayback),
]


def _discover_subdomains_aggregate(root_domain: str) -> Dict[str, Dict[str, Any]]:
    """Query every passive source in parallel, merge & dedupe. Returns a dict:
    {hostname: {'sources': ['crt.sh', 'sublist3r', ...]}}. Sources that fail
    are silently skipped — having 8 means losing one or two is harmless."""
    from concurrent.futures import ThreadPoolExecutor, as_completed

    root_lower = root_domain.lower().rstrip('.')
    aggregated: Dict[str, Dict[str, Any]] = {}

    with ThreadPoolExecutor(max_workers=len(_DISCOVERY_SOURCES)) as pool:
        futures = {pool.submit(fn, root_domain): name for name, fn in _DISCOVERY_SOURCES}
        for fut in as_completed(futures):
            source_name = futures[fut]
            try:
                hosts = fut.result() or []
            except Exception as exc:
                logging.warning("source %s raised: %s", source_name, exc)
                continue
            for raw in hosts:
                host = _clean_host(raw, root_lower)
                if not host:
                    continue
                if host not in aggregated:
                    aggregated[host] = {'sources': []}
                if source_name not in aggregated[host]['sources']:
                    aggregated[host]['sources'].append(source_name)
    return aggregated


def _resolve_host_active(host: str, timeout: float = 3.0):
    """Return the resolved IPv4 address if the host has an A record, else None."""
    socket.setdefaulttimeout(timeout)
    try:
        ip = socket.gethostbyname(host)
        return ip
    except (socket.gaierror, socket.timeout, OSError):
        return None
    finally:
        socket.setdefaulttimeout(None)


def discover_active_subdomains(root_domain: str, max_hosts: int = 2000,
                                 resolve_workers: int = 64, timeout: float = 3.0):
    """Discover candidates via 8 parallel passive sources (Sublist3r + 7 online
    APIs), DNS-resolve in parallel, return only the ones with an A record.
    Each entry: {'value': host, 'ip_address': ip, 'source': 'a,b,c'} where
    'source' is a comma-separated list of every aggregator that saw the host."""
    from concurrent.futures import ThreadPoolExecutor, as_completed

    aggregated = _discover_subdomains_aggregate(root_domain)
    if not aggregated:
        return []

    # Cap to avoid blasting DNS resolver on huge result sets.
    candidates = sorted(aggregated.keys())[:max_hosts]

    active = []
    with ThreadPoolExecutor(max_workers=resolve_workers) as pool:
        futures = {pool.submit(_resolve_host_active, h, timeout): h for h in candidates}
        for fut in as_completed(futures):
            host = futures[fut]
            ip = fut.result()
            if ip:
                src_list = aggregated[host]['sources']
                active.append({
                    'value': host,
                    'ip_address': ip,
                    'source': ','.join(sorted(src_list)),
                })
    active.sort(key=lambda x: x['value'])
    return active


# ============================================
# DOMAIN & ASSET VIEWSETS
# ============================================

class DomainViewSet(OrganisationScopedMixin, viewsets.ModelViewSet):
    """
    ViewSet for managing domains (scoped to the caller's organisation)
    """
    queryset = Domain.objects.all()
    serializer_class = DomainSerializer
    pagination_class = StandardPagination

    def get_permissions(self):
        # Domains are assigned by the platform admin only. Org users get read
        # access plus the discover-subdomains / posture actions, but cannot
        # create, edit, delete, or re-assign the primary domain.
        if self.action in ('create', 'update', 'partial_update', 'destroy', 'set_primary'):
            return [IsAuthenticated(), IsPlatformAdmin()]
        return [IsAuthenticated()]

    def _apply_primary(self, domain):
        """Enforce a single primary domain per organisation."""
        if domain.is_primary:
            Domain.objects.filter(
                organisation=domain.organisation, is_primary=True
            ).exclude(pk=domain.pk).update(is_primary=False)

    def _ensure_apex_asset(self, domain):
        """Auto-create the apex (root_domain) asset so the domain is scannable."""
        Asset.objects.get_or_create(
            domain=domain, value=domain.root_domain,
            defaults={'asset_type': 'root_domain', 'source': 'platform-admin'},
        )

    def perform_create(self, serializer):
        obj = serializer.save(organisation=self.get_organisation())
        self._apply_primary(obj)
        self._ensure_apex_asset(obj)
        log_action('domain.create', request=self.request,
                   organisation=obj.organisation, target=obj.root_domain,
                   is_primary=obj.is_primary)

    def perform_update(self, serializer):
        obj = serializer.save()
        self._apply_primary(obj)

    @action(detail=True, methods=['post'], url_path='set-primary')
    def set_primary(self, request, pk=None):
        """Mark this domain as the organisation's primary (platform admin only)."""
        domain = self.get_object()
        Domain.objects.filter(
            organisation=domain.organisation, is_primary=True
        ).exclude(pk=domain.pk).update(is_primary=False)
        domain.is_primary = True
        domain.save(update_fields=['is_primary'])
        log_action('domain.set_primary', request=request,
                   organisation=domain.organisation, target=domain.root_domain)
        return Response(DomainSerializer(domain).data)

    @action(detail=True, methods=['post'], url_path='discover-subdomains')
    def discover_subdomains(self, request, pk=None):
        """Discover active subdomains for a domain and auto-populate assets.

        POST /api/domains/{id}/discover-subdomains/

        Optional JSON body:
          {
            "max_hosts": 500,        # cap candidates from crt.sh (default 500)
            "resolve_timeout": 3.0,  # per-host DNS timeout in seconds (default 3.0)
            "dry_run": false         # if true, just return discovery without DB writes
          }

        Returns:
          {
            "domain": "<root_domain>",
            "candidates_found": <int>,   # total from crt.sh after dedupe
            "active": <int>,             # number with an A record
            "created": <int>,            # new Asset rows inserted
            "skipped_existing": <int>,   # already in Asset table
            "subdomains": [
              {"value": "...", "ip_address": "...", "source": "crt.sh", "created": true/false}
            ]
          }
        """
        domain = self.get_object()
        body = request.data if isinstance(request.data, dict) else {}
        max_hosts = int(body.get('max_hosts') or 500)
        resolve_timeout = float(body.get('resolve_timeout') or 3.0)
        dry_run = bool(body.get('dry_run'))

        active = discover_active_subdomains(
            domain.root_domain,
            max_hosts=max_hosts,
            timeout=resolve_timeout,
        )

        existing = set(
            domain.assets.filter(asset_type='subdomain')
            .values_list('value', flat=True)
        )

        results = []
        created_count = 0
        skipped = 0
        for entry in active:
            host = entry['value']
            already = host in existing
            if already:
                skipped += 1
            elif not dry_run:
                try:
                    Asset.objects.create(
                        domain=domain,
                        asset_type='subdomain',
                        value=host,
                        source=entry['source'],
                        ip_address=entry['ip_address'],
                        last_verified=timezone.now(),
                    )
                    created_count += 1
                except IntegrityError:
                    skipped += 1
                    already = True
            results.append({
                'value': host,
                'ip_address': entry['ip_address'],
                'source': entry['source'],
                'created': not already and not dry_run,
            })

        return Response({
            'domain': domain.root_domain,
            'candidates_found': len(active),
            'active': len(active),
            'created': created_count,
            'skipped_existing': skipped,
            'dry_run': dry_run,
            'subdomains': results,
        }, status=status.HTTP_200_OK)

    @action(detail=True, methods=['get'])
    def posture(self, request, pk=None):
        """
        Get comprehensive security posture for a domain
        GET /api/domains/{id}/posture/
        
        Returns:
        - Domain info
        - Latest scan summary
        - Issues/findings organized by severity and category
        - Email, SSL, DNS, Headers status across all assets
        - List of scanned assets
        """
        domain = self.get_object()
        
        # Get the latest scan for this domain (across all its assets)
        latest_scan = Scan.objects.filter(
            scan_assets__asset__domain=domain,
            status='completed'
        ).order_by('-finished_at').first()
        
        if not latest_scan:
            return Response({
                'domain_id': domain.id,
                'domain': domain.root_domain,
                'status': 'no_scans',
                'message': 'No completed scans found for this domain',
                'last_scan_at': None,
                'last_scan_date': None,
                'last_scan_time': None,
                'assets': []
            }, status=status.HTTP_200_OK)
        
        # Get all assets under this domain
        assets = domain.assets.all()
        
        # Collect all findings for this domain from the latest scan
        findings = Finding.objects.filter(
            asset__domain=domain,
            finding_scans__scan=latest_scan
        ).select_related('asset')
        
        # Aggregate statistics
        finding_counts = {
            'total': findings.count(),
            'critical': findings.filter(risk_rating='Critical').count(),
            'high': findings.filter(risk_rating='High').count(),
            'medium': findings.filter(risk_rating='Medium').count(),
            'low': findings.filter(risk_rating='Low').count(),
        }
        
        # Group findings by category and severity
        findings_by_category = {}
        for category in ['SSL', 'DNS', 'Email', 'CVE', 'Misconfiguration']:
            category_findings = findings.filter(category=category)
            if category_findings.exists():
                findings_by_category[category] = {
                    'count': category_findings.count(),
                    'by_severity': {
                        'critical': category_findings.filter(risk_rating='Critical').count(),
                        'high': category_findings.filter(risk_rating='High').count(),
                        'medium': category_findings.filter(risk_rating='Medium').count(),
                        'low': category_findings.filter(risk_rating='Low').count(),
                    },
                    'issues': FindingSerializer(
                        category_findings[:10],  # Limit to top 10 per category
                        many=True
                    ).data
                }
        
        # Get email security checks status
        email_checks = EmailSecurityCheck.objects.filter(
            asset__domain=domain,
            scan=latest_scan
        )
        email_status = {
            'spf': self._get_email_status(email_checks, 'SPF'),
            'dkim': self._get_email_status(email_checks, 'DKIM'),
            'dmarc': self._get_email_status(email_checks, 'DMARC'),
            'total_checks': email_checks.count()
        }
        
        # Get SSL checks status
        ssl_checks = SSLTLSCheck.objects.filter(
            asset__domain=domain,
            scan=latest_scan
        )
        ssl_passed = ssl_checks.filter(risk_rating='Low').count()
        ssl_failed = ssl_checks.exclude(risk_rating='Low').count()
        ssl_status = {
            'passed': ssl_passed,
            'failed': ssl_failed,
            'score': 100 if not ssl_checks else round(100 * (ssl_passed / ssl_checks.count())),
            'total_checks': ssl_checks.count()
        }
        
        # Get DNS checks status
        dns_checks = DNSSecurityCheck.objects.filter(
            asset__domain=domain,
            scan=latest_scan
        )
        dns_status = {
            'total_checks': dns_checks.count(),
            'issues': dns_checks.exclude(risk_rating='Low').count()
        }
        
        # Get security headers status
        header_checks = SecurityHeaderCheck.objects.filter(
            asset__domain=domain,
            scan=latest_scan
        )
        headers_present = header_checks.filter(status='present').count()
        headers_missing = header_checks.filter(status='missing').count()
        header_status = {
            'present': headers_present,
            'missing': headers_missing,
            'total': header_checks.count()
        }
        
        # Get library checks status
        library_checks = FrontendLibraryCheck.objects.filter(
            asset__domain=domain,
            scan=latest_scan
        )
        library_status = {
            'up_to_date': library_checks.filter(vulnerability_status='up-to-date').count(),
            'outdated': library_checks.filter(vulnerability_status='outdated').count(),
            'vulnerable': library_checks.filter(vulnerability_status='vulnerable').count(),
            'total': library_checks.count()
        }
        
        # Determine overall risk rating
        risk_ratings = [f.risk_rating for f in findings]
        risk_priority = {'Critical': 0, 'High': 1, 'Medium': 2, 'Low': 3}
        overall_risk = min(risk_ratings, key=lambda x: risk_priority.get(x, 999)) if risk_ratings else 'Low'

        # Use a safe fallback in case finished_at is not populated for historical scans.
        last_scan_at = latest_scan.finished_at or latest_scan.started_at or latest_scan.created_at
        last_scan_date = last_scan_at.date().isoformat() if last_scan_at else None
        last_scan_time = last_scan_at.time().replace(microsecond=0).isoformat() if last_scan_at else None
        
        posture_data = {
            'domain_id': domain.id,
            'domain': domain.root_domain,
            'status': latest_scan.status,
            'last_scan': last_scan_at,
            'last_scan_at': last_scan_at,
            'last_scan_date': last_scan_date,
            'last_scan_time': last_scan_time,
            'scan_duration_seconds': latest_scan.duration_seconds,
            'overall_risk_rating': overall_risk,
            'findings_summary': finding_counts,
            'findings_by_category': findings_by_category,
            'email_security': email_status,
            'ssl_tls': ssl_status,
            'dns_security': dns_status,
            'security_headers': header_status,
            'frontend_libraries': library_status,
            'assets_scanned': [
                {
                    'id': asset.id,
                    'value': asset.value,
                    'type': asset.asset_type,
                    'issues_found': findings.filter(asset=asset).count()
                }
                for asset in assets
            ]
        }
        
        return Response(posture_data, status=status.HTTP_200_OK)
    
    def _get_email_status(self, checks, check_type):
        """Helper to get email check status"""
        check = checks.filter(check_type=check_type).first()
        if not check:
            return 'not_checked'
        return check.status.lower() if check.status else 'unknown'


class AssetViewSet(OrganisationScopedMixin, viewsets.ModelViewSet):
    """
    ViewSet for managing assets (scoped to the caller's organisation)
    """
    queryset = Asset.objects.all()
    serializer_class = AssetSerializer
    pagination_class = StandardPagination

    def get_queryset(self):
        """Filter assets by the caller's organisation (platform admins: all orgs),
        then by domain when requested."""
        queryset = self.scope_queryset(Asset.objects.select_related('domain'))

        domain = self.request.GET.get('domain') or self.request.GET.get('domain_id')
        if domain:
            if not str(domain).isdigit():
                raise ValidationError({'domain': 'domain must be a numeric id'})
            queryset = queryset.filter(domain_id=int(domain))

        return queryset

    def perform_create(self, serializer):
        # Stamp the asset with its domain's organisation. For platform admins
        # the domain may belong to any org (they are unrestricted); for regular
        # users the create() guard already confirmed the domain is in their org.
        domain = serializer.validated_data.get('domain')
        if domain is not None and is_platform_admin(getattr(self.request, 'user', None)):
            serializer.save(organisation=domain.organisation)
        else:
            serializer.save(organisation=self.get_organisation())

    def create(self, request, *args, **kwargs):
        """Create an asset, but reuse an existing one if it already exists for the domain."""
        domain_id = request.data.get('domain')
        raw_value = request.data.get('value', '')

        # Resolve the organisation the asset belongs to. Platform admins are
        # unrestricted: they can create assets under any organisation's domain,
        # and the asset is stamped with that domain's org. Regular users are
        # scoped to their own organisation.
        if is_platform_admin(getattr(request, 'user', None)):
            domain = Domain.objects.filter(id=domain_id).first() if domain_id else None
            if domain_id and domain is None:
                raise ValidationError({'domain': 'Domain not found.'})
            org = domain.organisation if domain else self.get_organisation()
        else:
            org = self.get_organisation()
            if domain_id and not Domain.objects.filter(id=domain_id, organisation=org).exists():
                raise ValidationError({'domain': 'Domain not found in your organisation.'})

        if domain_id and raw_value:
            candidates = _asset_value_candidates(raw_value)
            if candidates:
                query = Q()
                for candidate in candidates:
                    query |= Q(value__iexact=candidate)

                existing = Asset.objects.filter(
                    organisation=org, domain_id=domain_id
                ).filter(query).order_by('id').first()
                if existing:
                    serializer = self.get_serializer(existing)
                    return Response(serializer.data, status=status.HTTP_200_OK)

        try:
            return super().create(request, *args, **kwargs)
        except IntegrityError:
            # Race condition: another request created the same asset between our
            # check above and this insert. Return the row that won the race.
            if domain_id and raw_value:
                candidates = _asset_value_candidates(raw_value)
                query = Q()
                for candidate in candidates:
                    query |= Q(value__iexact=candidate)
                existing = Asset.objects.filter(
                    organisation=org, domain_id=domain_id
                ).filter(query).order_by('id').first()
                if existing:
                    serializer = self.get_serializer(existing)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            raise


# ============================================
# BACKGROUND SCAN EXECUTION
# ============================================
#
# A scan runs the Nuclei binary plus header/DNS/email/library checks across
# every asset and takes minutes. Running it inside the HTTP request held a web
# worker the whole time, so a handful of concurrent scans could starve the API.
# We hand execution to a bounded background worker pool instead: the endpoint
# enqueues the job and returns immediately, and the client polls scan status.
#
# `submit_scan_job` is the single seam between "API request" and "scan runs".
# To scale beyond one host, swap its body to enqueue a Celery task (the worker
# would just call `ScanViewSet()._run_scan_job(scan, params)` the same way).

_SCAN_WORKER_THREADS = getattr(settings, 'SCAN_WORKER_THREADS', None) or int(
    os.getenv('SCAN_WORKER_THREADS', '2')
)
_scan_pool = ThreadPoolExecutor(
    max_workers=_SCAN_WORKER_THREADS, thread_name_prefix='scan-worker'
)


def _scan_job_worker(scan_id, params):
    """Run one scan in a background thread, with its own DB connection.

    The pipeline persists all state on the Scan row, so the worker only needs to
    reload the scan, run it, and make sure a crash still finalises the row
    instead of leaving it stuck in 'running'.
    """
    try:
        scan = Scan.objects.get(id=scan_id)
        ScanViewSet()._run_scan_job(scan, params)
    except Exception:
        logger.exception("Background scan job %s crashed", scan_id)
        try:
            scan = Scan.objects.get(id=scan_id)
            if scan.status in ('queued', 'running'):
                scan.status = 'failed'
                scan.error_message = 'Scan worker crashed before completion.'
                scan.finished_at = timezone.now()
                if scan.started_at:
                    scan.duration_seconds = int(
                        (scan.finished_at - scan.started_at).total_seconds()
                    )
                scan.save(update_fields=[
                    'status', 'error_message', 'finished_at', 'duration_seconds'
                ])
        except Exception:
            logger.exception("Failed to finalise crashed scan %s", scan_id)
    finally:
        # Django opens a per-thread connection on first query; close it so the
        # pool's long-lived threads don't hold connections open indefinitely.
        connection.close()


def submit_scan_job(scan_id, params):
    """Enqueue a queued scan for background execution and return immediately."""
    _scan_pool.submit(_scan_job_worker, scan_id, params)


# ============================================
# SCAN MANAGEMENT VIEWSET
# ============================================

class ScanViewSet(OrganisationScopedMixin, viewsets.ModelViewSet):
    """
    ViewSet for managing scans (scoped to the caller's organisation)
    """
    queryset = Scan.objects.all()
    pagination_class = StandardPagination

    def get_queryset(self):
        """
        Filter scans based on query parameters (already org-scoped by the mixin)
        """
        # Auto-finalize orphaned scans (e.g. left 'running' by a server restart)
        # so the UI never shows a permanently-running scan.
        if is_platform_admin(self.request.user):
            finalize_stale_scans()  # sweep across all organisations
        else:
            finalize_stale_scans(organisation=self.get_organisation())
        queryset = super().get_queryset()
        
        # Filter by domain_id: returns scans that scanned assets belonging to this domain
        domain_id = self.request.GET.get('domain_id') or self.request.GET.get('domain')
        if domain_id:
            if not str(domain_id).isdigit():
                raise ValidationError({'domain_id': 'domain_id must be a numeric id'})
            queryset = queryset.filter(
                scan_assets__asset__domain_id=int(domain_id)
            ).distinct()
        
        # Filter by asset_id: returns scans that scanned this specific asset
        asset_id = self.request.GET.get('asset_id')
        if asset_id:
            if not str(asset_id).isdigit():
                raise ValidationError({'asset_id': 'asset_id must be a numeric id'})
            queryset = queryset.filter(
                scan_assets__asset_id=int(asset_id)
            ).distinct()
        
        return queryset
    
    def get_serializer_class(self):
        if self.action == 'list':
            return ScanListSerializer
        elif self.action == 'create':
            return ScanCreateSerializer
        elif self.action == 'retrieve':
            return ScanResultSerializer
        return ScanDetailSerializer
    
    def create(self, request, *args, **kwargs):
        """
        Create and initiate a new scan
        POST /api/scans/
        Body: {
            "asset_ids": [1, 2, 3],
            "scan_type": "on-demand",
            "template_categories": ["ssl", "dns", "email"]
        }
        """
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        
        asset_ids = serializer.validated_data['asset_ids']
        scan_type = serializer.validated_data.get('scan_type', 'on-demand')
        template_categories = serializer.validated_data.get('template_categories', [])

        # Resolve assets and the scan's organisation. Platform admins are
        # unrestricted: they can scan any organisation's assets across one or
        # more domains. The scan is stamped with the org that owns those assets.
        # A single scan record holds one org (Scan.organisation is required), so
        # a cross-org batch is rejected with a clear message rather than silently
        # dropping targets. Regular users are scoped to their own organisation.
        if is_platform_admin(getattr(request, 'user', None)):
            assets = Asset.objects.filter(id__in=asset_ids)
            if assets.count() != len(set(asset_ids)):
                return Response(
                    {'error': 'One or more asset IDs not found'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            org_ids = set(assets.values_list('organisation_id', flat=True))
            if len(org_ids) > 1:
                return Response(
                    {'error': 'A single scan cannot span multiple organisations. '
                              'Submit one scan per organisation.'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            org = assets.first().organisation
        else:
            org = self.get_organisation()
            assets = Asset.objects.filter(id__in=asset_ids, organisation=org)
            if assets.count() != len(set(asset_ids)):
                return Response(
                    {'error': 'One or more asset IDs not found in your organisation'},
                    status=status.HTTP_400_BAD_REQUEST
                )

        # Create scan
        scan = Scan.objects.create(
            organisation=org,
            scan_type=scan_type,
            status='queued',
            initiated_by=request.user.id if request.user.is_authenticated else None
        )
        
        # Link assets to scan
        for asset in assets:
            ScanAsset.objects.create(scan=scan, asset=asset)

        log_action('scan.create', request=request, organisation=org,
                   target=f"scan #{scan.id}", scan_id=scan.id, asset_count=len(assets))

        return Response(
            ScanDetailSerializer(scan).data,
            status=status.HTTP_201_CREATED
        )

    @action(detail=True, methods=['post'])
    def cancel(self, request, pk=None):
        """
        Request cancellation for a queued/running scan.
        POST /api/scans/{id}/cancel/
        """
        scan = self.get_object()

        if scan.status in ['completed', 'failed', 'cancelled']:
            return Response(
                {
                    'error': f"Cannot cancel scan with status: {scan.status}",
                    'scan_id': scan.id,
                    'status': scan.status,
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        scan.cancel_requested = True

        # If not running yet, mark immediately as cancelled.
        if scan.status == 'queued':
            now = timezone.now()
            scan.status = 'cancelled'
            scan.finished_at = now
            if scan.started_at:
                scan.duration_seconds = int((now - scan.started_at).total_seconds())
            else:
                scan.duration_seconds = 0
            scan.error_message = 'Cancelled by user before execution started.'

        scan.save()

        return Response(
            {
                'message': 'Cancellation requested successfully.',
                'scan_id': scan.id,
                'status': scan.status,
                'effective_status': 'cancelling' if scan.status == 'running' else scan.status,
                'cancel_requested': scan.cancel_requested,
            },
            status=status.HTTP_200_OK,
        )
    
    @action(detail=True, methods=['get'])
    def summary(self, request, pk=None):
        """
        Get summary statistics for a scan using the single source of truth
        GET /api/scans/{id}/summary/
        """
        from .findings_aggregation import get_normalized_findings_and_counts
        scan = self.get_object()
        scan_assets = scan.scan_assets.select_related('asset', 'asset__domain')
        domain_obj = scan_assets[0].asset.domain if scan_assets else None
        asset_obj = None
        findings, summary_counts, summary_total = get_normalized_findings_and_counts(domain_obj, asset_obj, scan)

        summary_data = {
            'scan_id': scan.id,
            'total_assets_scanned': scan.scan_assets.count(),
            'total_findings': summary_total,
            'critical_findings': summary_counts.get('Critical', 0),
            'high_findings': summary_counts.get('High', 0),
            'medium_findings': summary_counts.get('Medium', 0),
            'low_findings': summary_counts.get('Low', 0),
            # Check type counts
            'library_checks_count': scan.library_checks.count(),
            'ssl_checks_count': scan.ssl_checks.count(),
            'email_checks_count': scan.email_checks.count(),
            'header_checks_count': scan.header_checks.count(),
            'dns_checks_count': scan.dns_checks.count(),
            'technology_checks_count': scan.technology_checks.count(),
            # Library statistics
            'libraries_up_to_date': scan.library_checks.filter(
                vulnerability_status='up-to-date'
            ).count(),
            'libraries_outdated': scan.library_checks.filter(
                vulnerability_status='outdated'
            ).count(),
            'libraries_vulnerable': scan.library_checks.filter(
                vulnerability_status='vulnerable'
            ).count(),
            # Email security status
            'spf_status': self._get_email_check_status(scan, 'SPF'),
            'dkim_status': self._get_email_check_status(scan, 'DKIM'),
            'dmarc_status': self._get_email_check_status(scan, 'DMARC'),
            # SSL statistics
            'ssl_issues_found': scan.ssl_checks.filter(
                cvss_score__gte=4.0
            ).count(),
            'weak_ciphers_detected': scan.ssl_checks.filter(
                check_type='cipher', 
                risk_rating__in=['High', 'Critical']
            ).exists(),
            'certificate_expiring_soon': scan.ssl_checks.filter(
                check_type='certificate',
                certificate_days_remaining__lte=30
            ).exists(),
            # Header statistics
            'missing_headers': scan.header_checks.filter(status='missing').count(),
            'present_headers': scan.header_checks.filter(status='present').count(),
        }
        serializer = ScanSummarySerializer(summary_data)
        return Response(serializer.data)
    
    def _get_email_check_status(self, scan, check_type):
        """Helper to get email check status"""
        check = scan.email_checks.filter(check_type=check_type).first()
        return check.status if check else 'NOT_CHECKED'

    def _build_scan_report_payload(self, scan, domain_name=None):
        """Build normalized report payload for export/download using single source of truth."""
        from .findings_aggregation import get_normalized_findings_and_counts
        # Get domain and asset context from scan, optionally scoped to a single domain.
        scan_assets_qs = ScanAsset.objects.filter(scan=scan).select_related('asset', 'asset__domain')
        if domain_name:
            scan_assets_qs = scan_assets_qs.filter(asset__domain__root_domain__iexact=domain_name)

        scan_assets = list(scan_assets_qs)
        if domain_name and not scan_assets:
            raise ValueError(
                f"Domain '{domain_name}' is not part of scan {scan.id}"
            )

        domain_obj = scan_assets[0].asset.domain if scan_assets else None
        asset_obj = None  # For full scan, not asset-specific
        # Get findings, counts, total
        findings, summary_counts, summary_total = get_normalized_findings_and_counts(domain_obj, asset_obj, scan)

        assets = [
            {
                'asset_id': item.asset.id,
                'asset_type': item.asset.asset_type,
                'value': item.asset.value,
                'domain': item.asset.domain.root_domain if item.asset.domain else None,
            }
            for item in scan_assets
        ]

        # Keep checks for compatibility
        library_checks_qs = scan.library_checks.select_related('asset').all()
        technology_checks_qs = scan.technology_checks.select_related('asset').all()
        ssl_checks_qs = scan.ssl_checks.select_related('asset').all()
        email_checks_qs = scan.email_checks.select_related('asset').all()
        header_checks_qs = scan.header_checks.select_related('asset').all()
        dns_checks_qs = scan.dns_checks.select_related('asset').all()

        if domain_obj:
            library_checks_qs = library_checks_qs.filter(asset__domain=domain_obj)
            technology_checks_qs = technology_checks_qs.filter(asset__domain=domain_obj)
            ssl_checks_qs = ssl_checks_qs.filter(asset__domain=domain_obj)
            email_checks_qs = email_checks_qs.filter(asset__domain=domain_obj)
            header_checks_qs = header_checks_qs.filter(asset__domain=domain_obj)
            dns_checks_qs = dns_checks_qs.filter(asset__domain=domain_obj)

        library_checks = [
            {
                'name': item.library_name,
                'detected_version': item.detected_version,
                'latest_version': item.latest_version,
                'status': item.vulnerability_status,
                'risk_rating': item.risk_level,
                'asset': item.asset.value if item.asset else None,
                'recommendation': item.recommendation,
                'checked_at': item.checked_at,
            }
            for item in library_checks_qs.order_by('library_name', 'id')
        ]
        technology_checks = [
            {
                'name': item.technology_name,
                'version': item.version,
                'latest_version': item.latest_version,
                'category': item.category,
                'risk_rating': item.risk_level,
                'asset': item.asset.value if item.asset else None,
                'checked_at': item.created_at,
            }
            for item in technology_checks_qs.order_by('category', 'technology_name', 'id')
        ]
        ssl_checks = [
            {
                'name': item.check_type,
                'status': item.finding,
                'risk_rating': item.risk_rating,
                'cvss_score': item.cvss_score,
                'asset': item.asset.value if item.asset else None,
                'details': item.example,
                'recommendation': item.recommendation,
                'checked_at': item.checked_at,
            }
            for item in ssl_checks_qs.order_by('check_type', 'id')
        ]
        email_checks = [
            {
                'name': item.check_type,
                'status': item.status,
                'risk_rating': item.risk_rating,
                'cvss_score': item.cvss_score,
                'asset': item.asset.value if item.asset else None,
                'details': item.details,
                'recommendation': item.recommendation,
                'checked_at': item.checked_at,
            }
            for item in email_checks_qs.order_by('check_type', 'id')
        ]
        header_checks = [
            {
                'name': item.header,
                'status': item.status,
                'risk_rating': item.risk_rating,
                'cvss_score': item.cvss_score,
                'asset': item.asset.value if item.asset else None,
                'details': item.header_value,
                'recommendation': item.recommendation,
                'checked_at': item.checked_at,
            }
            for item in header_checks_qs.order_by('header', 'id')
        ]
        dns_checks = [
            {
                'name': item.check_type,
                'status': item.finding,
                'risk_rating': item.risk_rating,
                'cvss_score': item.cvss_score,
                'asset': item.asset.value if item.asset else None,
                'details': item.example,
                'recommendation': item.recommendation,
                'checked_at': item.checked_at,
            }
            for item in dns_checks_qs.order_by('check_type', 'id')
        ]

        root_domain = domain_obj.root_domain if domain_obj else None
        summary = {
            'scan_id': scan.id,
            'scan_type': scan.scan_type,
            'status': scan.status,
            'started_at': scan.started_at,
            'finished_at': scan.finished_at,
            'duration_seconds': scan.duration_seconds,
            'total_assets_scanned': len(assets),
            'total_findings': summary_total,
            'critical_findings': summary_counts.get('Critical', 0),
            'high_findings': summary_counts.get('High', 0),
            'medium_findings': summary_counts.get('Medium', 0),
            'low_findings': summary_counts.get('Low', 0),
            'library_checks_count': library_checks_qs.count(),
            'ssl_checks_count': ssl_checks_qs.count(),
            'email_checks_count': email_checks_qs.count(),
            'header_checks_count': header_checks_qs.count(),
            'dns_checks_count': dns_checks_qs.count(),
            'technology_checks_count': technology_checks_qs.count(),
            'domain_name': root_domain,
            'org_name': _domain_to_org_name(root_domain),
        }

        return {
            'generated_at': timezone.now(),
            'summary': summary,
            'assets': assets,
            'findings': findings,
            'checks': {
                'libraries': library_checks,
                'technologies': technology_checks,
                'ssl': ssl_checks,
                'email': email_checks,
                'headers': header_checks,
                'dns': dns_checks,
            },
        }

    def _generate_pdf_report(self, payload):
        return generate_pdf_report(payload)

    def _generate_docx_report(self, payload):
        return generate_docx_report(payload)

    @action(detail=True, methods=['get'], url_path='report/download')
    def download_report(self, request, pk=None):
        """
        Download generated report for a scan.
        GET /api/scans/{id}/report/download/?format=json|csv|pdf|docx
        """
        scan = self.get_object()
        report_format = request.GET.get('format', 'json').lower()

        if report_format not in {'json', 'csv', 'pdf', 'docx'}:
            return Response(
                {'error': "Invalid format. Use 'json', 'csv', 'pdf', or 'docx'."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        payload = self._build_scan_report_payload(scan)
        timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
        org_slug = re.sub(r'[^\w\-]', '_', payload['summary'].get('org_name') or f'scan_{scan.id}')

        if report_format == 'json':
            response = HttpResponse(
                json.dumps(payload, default=str, indent=2),
                content_type='application/json',
            )
            response['Content-Disposition'] = (
                f'attachment; filename="{org_slug}_security_report_{timestamp}.json"'
            )
            return response

        if report_format == 'pdf':
            content, error = self._generate_pdf_report(payload)
            if error:
                return Response({'error': error}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            response = HttpResponse(content, content_type='application/pdf')
            response['Content-Disposition'] = (
                f'attachment; filename="{org_slug}_security_report_{timestamp}.pdf"'
            )
            return response

        if report_format == 'docx':
            content, error = self._generate_docx_report(payload)
            if error:
                return Response({'error': error}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            response = HttpResponse(
                content,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
            response['Content-Disposition'] = (
                f'attachment; filename="{org_slug}_security_report_{timestamp}.docx"'
            )
            return response

        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = (
            f'attachment; filename="{org_slug}_security_report_{timestamp}.csv"'
        )

        writer = csv.writer(response)

        writer.writerow(['scan_id', payload['summary']['scan_id']])
        writer.writerow(['scan_type', payload['summary']['scan_type']])
        writer.writerow(['status', payload['summary']['status']])
        writer.writerow(['started_at', payload['summary']['started_at']])
        writer.writerow(['finished_at', payload['summary']['finished_at']])
        writer.writerow(['duration_seconds', payload['summary']['duration_seconds']])
        writer.writerow(['total_assets_scanned', payload['summary']['total_assets_scanned']])
        writer.writerow(['total_findings', payload['summary']['total_findings']])
        writer.writerow(['critical_findings', payload['summary']['critical_findings']])
        writer.writerow(['high_findings', payload['summary']['high_findings']])
        writer.writerow(['medium_findings', payload['summary']['medium_findings']])
        writer.writerow(['low_findings', payload['summary']['low_findings']])
        writer.writerow([])

        writer.writerow([
            'finding_id',
            'title',
            'category',
            'risk_rating',
            'status',
            'cvss_score',
            'nuclei_template_id',
            'asset',
            'domain',
            'cve_ids',
            'recommendation',
            'first_seen',
            'last_seen',
        ])

        for finding in payload['findings']:
            writer.writerow([
                finding['id'],
                finding['title'],
                finding['category'],
                finding['risk_rating'],
                finding['status'],
                finding['cvss_score'],
                finding['nuclei_template_id'],
                finding['asset'],
                finding['domain'],
                ';'.join(finding['cve_ids']),
                finding['recommendation'],
                finding['first_seen'],
                finding['last_seen'],
            ])

        return response

    # @action(detail=True, methods=['post'])
    @action(detail=True, methods=['post'])
    def execute(self, request, pk=None):
        """
        Execute a queued scan
        POST /api/scans/{id}/execute/
        """
        scan = self.get_object()

        # Respect prior cancellation requests.
        if scan.cancel_requested:
            if scan.status != 'cancelled':
                now = timezone.now()
                scan.status = 'cancelled'
                scan.finished_at = now
                if scan.started_at:
                    scan.duration_seconds = int((now - scan.started_at).total_seconds())
                else:
                    scan.duration_seconds = 0
                if not scan.error_message:
                    scan.error_message = 'Cancelled by user.'
                scan.save()

            return Response(
                {
                    'error': 'Scan execution cancelled before start.',
                    'scan_id': scan.id,
                    'status': scan.status,
                    'cancel_requested': scan.cancel_requested,
                },
                status=status.HTTP_400_BAD_REQUEST,
            )
        
        if scan.status != 'queued':
            return Response(
                {'error': f'Cannot execute scan with status: {scan.status}'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Hand the heavy pipeline to a background worker so this request returns
        # immediately instead of holding a web worker for the minutes a full scan
        # takes. The client polls GET /api/scans/{id}/ for status transitions
        # (queued -> running -> completed/failed/cancelled).
        params = {
            'templates': request.data.get('templates', ['ssl', 'dns', 'email']),
            'check_libraries': request.data.get('check_libraries', True),
            'check_cves': request.data.get('check_cves', True),
            'extract_org': request.data.get('extract_org', True),
        }
        submit_scan_job(scan.id, params)
        return Response(
            {
                'message': 'Scan queued for execution.',
                'scan_id': scan.id,
                'status': scan.status,
            },
            status=status.HTTP_202_ACCEPTED,
        )

    def _run_scan_job(self, scan, params):
        """Execute the full scan pipeline for a queued scan.

        Runs in a background worker (see submit_scan_job / _scan_job_worker),
        never in the web request, so a slow scan never blocks the API. All state
        is persisted on the Scan row; the return value is ignored by the worker.
        """
        # Expand requested template categories into concrete template paths.
        templates = params.get('templates', ['ssl', 'dns', 'email'])
        expanded_templates = []
        for template in templates:
            mapped = get_template_path(template)
            if not mapped:
                continue
            if "," in mapped:
                expanded_templates.extend([item for item in mapped.split(",") if item])
            else:
                expanded_templates.append(mapped)
        templates = expanded_templates
        check_libraries = params.get('check_libraries', True)
        check_cves = params.get('check_cves', True)
        extract_org = params.get('extract_org', True)

        # Get assets
        scan_assets = ScanAsset.objects.filter(scan=scan).select_related('asset')
        assets = [sa.asset for sa in scan_assets]

        active_assets, inactive_assets = _split_active_assets(assets)
        if inactive_assets:
            logger.warning(
                f"Scan {scan.id} precheck: skipping {len(inactive_assets)} inactive targets"
            )

        if not active_assets:
            scan.status = 'failed'
            scan.error_message = 'All targets are inactive/unreachable. Nothing to scan.'
            scan.finished_at = timezone.now()
            if scan.started_at:
                scan.duration_seconds = int((scan.finished_at - scan.started_at).total_seconds())
            else:
                scan.duration_seconds = 0
            scan.save()
            return Response(
                {
                    'error': 'All targets are inactive/unreachable. Nothing to scan.',
                    'scan_id': scan.id,
                    'status': 'failed',
                    'inactive_targets': inactive_assets,
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        if inactive_assets:
            ScanAsset.objects.filter(scan=scan, asset_id__in=[item['asset_id'] for item in inactive_assets]).delete()

        assets = active_assets
        
        # Update status
        scan.cancel_requested = False
        scan.status = 'running'
        scan.started_at = timezone.now()
        scan.save()

        def _cancel_if_requested(reason='Cancelled by user during execution.'):
            scan.refresh_from_db(fields=['cancel_requested'])
            if not scan.cancel_requested:
                return None

            now = timezone.now()
            scan.status = 'cancelled'
            scan.finished_at = now
            if scan.started_at:
                scan.duration_seconds = int((now - scan.started_at).total_seconds())
            else:
                scan.duration_seconds = 0
            scan.error_message = reason
            scan.save(update_fields=['status', 'finished_at', 'duration_seconds', 'error_message'])
            return Response(
                {
                    'message': 'Scan cancelled successfully.',
                    'scan_id': scan.id,
                    'status': scan.status,
                    'duration_seconds': scan.duration_seconds,
                }
            )
        
        try:
            # Execute scans for each asset
            for asset in assets:
                cancelled = _cancel_if_requested()
                if cancelled:
                    return cancelled

                logger.info(f"=== STARTING SCAN FOR {asset.value} ===")
                
                # 1. Extract organization name if requested
                if extract_org:
                    logger.info(f"🔍 Extracting organization name for {asset.value}")
                    try:
                        # Ensure we have a domain object
                        if not hasattr(asset, 'domain') or asset.domain is None:
                            logger.warning(f"⚠️ Asset {asset.id} has no associated domain - skipping org extraction")
                        else:
                            # Only extract if not already set
                            if not asset.domain.owner:
                                # Ensure URL has protocol for org extraction
                                asset_url = asset.value
                                if not asset_url.startswith(('http://', 'https://')):
                                    asset_url = f'https://{asset_url}'
                                    logger.info(f"Added protocol: {asset_url}")
                                
                                logger.info(f"Calling extract_organization_name for {asset_url}")
                                org_name = extract_organization_name(asset_url)
                                
                                if org_name:
                                    asset.domain.owner = org_name
                                    asset.domain.save()
                                    logger.info(f"✅ Organization saved: '{org_name}' for domain {asset.domain.root_domain}")
                                else:
                                    logger.warning(f"⚠️ Could not extract organization name for {asset_url}")
                            else:
                                logger.info(f"ℹ️ Organization already set: {asset.domain.owner}")
                    except Exception as e:
                        logger.error(f"❌ Org extraction failed for {asset.value}: {e}")
                        import traceback
                        logger.error(f"Traceback:\n{traceback.format_exc()}")

                cancelled = _cancel_if_requested()
                if cancelled:
                    return cancelled
                
                # 2. Run Nuclei scan
                logger.info(f"Running Nuclei scan on {asset.value}")
                try:
                    results = run_nuclei_scan(asset.value, templates=templates)
                    if results:
                        _process_nuclei_results(scan, asset, results, templates)
                        logger.info(f"✓ Nuclei scan completed: {len(results)} results")
                    else:
                        logger.info(f"ℹ️ Nuclei scan skipped or returned no results (templates may be unavailable on Windows)")
                except Exception as e:
                    logger.warning(f"⚠️ Nuclei scan failed: {e}. Continuing with other checks...")

                cancelled = _cancel_if_requested()
                if cancelled:
                    return cancelled
                
                # 2.5. Check security headers directly via HTTP
                logger.info(f"🔍 === SECURITY HEADERS CHECK STARTED ===")
                try:
                    asset_url = asset.value
                    if not asset_url.startswith(('http://', 'https://')):
                        asset_url = f'https://{asset_url}'
                        logger.info(f"Added protocol: {asset_url}")
                    
                    headers_findings = check_security_headers(asset_url)
                    _process_header_checks(scan, asset, headers_findings)
                    
                    missing_count = sum(1 for f in headers_findings.values() if f['status'] == 'missing')
                    present_count = sum(1 for f in headers_findings.values() if f['status'] == 'present')
                    logger.info(f"✓ Security headers check completed: {present_count} present, {missing_count} missing")
                except Exception as e:
                    logger.error(f"❌ Security headers check failed: {e}")
                    import traceback
                    logger.error(f"Traceback:\n{traceback.format_exc()}")

                cancelled = _cancel_if_requested()
                if cancelled:
                    return cancelled
                
                # 3. Check for libraries and CVEs
                if check_libraries:
                    logger.info(f"🔍 === LIBRARY DETECTION STARTED ===")
                    
                    try:
                        # Ensure URL has protocol
                        asset_url = asset.value
                        if not asset_url.startswith(('http://', 'https://')):
                            asset_url = f'https://{asset_url}'
                            logger.info(f"Added protocol: {asset_url}")
                        
                        # Detect libraries and CMS
                        logger.info(f"Calling detect_technologies for {asset_url}")
                        tech_result = detect_technologies(asset_url)

                        cancelled = _cancel_if_requested()
                        if cancelled:
                            return cancelled

                        _process_technology_results(scan, asset, tech_result)
                        
                        detected_libraries = tech_result.get('libraries', [])
                        detected_libraries = _augment_library_detections_with_analytics(tech_result, detected_libraries)
                        detected_cms = tech_result.get('cms')
                        
                        logger.info(f"✓ Detected {len(detected_libraries)} libraries")
                        
                        if not detected_libraries:
                            logger.warning("⚠️ No libraries detected - this might be normal for static sites")
                        else:
                            # Build library list string
                            lib_list = ', '.join([f"{lib.get('name')}@{lib.get('version')}" for lib in detected_libraries])
                            logger.info(f"Libraries found: {lib_list}")
                        
                        # Process each library - NEVER SKIP!
                        for lib in detected_libraries:
                            cancelled = _cancel_if_requested()
                            if cancelled:
                                return cancelled

                            lib_name = lib.get('name')
                            lib_version = lib.get('version', 'Unknown')  # Default to 'Unknown'
                            lib_source = lib.get('source', 'unknown')
                            latest_version = 'Unknown'
                            
                            logger.info(f"\n--- Processing: {lib_name}@{lib_version} (source: {lib_source}) ---")
                            
                            # Initialize defaults
                            vulnerabilities = []
                            vuln_status = 'unknown'
                            risk_level = 'Low'
                            max_cvss = 0.0
                            recommendation = f"Version detection incomplete for {lib_name}"
                            
                            # Only check CVEs if we have a valid version
                            can_check_cve = (
                                lib_version and 
                                lib_version.strip() != '' and
                                lib_version.lower() not in ['unknown', 'latest', 'saas', 'n/a']
                            )

                            try:
                                ecosystem = get_ecosystem_for_library(lib_name)
                                resolved_latest_version = fetch_latest_library_version(lib_name, ecosystem)
                                if resolved_latest_version:
                                    latest_version = resolved_latest_version
                                    logger.info(f"📦 Latest registry version for {lib_name}: {latest_version}")
                            except Exception as e:
                                logger.warning(f"⚠️ Latest version lookup failed for {lib_name}: {e}")
                            
                            if can_check_cve and check_cves:
                                logger.info(f"🔍 Checking CVEs for {lib_name}@{lib_version}")
                                
                                try:
                                    # Normalize library name for CVE lookup
                                    lib_name_normalized = lib_name.lower().strip()
                                    lib_name_normalized = re.sub(r'\.(js|css)$', '', lib_name_normalized)
                                    lib_name_normalized = re.sub(r'\s+', '-', lib_name_normalized)
                                    
                                    logger.info(f"Normalized: '{lib_name}' → '{lib_name_normalized}'")
                                    
                                    # Get ecosystem
                                    ecosystem = get_ecosystem_for_library(lib_name_normalized)
                                    logger.info(f"Ecosystem: {ecosystem}")
                                    
                                    # Check vulnerabilities
                                    cve_result = check_library_vulnerabilities(
                                        lib_name_normalized,
                                        lib_version,
                                        ecosystem,
                                        use_all_sources=True
                                    )
                                    
                                    vulnerabilities = cve_result.get('vulnerabilities', [])
                                    max_cvss = cve_result.get('max_cvss_score', 0.0)
                                    
                                    logger.info(f"CVE check complete: {len(vulnerabilities)} vulns, max CVSS: {max_cvss}")
                                    
                                    if vulnerabilities:
                                        vuln_status = 'vulnerable'
                                        
                                        # Determine risk level based on CVSS
                                        if max_cvss >= 9.0:
                                            risk_level = 'Critical'
                                        elif max_cvss >= 7.0:
                                            risk_level = 'High'
                                        elif max_cvss >= 4.0:
                                            risk_level = 'Medium'
                                        else:
                                            risk_level = 'Low'
                                        
                                        recommendation = f"Update {lib_name} immediately - {len(vulnerabilities)} CVEs found"
                                        logger.info(f"⚠️ {len(vulnerabilities)} CVEs found - Risk: {risk_level}")
                                        
                                        # Log first 3 CVEs
                                        for i, vuln in enumerate(vulnerabilities[:3], 1):
                                            cve_id = vuln.get('cve_id') or vuln.get('primary_cve') or vuln.get('id')
                                            logger.info(f"  {i}. {cve_id}: CVSS {vuln.get('cvss_score', 0.0)}")
                                    else:
                                        if latest_version != 'Unknown' and version_compare(lib_version, latest_version) < 0:
                                            vuln_status = 'outdated'
                                            recommendation = f"Update {lib_name} from {lib_version} to {latest_version}"
                                            logger.info(f"⚠️ No CVEs found, but {lib_name} is outdated ({lib_version} < {latest_version})")
                                        else:
                                            vuln_status = 'up-to-date'
                                            recommendation = f"No known vulnerabilities for {lib_name} {lib_version}"
                                            logger.info(f"✓ No CVEs found")
                                        
                                except Exception as e:
                                    logger.error(f"❌ CVE check failed for {lib_name}: {e}")
                                    import traceback
                                    logger.error(f"Traceback:\n{traceback.format_exc()}")
                                    vuln_status = 'check-failed'
                                    recommendation = f"CVE check failed for {lib_name} - manual review recommended"
                            
                            else:
                                # Can't check CVEs, but still save the library
                                logger.warning(f"⚠️ Cannot check CVEs for {lib_name} (version: {lib_version})")
                                vuln_status = 'unknown'
                                recommendation = f"Manual version verification needed for {lib_name}"
                            
                            # ALWAYS create the library check record - NEVER SKIP
                            library_check, created = FrontendLibraryCheck.objects.get_or_create(
                                scan=scan,
                                asset=asset,
                                library_name=lib_name,
                                detected_version=lib_version,
                                defaults={
                                    'latest_version': latest_version,
                                    'vulnerability_status': vuln_status,
                                    'risk_level': risk_level,
                                    'source_urls': [lib.get('source_url', asset_url)],
                                    'recommendation': recommendation
                                }
                            )

                            if not created:
                                updates = []
                                if library_check.latest_version != latest_version:
                                    library_check.latest_version = latest_version
                                    updates.append('latest_version')
                                if library_check.vulnerability_status != vuln_status:
                                    library_check.vulnerability_status = vuln_status
                                    updates.append('vulnerability_status')
                                if library_check.risk_level != risk_level:
                                    library_check.risk_level = risk_level
                                    updates.append('risk_level')
                                if library_check.recommendation != recommendation:
                                    library_check.recommendation = recommendation
                                    updates.append('recommendation')
                                if updates:
                                    library_check.save(update_fields=updates)
                            
                            if created:
                                logger.info(f"✓✓✓ SAVED to DB: FrontendLibraryCheck id={library_check.id}")
                                logger.info(f"✓ Created library check: {lib_name} v{lib_version} - {vuln_status} ({risk_level})")
                            else:
                                logger.info(f"ℹ️ Already exists: {lib_name} v{lib_version}")

                            # ============================================================
                            # VULNERABILITY PROCESSING
                            # ============================================================
                            logger.info(f"\n{'='*80}")
                            logger.info(f"VULNERABILITY PROCESSING FOR {lib_name} v{lib_version}")
                            logger.info(f"{'='*80}")
                            logger.info(f"Total vulnerabilities found: {len(vulnerabilities)}")
                            logger.info(f"Max CVSS score: {max_cvss}")
                            logger.info(f"Risk level: {risk_level}")
                            logger.info(f"Asset: {asset.value} (ID={asset.id})")
                            logger.info(f"Scan: {scan.id}")

                            if not vulnerabilities:
                                logger.warning(f"⚠️ No vulnerabilities to process for {lib_name} v{lib_version}")
                            else:
                                logger.info(f"Processing {len(vulnerabilities)} vulnerabilities...")
                                for i, v in enumerate(vulnerabilities, 1):
                                    v_id = (
                                        v.get('cve_id') or 
                                        v.get('primary_cve') or 
                                        (v.get('cve_ids', [None])[0] if v.get('cve_ids') else None) or
                                        v.get('id', 'UNKNOWN')
                                    )
                                    v_cvss = v.get('cvss_score', 0.0)
                                    v_source = v.get('source', 'UNKNOWN')
                                    logger.info(f"  {i}. {v_id} (CVSS: {v_cvss}) [Source: {v_source}]")
                            
                            # Create findings for CVEs (only if vulnerabilities found)
                            for idx, vuln in enumerate(vulnerabilities, 1):
                                cancelled = _cancel_if_requested()
                                if cancelled:
                                    return cancelled

                                logger.info(f"\n{'='*60}")
                                logger.info(f"Processing vulnerability {idx}/{len(vulnerabilities)} for {lib_name} v{lib_version}")
                                
                                # Extract vulnerability ID
                                vuln_id = (
                                    vuln.get('cve_id') or 
                                    vuln.get('primary_cve') or 
                                    (vuln.get('cve_ids', [None])[0] if vuln.get('cve_ids') else None) or
                                    vuln.get('id', 'UNKNOWN')
                                )

                                # Determine if this is a CVE-based vulnerability
                                is_cve_vulnerability = vuln_id and vuln_id.startswith('CVE-')

                                logger.info(f"Vulnerability ID: {vuln_id} (CVE: {is_cve_vulnerability})")
                                logger.info(f"Source: {vuln.get('source', 'UNKNOWN')}")
                                
                                # Extract CVSS score with fallback to severity mapping
                                cvss_score = vuln.get('cvss_score', 0.0)
                                if cvss_score is None or cvss_score == 0.0:
                                    cvss_score = max_cvss

                                try:
                                    cvss_score = float(cvss_score)
                                except (ValueError, TypeError):
                                    cvss_score = 0.0

                                # If still 0, use severity mapping
                                if cvss_score == 0.0:
                                    severity = vuln.get('severity', 'UNKNOWN').upper()
                                    severity_map = {
                                        'CRITICAL': 9.5,
                                        'HIGH': 7.5,
                                        'MEDIUM': 5.0,
                                        'MODERATE': 5.0,
                                        'LOW': 3.0,
                                        'UNKNOWN': 0.0
                                    }
                                    cvss_score = severity_map.get(severity, 5.0)
                                    logger.info(f"Using severity-based CVSS: {cvss_score} (from {severity})")

                                cvss_vector = vuln.get('cvss_vector', '')

                                logger.info(f"📊 CVSS Score: {cvss_score}")
                                
                                # Step 1: Create CVE record ONLY if this is a CVE vulnerability
                                cve_record = None
                                if is_cve_vulnerability:
                                    logger.info(f"📝 Creating/updating CVE record for {vuln_id}")
                                    cve_record, cve_created = CVE.objects.get_or_create(
                                        cve_id=vuln_id,
                                        defaults={
                                            'cvss_score': cvss_score,
                                            'cvss_vector': cvss_vector or '',
                                            'description': (
                                                vuln.get('summary') or 
                                                vuln.get('description') or 
                                                vuln.get('details', '')
                                            )[:500],
                                            'published_date': vuln.get('published'),
                                            'last_modified': vuln.get('modified')
                                        }
                                    )
                                    
                                    if cve_created:
                                        logger.info(f"✅ Created new CVE record: {vuln_id}")
                                    else:
                                        # Update existing CVE if we have better data
                                        updated = False
                                        if cvss_score > 0.0 and (not cve_record.cvss_score or cve_record.cvss_score == 0.0):
                                            cve_record.cvss_score = cvss_score
                                            updated = True
                                        if cvss_vector and not cve_record.cvss_vector:
                                            cve_record.cvss_vector = cvss_vector
                                            updated = True
                                        if updated:
                                            cve_record.save()
                                            logger.info(f"Updated CVE record: {vuln_id}")
                                        else:
                                            logger.info(f"CVE record already exists: {vuln_id}")
                                else:
                                    logger.info(f"Non-CVE vulnerability: {vuln_id} - will create Finding without CVE link")
                                
                                # Step 2: Create Finding - IMPROVED title for non-CVE vulnerabilities
                                if is_cve_vulnerability:
                                    finding_title = f"{lib_name} {lib_version} - {vuln_id}"
                                else:
                                    # For non-CVE vulnerabilities, use source and summary
                                    source = vuln.get('source', 'Security')
                                    summary = vuln.get('summary', 'Vulnerability')
                                    
                                    # Clean up summary for title (take first line or first 60 chars)
                                    if summary:
                                        summary_clean = summary.split('\n')[0].split('http')[0].strip()
                                        if len(summary_clean) > 60:
                                            summary_clean = summary_clean[:57] + '...'
                                    else:
                                        summary_clean = 'Vulnerability'
                                    
                                    finding_title = f"{lib_name} {lib_version} - {source} Vulnerability: {summary_clean}"

                                logger.info(f"Finding title: {finding_title}")
                                
                                # FIXED: More specific deduplication - match exact title + asset
                                existing_finding = Finding.objects.filter(
                                    asset=asset,
                                    title=finding_title,
                                    category='CVE'
                                ).first()
                                
                                if existing_finding:
                                    logger.info(f"📎 Finding already exists (ID={existing_finding.id}): {finding_title}")
                                    
                                    # Link existing finding to current scan
                                    scan_finding, sf_created = ScanFinding.objects.get_or_create(
                                        scan=scan,
                                        finding=existing_finding
                                    )
                                    if sf_created:
                                        logger.info(f"✅ Linked existing finding to scan {scan.id}")
                                    else:
                                        logger.info(f"ℹ️ Finding already linked to scan {scan.id}")
                                else:
                                    # CREATE NEW FINDING
                                    logger.info(f"💾 Creating new finding: {finding_title}")
                                    logger.info(f"   Asset: {asset.value}")
                                    logger.info(f"   CVSS Score: {cvss_score}")
                                    logger.info(f"   Risk Level: {risk_level}")
                                    
                                    try:
                                        # Get recommendation
                                        recommendation_text = (
                                            vuln.get('recommendation') or
                                            vuln.get('remediation') or
                                            vuln.get('details') or 
                                            vuln.get('summary') or 
                                            f'Update {lib_name} to a patched version'
                                        )
                                        
                                        # Get evidence/description
                                        evidence_text = (
                                            vuln.get('summary') or
                                            vuln.get('description') or
                                            vuln.get('details') or
                                            f"Vulnerable library: {lib_name} v{lib_version}"
                                        )

                                        # Add vulnerability source and references to evidence for non-CVE vulns
                                        if not is_cve_vulnerability:
                                            refs = vuln.get('references', [])
                                            if refs:
                                                evidence_text += f"\n\nReferences:\n" + "\n".join(refs[:3])

                                        finding = Finding.objects.create(
                                            asset=asset,
                                            title=finding_title,
                                            category='CVE',
                                            nuclei_template_id='library-vuln-check',
                                            nuclei_severity=vuln.get('severity', 'medium').lower(),
                                            cvss_score=cvss_score,
                                            cvss_vector=cvss_vector or '',
                                            risk_rating=risk_level,
                                            scoring_confidence='High',
                                            evidence=evidence_text[:1000],
                                            recommendation=recommendation_text[:500],
                                            status='open'
                                        )
                                        
                                        logger.info(f"✅✅✅ CREATED Finding ID={finding.id}: {finding_title}")
                                        logger.info(f"   Database ID: {finding.id}")
                                        logger.info(f"   CVSS Score in DB: {finding.cvss_score}")
                                        
                                        # Link finding to scan
                                        scan_finding = ScanFinding.objects.create(
                                            scan=scan, 
                                            finding=finding
                                        )
                                        logger.info(f"✅ Linked finding to scan via ScanFinding ID={scan_finding.id}")
                                        
                                        # Link finding to CVE (ONLY if we have a CVE record)
                                        if cve_record:
                                            FindingCVE.objects.create(
                                                finding=finding,
                                                cve=cve_record,
                                                relevance='direct'
                                            )
                                            logger.info(f"✅ Linked finding to CVE {vuln_id}")
                                        else:
                                            logger.info(f"ℹ️ Non-CVE vulnerability - no CVE link created")
                                        
                                        # Verify the finding was created
                                        verify_finding = Finding.objects.filter(id=finding.id).first()
                                        if verify_finding:
                                            logger.info(f"✅ VERIFICATION: Finding exists in database")
                                            logger.info(f"   Title: {verify_finding.title}")
                                            logger.info(f"   CVSS: {verify_finding.cvss_score}")
                                            logger.info(f"   Category: {verify_finding.category}")
                                        else:
                                            logger.error(f"❌ VERIFICATION FAILED: Finding not found in database!")
                                            
                                    except Exception as e:
                                        logger.error(f"❌ FAILED to create finding: {e}")
                                        import traceback
                                        logger.error(f"Traceback:\n{traceback.format_exc()}")
                            
                            logger.info(f"\n{'='*80}")
                            logger.info(f"FINISHED processing {len(vulnerabilities)} vulnerabilities for {lib_name}")
                            logger.info(f"{'='*80}\n")
                        
                        # Store CMS info if detected
                        if detected_cms:
                            logger.info(f"🔍 Detected CMS: {detected_cms['name']} v{detected_cms['version']}")
                            
                            # Create informational finding for CMS
                            if detected_cms['version'] not in ['Unknown', 'SaaS']:
                                existing_cms_finding = Finding.objects.filter(
                                    asset=asset,
                                    title__icontains=f"{detected_cms['name']} CMS"
                                ).first()
                                
                                if not existing_cms_finding:
                                    finding = Finding.objects.create(
                                        asset=asset,
                                        title=f"{detected_cms['name']} CMS Detected - {detected_cms['version']}",
                                        category='Misconfiguration',
                                        nuclei_template_id='cms-detection',
                                        nuclei_severity='info',
                                        risk_rating='Low',
                                        scoring_confidence='High',
                                        evidence=f"CMS: {detected_cms['name']} v{detected_cms['version']}",
                                        recommendation=f"Ensure {detected_cms['name']} is up to date and properly configured",
                                        status='open'
                                    )
                                    ScanFinding.objects.create(scan=scan, finding=finding)
                                    logger.info(f"✓ Created CMS finding")
                                else:
                                    # Link existing CMS finding to current scan
                                    ScanFinding.objects.get_or_create(
                                        scan=scan,
                                        finding=existing_cms_finding
                                    )
                                    logger.info(f"📎 Linked existing CMS finding")
                    
                    except Exception as e:
                        logger.error(f"❌ Library/CVE detection failed for {asset.value}: {e}")
                        import traceback
                        logger.error(f"Full traceback:\n{traceback.format_exc()}")

            scan.refresh_from_db(fields=['cancel_requested'])
            scan.finished_at = timezone.now()
            scan.duration_seconds = int((scan.finished_at - scan.started_at).total_seconds())
            if scan.cancel_requested:
                scan.status = 'cancelled'
                scan.error_message = 'Cancelled by user during execution.'
            else:
                scan.status = 'completed'
            scan.save()
            
            return Response({
                'message': 'Scan executed successfully' if scan.status == 'completed' else 'Scan cancelled successfully.',
                'scan_id': scan.id,
                'status': scan.status,
                'duration_seconds': scan.duration_seconds,
                'scanned_assets_count': len(assets),
                'skipped_assets_count': len(inactive_assets),
                'skipped_targets': inactive_assets,
            })
        
        except Exception as e:
            scan.status = 'failed'
            scan.error_message = str(e)
            scan.finished_at = timezone.now()
            if scan.started_at:
                scan.duration_seconds = int((scan.finished_at - scan.started_at).total_seconds())
            else:
                scan.duration_seconds = 0
            scan.save()
            return Response({'error': str(e)}, status=500)

# ============================================
# SEPARATE ENDPOINTS FOR EACH CHECK TYPE
# ============================================

@api_view(['GET'])
def scan_library_checks(request, scan_id):
    """
    Get frontend library checks for a scan
    GET /api/scans/{scan_id}/library-checks/
    """
    scan = _scoped_scan(request, scan_id)
    checks = FrontendLibraryCheck.objects.filter(scan=scan).select_related('asset')
    
    # Optional filtering
    risk_level = request.GET.get('risk_level')
    if risk_level:
        checks = checks.filter(risk_level=risk_level)
    
    status_filter = request.GET.get('status')
    if status_filter:
        checks = checks.filter(vulnerability_status=status_filter)
    
    serializer = FrontendLibraryCheckSerializer(checks, many=True)
    return Response({
        'scan_id': scan_id,
        'count': checks.count(),
        'checks': serializer.data
    })


@api_view(['GET'])
def scan_ssl_checks(request, scan_id):
    """
    Get SSL/TLS checks for a scan
    GET /api/scans/{scan_id}/ssl-checks/
    """
    scan = _scoped_scan(request, scan_id)
    checks = SSLTLSCheck.objects.filter(scan=scan).select_related('asset')
    
    # Optional filtering
    check_type = request.GET.get('check_type')
    if check_type:
        checks = checks.filter(check_type=check_type)
    
    risk_level = request.GET.get('risk_level')
    if risk_level:
        checks = checks.filter(risk_rating=risk_level)
    
    serializer = SSLTLSCheckSerializer(checks, many=True)
    return Response({
        'scan_id': scan_id,
        'count': checks.count(),
        'checks': serializer.data
    })


@api_view(['GET'])
def scan_email_checks(request, scan_id):
    """
    Get email security checks for a scan
    GET /api/scans/{scan_id}/email-checks/
    """
    scan = _scoped_scan(request, scan_id)
    checks = EmailSecurityCheck.objects.filter(scan=scan).select_related('asset')
    
    # Optional filtering
    check_type = request.GET.get('check_type')
    if check_type:
        checks = checks.filter(check_type=check_type)
    
    status_filter = request.GET.get('status')
    if status_filter:
        checks = checks.filter(status=status_filter)
    
    serializer = EmailSecurityCheckSerializer(checks, many=True)
    return Response({
        'scan_id': scan_id,
        'count': checks.count(),
        'checks': serializer.data
    })


@api_view(['GET'])
def scan_header_checks(request, scan_id):
    """
    Get security header checks for a scan
    GET /api/scans/{scan_id}/header-checks/
    """
    scan = _scoped_scan(request, scan_id)
    checks = SecurityHeaderCheck.objects.filter(scan=scan).select_related('asset')
    
    # Optional filtering
    status_filter = request.GET.get('status')
    if status_filter:
        checks = checks.filter(status=status_filter)
    
    risk_level = request.GET.get('risk_level')
    if risk_level:
        checks = checks.filter(risk_rating=risk_level)
    
    serializer = SecurityHeaderCheckSerializer(checks, many=True)
    return Response({
        'scan_id': scan_id,
        'count': checks.count(),
        'checks': serializer.data
    })


def _score_to_grade(score: int) -> str:
    """Convert numeric score to simple A-F grade."""
    if score >= 90:
        return 'A'
    if score >= 80:
        return 'B'
    if score >= 70:
        return 'C'
    if score >= 60:
        return 'D'
    return 'F'


_SEVERITY_WEIGHT = {
    'Critical': 10,
    'High': 7,
    'Medium': 4,
    'Low': 2,
}


def _normalize_risk(risk_value, fallback='Medium'):
    risk = (risk_value or fallback).strip().title() if isinstance(risk_value, str) else fallback
    return risk if risk in _SEVERITY_WEIGHT else fallback


def _empty_severity_counts():
    return {
        'critical': 0,
        'high': 0,
        'medium': 0,
        'low': 0,
    }


def _add_severity_count(container, risk):
    key = _normalize_risk(risk).lower()
    if key in container:
        container[key] += 1


def _compute_grade_score(issue_points, max_points):
    score = int(round(100 * (1 - (issue_points / max_points)))) if max_points else 0
    score = max(0, min(score, 100))
    impact = round((issue_points / max_points) * 100, 1) if max_points else 0.0
    return score, impact, _score_to_grade(score)


def _build_header_scorecard_data(scan_id, checks):
    if not checks:
        return {
            'scan_id': scan_id,
            'factor': 'Security Headers',
            'grade': 'N/A',
            'score': 0,
            'impact': 0.0,
            'summary': {
                'total_headers': 0,
                'passed': 0,
                'failed': 0,
                'missing': 0,
                'misconfigured': 0,
                'present': 0,
            },
            'issues_by_severity': _empty_severity_counts(),
            'findings': [],
        }

    findings = []
    issue_points = 0
    max_points = 0
    missing = 0
    misconfigured = 0
    present = 0
    severity_counts = _empty_severity_counts()

    for item in checks:
        risk = _normalize_risk(item.risk_rating)
        weight = _SEVERITY_WEIGHT.get(risk, 4)
        max_points += weight

        is_issue = item.status in {'missing', 'misconfigured'}
        if is_issue:
            issue_points += weight
            if item.status == 'missing':
                missing += 1
            if item.status == 'misconfigured':
                misconfigured += 1
            _add_severity_count(severity_counts, risk)
        elif item.status == 'present':
            present += 1

        findings.append(
            {
                'header': item.header,
                'status': item.status,
                'risk_rating': risk,
                'cvss_score': item.cvss_score,
                'asset': item.asset.value if item.asset else None,
                'recommendation': item.recommendation,
                'header_value': item.header_value,
                'checked_at': item.checked_at,
                'is_issue': is_issue,
            }
        )

    score, impact, grade = _compute_grade_score(issue_points, max_points)
    return {
        'scan_id': scan_id,
        'factor': 'Security Headers',
        'grade': grade,
        'score': score,
        'impact': impact,
        'summary': {
            'total_headers': len(checks),
            'passed': present,
            'failed': missing + misconfigured,
            'missing': missing,
            'misconfigured': misconfigured,
            'present': present,
        },
        'issues_by_severity': severity_counts,
        'findings': findings,
    }


def _build_ssl_scorecard_data(scan_id, checks):
    if not checks:
        return {
            'scan_id': scan_id,
            'factor': 'SSL/TLS',
            'grade': 'N/A',
            'score': 0,
            'impact': 0.0,
            'summary': {'total_checks': 0, 'passed': 0, 'failed': 0},
            'issues_by_severity': _empty_severity_counts(),
            'findings': [],
        }

    severity_counts = _empty_severity_counts()
    findings = []
    issue_points = 0
    max_points = 0
    passed = 0
    failed = 0

    for item in checks:
        risk = _normalize_risk(item.risk_rating)
        weight = _SEVERITY_WEIGHT.get(risk, 4)
        max_points += weight
        cvss = float(item.cvss_score) if item.cvss_score is not None else 0.0
        is_issue = not (risk == 'Low' and cvss <= 0.0)
        if is_issue:
            failed += 1
            issue_points += weight
            _add_severity_count(severity_counts, risk)
        else:
            passed += 1
        findings.append(
            {
                'check_type': item.check_type,
                'finding': item.finding,
                'status': 'issue' if is_issue else 'pass',
                'risk_rating': risk,
                'cvss_score': item.cvss_score,
                'asset': item.asset.value if item.asset else None,
                'recommendation': item.recommendation,
                'checked_at': item.checked_at,
                'is_issue': is_issue,
            }
        )

    score, impact, grade = _compute_grade_score(issue_points, max_points)
    return {
        'scan_id': scan_id,
        'factor': 'SSL/TLS',
        'grade': grade,
        'score': score,
        'impact': impact,
        'summary': {'total_checks': len(checks), 'passed': passed, 'failed': failed},
        'issues_by_severity': severity_counts,
        'findings': findings,
    }


def _build_email_scorecard_data(scan_id, checks):
    if not checks:
        return {
            'scan_id': scan_id,
            'factor': 'Email Security',
            'grade': 'N/A',
            'score': 0,
            'impact': 0.0,
            'summary': {'total_checks': 0, 'passed': 0, 'failed': 0},
            'issues_by_severity': _empty_severity_counts(),
            'findings': [],
        }

    severity_counts = _empty_severity_counts()
    findings = []
    issue_points = 0
    max_points = 0
    passed = 0
    failed = 0

    for item in checks:
        risk = _normalize_risk(item.risk_rating)
        weight = _SEVERITY_WEIGHT.get(risk, 4)
        max_points += weight
        is_issue = (item.status or '').upper() != 'PASS'
        if is_issue:
            failed += 1
            issue_points += weight
            _add_severity_count(severity_counts, risk)
        else:
            passed += 1
        findings.append(
            {
                'check_type': item.check_type,
                'status': item.status,
                'risk_rating': risk,
                'cvss_score': item.cvss_score,
                'asset': item.asset.value if item.asset else None,
                'details': item.details,
                'record_value': item.record_value,
                'recommendation': item.recommendation,
                'checked_at': item.checked_at,
                'is_issue': is_issue,
            }
        )

    score, impact, grade = _compute_grade_score(issue_points, max_points)
    return {
        'scan_id': scan_id,
        'factor': 'Email Security',
        'grade': grade,
        'score': score,
        'impact': impact,
        'summary': {'total_checks': len(checks), 'passed': passed, 'failed': failed},
        'issues_by_severity': severity_counts,
        'findings': findings,
    }


def _build_dns_scorecard_data(scan_id, checks):
    if not checks:
        return {
            'scan_id': scan_id,
            'factor': 'DNS Health',
            'grade': 'N/A',
            'score': 0,
            'impact': 0.0,
            'summary': {'total_checks': 0, 'passed': 0, 'failed': 0},
            'issues_by_severity': _empty_severity_counts(),
            'findings': [],
        }

    severity_counts = _empty_severity_counts()
    findings = []
    issue_points = 0
    max_points = 0
    passed = 0
    failed = 0

    for item in checks:
        risk = _normalize_risk(item.risk_rating)
        weight = _SEVERITY_WEIGHT.get(risk, 4)
        max_points += weight
        cvss = float(item.cvss_score) if item.cvss_score is not None else 0.0
        is_issue = not (risk == 'Low' and cvss <= 0.0)
        if is_issue:
            failed += 1
            issue_points += weight
            _add_severity_count(severity_counts, risk)
        else:
            passed += 1
        findings.append(
            {
                'check_type': item.check_type,
                'finding': item.finding,
                'status': 'issue' if is_issue else 'pass',
                'risk_rating': risk,
                'cvss_score': item.cvss_score,
                'asset': item.asset.value if item.asset else None,
                'example': item.example,
                'recommendation': item.recommendation,
                'checked_at': item.checked_at,
                'is_issue': is_issue,
            }
        )

    score, impact, grade = _compute_grade_score(issue_points, max_points)
    return {
        'scan_id': scan_id,
        'factor': 'DNS Health',
        'grade': grade,
        'score': score,
        'impact': impact,
        'summary': {'total_checks': len(checks), 'passed': passed, 'failed': failed},
        'issues_by_severity': severity_counts,
        'findings': findings,
    }


def _build_library_scorecard_data(scan_id, checks):
    if not checks:
        return {
            'scan_id': scan_id,
            'factor': 'Library Security',
            'grade': 'A',
            'score': 100,
            'impact': 0.0,
            'summary': {'total_checks': 0, 'passed': 0, 'failed': 0},
            'issues_by_severity': _empty_severity_counts(),
            'findings': [],
        }

    severity_counts = _empty_severity_counts()
    findings = []
    issue_points = 0
    max_points = 0
    passed = 0
    failed = 0

    risk_to_cvss = {
        'Critical': 9.5,
        'High': 7.5,
        'Medium': 5.0,
        'Low': 2.5,
    }

    for item in checks:
        risk = _normalize_risk(item.risk_level)
        weight = _SEVERITY_WEIGHT.get(risk, 4)
        max_points += weight

        status_value = (item.vulnerability_status or '').lower()
        is_issue = status_value != 'up-to-date'

        if is_issue:
            failed += 1
            issue_points += weight
            _add_severity_count(severity_counts, risk)
        else:
            passed += 1

        findings.append(
            {
                'library_name': item.library_name,
                'detected_version': item.detected_version,
                'latest_version': item.latest_version,
                'status': item.vulnerability_status,
                'risk_rating': risk,
                'cvss_score': risk_to_cvss.get(risk, 0.0),
                'asset': item.asset.value if item.asset else None,
                'source_urls': item.source_urls,
                'recommendation': item.recommendation,
                'checked_at': item.checked_at,
                'is_issue': is_issue,
            }
        )

    score, impact, grade = _compute_grade_score(issue_points, max_points)
    return {
        'scan_id': scan_id,
        'factor': 'Library Security',
        'grade': grade,
        'score': score,
        'impact': impact,
        'summary': {'total_checks': len(checks), 'passed': passed, 'failed': failed},
        'issues_by_severity': severity_counts,
        'findings': findings,
    }


@api_view(['GET'])
def scan_header_scorecard(request, scan_id):
    """
    Security headers grading endpoint for frontend scorecard UIs.
    GET /api/scans/{scan_id}/header-scorecard/
    """
    scan = _scoped_scan(request, scan_id)
    checks = list(
        SecurityHeaderCheck.objects.filter(scan=scan)
        .select_related('asset')
        .order_by('header')
    )
    return Response(_build_header_scorecard_data(scan_id, checks))


@api_view(['GET'])
def scan_ssl_scorecard(request, scan_id):
    """
    SSL/TLS grading endpoint for frontend scorecard UIs.
    GET /api/scans/{scan_id}/ssl-scorecard/
    """
    scan = _scoped_scan(request, scan_id)
    checks = list(SSLTLSCheck.objects.filter(scan=scan).select_related('asset').order_by('check_type', 'id'))

    return Response(_build_ssl_scorecard_data(scan_id, checks))


@api_view(['GET'])
def scan_email_scorecard(request, scan_id):
    """
    Email security grading endpoint for frontend scorecard UIs.
    GET /api/scans/{scan_id}/email-scorecard/
    """
    scan = _scoped_scan(request, scan_id)
    checks = list(EmailSecurityCheck.objects.filter(scan=scan).select_related('asset').order_by('check_type', 'id'))

    return Response(_build_email_scorecard_data(scan_id, checks))


@api_view(['GET'])
def scan_dns_scorecard(request, scan_id):
    """
    DNS health grading endpoint for frontend scorecard UIs.
    GET /api/scans/{scan_id}/dns-scorecard/
    """
    scan = _scoped_scan(request, scan_id)
    checks = list(DNSSecurityCheck.objects.filter(scan=scan).select_related('asset').order_by('check_type', 'id'))

    return Response(_build_dns_scorecard_data(scan_id, checks))


@api_view(['GET'])
def scan_library_scorecard(request, scan_id):
    """
    Frontend library grading endpoint for frontend scorecard UIs.
    GET /api/scans/{scan_id}/library-scorecard/
    """
    scan = _scoped_scan(request, scan_id)
    checks = list(
        FrontendLibraryCheck.objects.filter(scan=scan)
        .select_related('asset')
        .order_by('library_name', 'id')
    )

    return Response(_build_library_scorecard_data(scan_id, checks))


@api_view(['GET'])
def scan_overall_scorecard(request, scan_id):
    """
    Combined scorecard across key factors for dashboard rendering.
    GET /api/scans/{scan_id}/overall-scorecard/
    """
    scan = _scoped_scan(request, scan_id)

    header_checks = list(SecurityHeaderCheck.objects.filter(scan=scan).select_related('asset').order_by('header'))
    ssl_checks = list(SSLTLSCheck.objects.filter(scan=scan).select_related('asset').order_by('check_type', 'id'))
    email_checks = list(EmailSecurityCheck.objects.filter(scan=scan).select_related('asset').order_by('check_type', 'id'))
    dns_checks = list(DNSSecurityCheck.objects.filter(scan=scan).select_related('asset').order_by('check_type', 'id'))
    library_checks = list(
        FrontendLibraryCheck.objects.filter(scan=scan)
        .select_related('asset')
        .order_by('library_name', 'id')
    )

    header_data = _build_header_scorecard_data(scan_id, header_checks)
    ssl_data = _build_ssl_scorecard_data(scan_id, ssl_checks)
    email_data = _build_email_scorecard_data(scan_id, email_checks)
    dns_data = _build_dns_scorecard_data(scan_id, dns_checks)
    library_data = _build_library_scorecard_data(scan_id, library_checks)

    factors = [header_data, ssl_data, email_data, dns_data, library_data]
    scored_factors = [item for item in factors if item.get('grade') != 'N/A']

    if scored_factors:
        avg_score = int(round(sum(item['score'] for item in scored_factors) / len(scored_factors)))
        avg_impact = round(sum(item['impact'] for item in scored_factors) / len(scored_factors), 1)
        overall_grade = _score_to_grade(avg_score)
    else:
        avg_score = 0
        avg_impact = 0.0
        overall_grade = 'N/A'

    return Response(
        {
            'scan_id': scan.id,
            'overall': {
                'grade': overall_grade,
                'score': avg_score,
                'impact': avg_impact,
                'factors_count': len(scored_factors),
            },
            'factors': factors,
        }
    )


@api_view(['GET'])
def scan_dns_checks(request, scan_id):
    """
    Get DNS security checks for a scan
    GET /api/scans/{scan_id}/dns-checks/
    """
    scan = _scoped_scan(request, scan_id)
    checks = DNSSecurityCheck.objects.filter(scan=scan).select_related('asset')
    
    # Optional filtering
    check_type = request.GET.get('check_type')
    if check_type:
        checks = checks.filter(check_type=check_type)
    
    risk_level = request.GET.get('risk_level')
    if risk_level:
        checks = checks.filter(risk_rating=risk_level)
    
    serializer = DNSSecurityCheckSerializer(checks, many=True)
    return Response({
        'scan_id': scan_id,
        'count': checks.count(),
        'checks': serializer.data
    })


@api_view(['GET'])
def scan_findings(request, scan_id):
    """
    Get all findings for a scan
    GET /api/scans/{scan_id}/findings/
    
    Query params:
    - category: Filter by category (CVE, SSL, DNS, etc.)
    - risk_rating: Filter by risk rating (Low, Medium, High, Critical)
    - status: Filter by status (open, in_progress, resolved, false_positive)
    - page: Page number (default: 1)
    - page_size: Results per page (default: 100, max: 1000)
    
    FIXED: Now properly returns all findings with correct filtering
    """
    scan = _scoped_scan(request, scan_id)
    
    # ✅ FIXED: Use QuerySet methods for filtering BEFORE converting to list
    scan_findings_qs = ScanFinding.objects.filter(
        scan=scan
    ).select_related('finding', 'finding__asset', 'finding__asset__domain')
    
    # Apply filters using QuerySet methods (DATABASE-LEVEL FILTERING)
    category = request.GET.get('category')
    if category:
        scan_findings_qs = scan_findings_qs.filter(finding__category=category)
        logger.info(f"Filtering by category: {category}")
    
    risk_rating = request.GET.get('risk_rating')
    if risk_rating:
        scan_findings_qs = scan_findings_qs.filter(finding__risk_rating=risk_rating)
        logger.info(f"Filtering by risk_rating: {risk_rating}")
    
    status_filter = request.GET.get('status')
    if status_filter:
        scan_findings_qs = scan_findings_qs.filter(finding__status=status_filter)
        logger.info(f"Filtering by status: {status_filter}")
    
    # Get total count BEFORE pagination
    total_count = scan_findings_qs.count()
    logger.info(f"Total findings matching filters: {total_count}")
    
    # Pagination
    page = int(request.GET.get('page', 1))
    page_size = min(int(request.GET.get('page_size', 100)), 1000)  # Max 1000
    
    start_idx = (page - 1) * page_size
    end_idx = start_idx + page_size
    
    # Apply pagination to QuerySet
    scan_findings_page = scan_findings_qs[start_idx:end_idx]
    
    # NOW convert to list (only for the current page)
    findings = [sf.finding for sf in scan_findings_page]
    
    # Serialize
    serializer = FindingSerializer(findings, many=True)
    
    logger.info(f"Returning {len(findings)} findings (page {page}/{(total_count + page_size - 1) // page_size})")
    
    return Response({
        'scan_id': scan_id,
        'total_count': total_count,
        'page': page,
        'page_size': page_size,
        'total_pages': (total_count + page_size - 1) // page_size if total_count > 0 else 0,
        'has_next': end_idx < total_count,
        'has_previous': page > 1,
        'findings': serializer.data
    })


@api_view(['GET'])
def scan_report_download(request, scan_id, report_format='json', domain_name=None):
    """
    Download generated report for a scan.
    GET /api/scans/{scan_id}/report/download/{format}/
    GET /api/report/download/{scan_id}/{format}/
    GET /api/report/download/{scan_id}/{format}/{domain_name}/

    Query params remain supported for backward compatibility:
    - format
    - domain
    """
    domain_name = (domain_name or request.GET.get('domain') or '').strip()
    report_format = (report_format or request.GET.get('format', 'json')).lower()
    logger.info(
        "scan_report_download called path=%s scan_id=%s format=%s domain=%s",
        request.path,
        scan_id,
        report_format,
        domain_name or None,
    )
    org = get_user_organisation(request)
    scan = Scan.objects.filter(id=scan_id, organisation=org).first()
    if not scan:
        return Response(
            {
                'error': 'Scan not found',
                'scan_id': scan_id,
            },
            status=status.HTTP_404_NOT_FOUND,
        )
    helper = ScanViewSet()

    if report_format not in {'json', 'csv', 'pdf', 'docx'}:
        return Response(
            {'error': "Invalid format. Use 'json', 'csv', 'pdf', or 'docx'."},
            status=status.HTTP_400_BAD_REQUEST,
        )

    try:
        payload = helper._build_scan_report_payload(scan, domain_name=domain_name or None)
    except ValueError as exc:
        return Response(
            {
                'error': str(exc),
                'scan_id': scan_id,
                'domain': domain_name or None,
            },
            status=status.HTTP_400_BAD_REQUEST,
        )
    except Exception as exc:
        logger.exception("Failed to build report payload for scan_id=%s", scan_id)
        return Response(
            {
                'error': 'Failed to build report payload',
                'detail': str(exc),
                'scan_id': scan_id,
            },
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )
    timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
    org_slug = re.sub(r'[^\w\-]', '_', payload['summary'].get('org_name') or f'scan_{scan.id}')

    if report_format == 'json':
        response = HttpResponse(
            json.dumps(payload, default=str, indent=2),
            content_type='application/json',
        )
        response['Content-Disposition'] = (
            f'attachment; filename="{org_slug}_security_report_{timestamp}.json"'
        )
        return response

    if report_format == 'pdf':
        content, error = helper._generate_pdf_report(payload)
        if error:
            return Response({'error': error}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        response = HttpResponse(content, content_type='application/pdf')
        response['Content-Disposition'] = (
            f'attachment; filename="{org_slug}_security_report_{timestamp}.pdf"'
        )
        return response

    if report_format == 'docx':
        content, error = helper._generate_docx_report(payload)
        if error:
            return Response({'error': error}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        response = HttpResponse(
            content,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        response['Content-Disposition'] = (
            f'attachment; filename="{org_slug}_security_report_{timestamp}.docx"'
        )
        return response

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = (
        f'attachment; filename="{org_slug}_security_report_{timestamp}.csv"'
    )

    writer = csv.writer(response)

    writer.writerow(['scan_id', payload['summary']['scan_id']])
    writer.writerow(['scan_type', payload['summary']['scan_type']])
    writer.writerow(['status', payload['summary']['status']])
    writer.writerow(['started_at', payload['summary']['started_at']])
    writer.writerow(['finished_at', payload['summary']['finished_at']])
    writer.writerow(['duration_seconds', payload['summary']['duration_seconds']])
    writer.writerow(['total_assets_scanned', payload['summary']['total_assets_scanned']])
    writer.writerow(['total_findings', payload['summary']['total_findings']])
    writer.writerow(['critical_findings', payload['summary']['critical_findings']])
    writer.writerow(['high_findings', payload['summary']['high_findings']])
    writer.writerow(['medium_findings', payload['summary']['medium_findings']])
    writer.writerow(['low_findings', payload['summary']['low_findings']])
    writer.writerow([])

    writer.writerow([
        'finding_id',
        'title',
        'category',
        'risk_rating',
        'status',
        'cvss_score',
        'nuclei_template_id',
        'asset',
        'domain',
        'cve_ids',
        'recommendation',
        'first_seen',
        'last_seen',
    ])

    for finding in payload['findings']:
        writer.writerow([
            finding['id'],
            finding['title'],
            finding['category'],
            finding['risk_rating'],
            finding['status'],
            finding['cvss_score'],
            finding['nuclei_template_id'],
            finding['asset'],
            finding['domain'],
            ';'.join(finding['cve_ids']),
            finding['recommendation'],
            finding['first_seen'],
            finding['last_seen'],
        ])

    return response


@api_view(['GET'])
def report_download_test(request):
    """Simple route test endpoint for download path troubleshooting."""
    logger.info("report_download_test endpoint hit")
    return Response(
        {
            'ok': True,
            'message': 'report download test route is reachable',
            'hint': 'If this works but download fails, issue is inside report handler logic.',
        },
        status=status.HTTP_200_OK,
    )


@api_view(['GET'])
def report_download_debug(request, scan_id, report_format='json'):

    """Debug endpoint to isolate where report generation fails."""
    report_format = (report_format or request.GET.get('format', 'json')).lower()
    org = get_user_organisation(request)
    scan = Scan.objects.filter(id=scan_id, organisation=org).first()

    if not scan:
        return Response(
            {
                'ok': False,
                'stage': 'scan_lookup',
                'error': 'Scan not found',
                'scan_id': scan_id,
            },
            status=status.HTTP_404_NOT_FOUND,
        )

    helper = ScanViewSet()

    try:
        payload = helper._build_scan_report_payload(scan)
    except Exception as exc:
        logger.exception("report_download_debug payload failure scan_id=%s", scan_id)
        return Response(
            {
                'ok': False,
                'stage': 'payload',
                'error': str(exc),
                'scan_id': scan_id,
            },
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )

    if report_format == 'pdf':
        content, error = helper._generate_pdf_report(payload)
        if error:
            return Response(
                {
                    'ok': False,
                    'stage': 'pdf_generation',
                    'error': error,
                    'scan_id': scan_id,
                },
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )
        return Response(
            {
                'ok': True,
                'stage': 'pdf_generation',
                'bytes': len(content),
                'scan_id': scan_id,
            },
            status=status.HTTP_200_OK,
        )

    if report_format == 'docx':
        content, error = helper._generate_docx_report(payload)
        if error:
            return Response(
                {
                    'ok': False,
                    'stage': 'docx_generation',
                    'error': error,
                    'scan_id': scan_id,
                },
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )
        return Response(
            {
                'ok': True,
                'stage': 'docx_generation',
                'bytes': len(content),
                'scan_id': scan_id,
            },
            status=status.HTTP_200_OK,
        )

    return Response(
        {
            'ok': True,
            'stage': 'payload',
            'scan_id': scan_id,
            'format': report_format,
            'summary': {
                'status': payload['summary']['status'],
                'total_assets_scanned': payload['summary']['total_assets_scanned'],
                'total_findings': payload['summary']['total_findings'],
            },
        },
        status=status.HTTP_200_OK,
    )

@api_view(['GET'])
def all_findings(request):
    """
    Get all findings across all scans
    GET /api/findings/
    
    NEW ENDPOINT: Allows querying all findings regardless of scan
    
    Query params:
    - asset_id: Filter by asset
    - domain_id: Filter by domain
    - category: Filter by category
    - risk_rating: Filter by risk rating
    - status: Filter by status
    - search: Search in title/evidence/recommendation
    - min_cvss: Minimum CVSS score
    - max_cvss: Maximum CVSS score
    - page: Page number
    - page_size: Results per page (max 1000)
    """
    org = get_user_organisation(request)
    findings_qs = Finding.objects.filter(organisation=org).select_related('asset', 'asset__domain')

    # Filters
    asset_id = request.GET.get('asset_id')
    if asset_id:
        findings_qs = findings_qs.filter(asset_id=asset_id)
    
    domain_id = request.GET.get('domain_id')
    if domain_id:
        findings_qs = findings_qs.filter(asset__domain_id=domain_id)
    
    category = request.GET.get('category')
    if category:
        findings_qs = findings_qs.filter(category=category)
    
    risk_rating = request.GET.get('risk_rating')
    if risk_rating:
        findings_qs = findings_qs.filter(risk_rating=risk_rating)
    
    status_filter = request.GET.get('status')
    if status_filter:
        findings_qs = findings_qs.filter(status=status_filter)
    
    search = request.GET.get('search')
    if search:
        findings_qs = findings_qs.filter(
            Q(title__icontains=search) | 
            Q(evidence__icontains=search) |
            Q(recommendation__icontains=search)
        )
    
    # CVSS range filters
    min_cvss = request.GET.get('min_cvss')
    if min_cvss:
        findings_qs = findings_qs.filter(cvss_score__gte=float(min_cvss))
    
    max_cvss = request.GET.get('max_cvss')
    if max_cvss:
        findings_qs = findings_qs.filter(cvss_score__lte=float(max_cvss))
    
    # Order by CVSS score (highest first), then by date
    findings_qs = findings_qs.order_by('-cvss_score', '-first_seen')
    
    # Count
    total_count = findings_qs.count()
    
    # Pagination
    page = int(request.GET.get('page', 1))
    page_size = min(int(request.GET.get('page_size', 100)), 1000)
    
    start_idx = (page - 1) * page_size
    end_idx = start_idx + page_size
    
    findings_page = findings_qs[start_idx:end_idx]
    
    serializer = FindingSerializer(findings_page, many=True)
    
    return Response({
        'total_count': total_count,
        'page': page,
        'page_size': page_size,
        'total_pages': (total_count + page_size - 1) // page_size if total_count > 0 else 0,
        'has_next': end_idx < total_count,
        'has_previous': page > 1,
        'findings': serializer.data
    })

@api_view(['GET'])
def scan_technology_checks(request, scan_id):

    scan = _scoped_scan(request, scan_id)

    checks = TechnologyCheck.objects.filter(scan=scan).select_related('asset')

    serializer = TechnologyCheckSerializer(checks, many=True)

    return Response({
        'scan_id': scan_id,
        'count': checks.count(),
        'checks': serializer.data
    })


# ============================================
# NUCLEI SCAN EXECUTION WITH LIBRARY & CVE CHECKS
# ============================================

@api_view(['POST'])
def execute_nuclei_scan(request):
    """
    Execute comprehensive security scan including:
    - Nuclei vulnerability scanning
    - Frontend library detection  
    - CVE vulnerability checking
    - Organization name extraction
    
    POST /api/nuclei/scan/
    Body: {
        "asset_ids": [1, 2],  # ← REQUIRED!
        "templates": ["ssl", "dns", "email"],
        "scan_type": "on-demand",
        "check_libraries": true,
        "check_cves": true,
        "extract_org": true
    }
    """
    asset_ids = request.data.get('asset_ids', [])
    templates = request.data.get('templates', ['ssl'])
    expanded_templates = []
    for template in templates:
        mapped = get_template_path(template)
        if not mapped:
            continue

        # Handle comma separated template paths
        if "," in mapped:
            expanded_templates.extend([item for item in mapped.split(",") if item])
        else:
            expanded_templates.append(mapped)

    templates = expanded_templates
    scan_type = request.data.get('scan_type', 'on-demand')
    check_libraries = request.data.get('check_libraries', True)
    check_cves = request.data.get('check_cves', True)
    extract_org = request.data.get('extract_org', True)
    
    # CRITICAL: Validate asset_ids
    if not asset_ids:
        return Response(
            {'error': 'asset_ids required - must provide at least one asset ID'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    # Get assets. Platform admins can quick-scan any organisation's assets, so
    # they are not bound to a single org — the scan is stamped with the org that
    # owns the targeted assets. An admin may still narrow to one org via
    # organisation_id (query param or X-Organisation-Id header). Regular users
    # are scoped to their own organisation.
    if is_platform_admin(getattr(request, 'user', None)):
        assets = Asset.objects.filter(id__in=asset_ids)
        requested = requested_org_id(request)
        if requested:
            assets = assets.filter(organisation_id=requested)
    else:
        org = get_user_organisation(request)
        assets = Asset.objects.filter(id__in=asset_ids, organisation=org)
    if not assets.exists():
        return Response(
            {'error': f'No valid assets found for IDs: {asset_ids}'},
            status=status.HTTP_400_BAD_REQUEST
        )

    # Scan record is stamped with the org that owns the scanned assets.
    org = assets.first().organisation

    active_assets, inactive_assets = _split_active_assets(assets)
    if inactive_assets:
        logger.warning(
            f"On-demand scan precheck: skipping {len(inactive_assets)} inactive targets"
        )

    if not active_assets:
        return Response(
            {
                'error': 'All targets are inactive/unreachable. Nothing to scan.',
                'inactive_targets': inactive_assets,
            },
            status=status.HTTP_400_BAD_REQUEST,
        )

    assets = active_assets
    
    # Create scan record
    scan = Scan.objects.create(
        organisation=org,
        scan_type=scan_type,
        status='running',
        started_at=timezone.now()
    )
    
    # Link assets
    for asset in assets:
        ScanAsset.objects.create(scan=scan, asset=asset)
    
    try:
        # Execute scans for each asset
        for asset in assets:
            logger.info(f"=== STARTING SCAN FOR {asset.value} ===")
            
            # 1. Extract organization name if requested
            if extract_org:
                logger.info(f"🔍 Extracting organization name for {asset.value}")
                try:
                    # Ensure we have a domain object
                    if not hasattr(asset, 'domain') or asset.domain is None:
                        logger.warning(f"⚠️ Asset {asset.id} has no associated domain - skipping org extraction")
                    else:
                        # Only extract if not already set
                        if not asset.domain.owner:
                            # Ensure URL has protocol for org extraction
                            asset_url = asset.value
                            if not asset_url.startswith(('http://', 'https://')):
                                asset_url = f'https://{asset_url}'
                                logger.info(f"Added protocol: {asset_url}")
                            
                            logger.info(f"Calling extract_organization_name for {asset_url}")
                            org_name = extract_organization_name(asset_url)
                            
                            if org_name:
                                asset.domain.owner = org_name
                                asset.domain.save()
                                logger.info(f"✅ Organization saved: '{org_name}' for domain {asset.domain.root_domain}")
                            else:
                                logger.warning(f"⚠️ Could not extract organization name for {asset_url}")
                        else:
                            logger.info(f"ℹ️ Organization already set: {asset.domain.owner}")
                except Exception as e:
                    logger.error(f"❌ Org extraction failed for {asset.value}: {e}")
                    import traceback
                    logger.error(f"Traceback:\n{traceback.format_exc()}")

            # 2. Run Nuclei scan
            logger.info(f"Running Nuclei scan on {asset.value}")
            try:
                results = run_nuclei_scan(asset.value, templates=templates)
                if results:
                    _process_nuclei_results(scan, asset, results, templates)
                    logger.info(f"✓ Nuclei scan completed: {len(results)} results")
                else:
                    logger.info(f"ℹ️ Nuclei scan skipped or returned no results (templates may be unavailable on Windows)")
            except Exception as e:
                logger.warning(f"⚠️ Nuclei scan failed: {e}. Continuing with other checks...")
            
            # 2.5. Check security headers directly via HTTP
            logger.info(f"🔍 === SECURITY HEADERS CHECK STARTED ===")
            try:
                asset_url = asset.value
                if not asset_url.startswith(('http://', 'https://')):
                    asset_url = f'https://{asset_url}'
                    logger.info(f"Added protocol: {asset_url}")
                
                headers_findings = check_security_headers(asset_url)
                _process_header_checks(scan, asset, headers_findings)
                
                missing_count = sum(1 for f in headers_findings.values() if f['status'] == 'missing')
                present_count = sum(1 for f in headers_findings.values() if f['status'] == 'present')
                logger.info(f"✓ Security headers check completed: {present_count} present, {missing_count} missing")
            except Exception as e:
                logger.error(f"❌ Security headers check failed: {e}")
                import traceback
                logger.error(f"Traceback:\n{traceback.format_exc()}")
            
            # 3. Check for libraries and CVEs
            if check_libraries:
                logger.info(f"🔍 === LIBRARY DETECTION STARTED ===")
                
                try:
                    # Ensure URL has protocol
                    asset_url = asset.value
                    if not asset_url.startswith(('http://', 'https://')):
                        asset_url = f'https://{asset_url}'
                        logger.info(f"Added protocol: {asset_url}")
                    
                    # Detect libraries and CMS
                    logger.info(f"Calling detect_technologies for {asset_url}")
                    tech_result = detect_technologies(asset_url)

                    _process_technology_results(scan, asset, tech_result)
                    
                    detected_libraries = tech_result.get('libraries', [])
                    detected_libraries = _augment_library_detections_with_analytics(tech_result, detected_libraries)
                    detected_cms = tech_result.get('cms')
                    
                    logger.info(f"✓ Detected {len(detected_libraries)} libraries")
                    
                    if not detected_libraries:
                        logger.warning("⚠️ No libraries detected - this might be normal for static sites")
                    else:
                        # Build library list string
                        lib_list = ', '.join([f"{lib.get('name')}@{lib.get('version')}" for lib in detected_libraries])
                        logger.info(f"Libraries found: {lib_list}")
                    
                    # Process each library - NEVER SKIP!
                    for lib in detected_libraries:
                        lib_name = lib.get('name')
                        lib_version = lib.get('version', 'Unknown')  # Default to 'Unknown'
                        lib_source = lib.get('source', 'unknown')
                        latest_version = 'Unknown'
                        
                        logger.info(f"\n--- Processing: {lib_name}@{lib_version} (source: {lib_source}) ---")
                        
                        # Initialize defaults
                        vulnerabilities = []
                        vuln_status = 'unknown'
                        risk_level = 'Low'
                        max_cvss = 0.0
                        recommendation = f"Version detection incomplete for {lib_name}"
                        
                        # Only check CVEs if we have a valid version
                        can_check_cve = (
                            lib_version and 
                            lib_version.strip() != '' and
                            lib_version.lower() not in ['unknown', 'latest', 'saas', 'n/a']
                        )

                        try:
                            ecosystem = get_ecosystem_for_library(lib_name)
                            resolved_latest_version = fetch_latest_library_version(lib_name, ecosystem)
                            if resolved_latest_version:
                                latest_version = resolved_latest_version
                                logger.info(f"📦 Latest registry version for {lib_name}: {latest_version}")
                        except Exception as e:
                            logger.warning(f"⚠️ Latest version lookup failed for {lib_name}: {e}")
                        
                        if can_check_cve and check_cves:
                            logger.info(f"🔍 Checking CVEs for {lib_name}@{lib_version}")
                            
                            try:
                                # Normalize library name for CVE lookup
                                lib_name_normalized = lib_name.lower().strip()
                                lib_name_normalized = re.sub(r'\.(js|css)$', '', lib_name_normalized)
                                lib_name_normalized = re.sub(r'\s+', '-', lib_name_normalized)
                                
                                logger.info(f"Normalized: '{lib_name}' → '{lib_name_normalized}'")
                                
                                # Get ecosystem
                                ecosystem = get_ecosystem_for_library(lib_name_normalized)
                                logger.info(f"Ecosystem: {ecosystem}")
                                
                                # Check vulnerabilities
                                cve_result = check_library_vulnerabilities(
                                    lib_name_normalized,
                                    lib_version,
                                    ecosystem,
                                    use_all_sources=True
                                )
                                
                                vulnerabilities = cve_result.get('vulnerabilities', [])
                                max_cvss = cve_result.get('max_cvss_score', 0.0)
                                
                                logger.info(f"CVE check complete: {len(vulnerabilities)} vulns, max CVSS: {max_cvss}")
                                
                                if vulnerabilities:
                                    vuln_status = 'vulnerable'
                                    
                                    # Determine risk level based on CVSS
                                    if max_cvss >= 9.0:
                                        risk_level = 'Critical'
                                    elif max_cvss >= 7.0:
                                        risk_level = 'High'
                                    elif max_cvss >= 4.0:
                                        risk_level = 'Medium'
                                    else:
                                        risk_level = 'Low'
                                    
                                    recommendation = f"Update {lib_name} immediately - {len(vulnerabilities)} CVEs found"
                                    logger.info(f"⚠️ {len(vulnerabilities)} CVEs found - Risk: {risk_level}")
                                    
                                    # Log first 3 CVEs
                                    for i, vuln in enumerate(vulnerabilities[:3], 1):
                                        cve_id = vuln.get('cve_id') or vuln.get('primary_cve') or vuln.get('id')
                                        logger.info(f"  {i}. {cve_id}: CVSS {vuln.get('cvss_score', 0.0)}")
                                else:
                                    if latest_version != 'Unknown' and version_compare(lib_version, latest_version) < 0:
                                        vuln_status = 'outdated'
                                        recommendation = f"Update {lib_name} from {lib_version} to {latest_version}"
                                        logger.info(f"⚠️ No CVEs found, but {lib_name} is outdated ({lib_version} < {latest_version})")
                                    else:
                                        vuln_status = 'up-to-date'
                                        recommendation = f"No known vulnerabilities for {lib_name} {lib_version}"
                                        logger.info(f"✓ No CVEs found")
                                    
                            except Exception as e:
                                logger.error(f"❌ CVE check failed for {lib_name}: {e}")
                                import traceback
                                logger.error(f"Traceback:\n{traceback.format_exc()}")
                                vuln_status = 'check-failed'
                                recommendation = f"CVE check failed for {lib_name} - manual review recommended"
                        
                        else:
                            # Can't check CVEs, but still save the library
                            logger.warning(f"⚠️ Cannot check CVEs for {lib_name} (version: {lib_version})")
                            vuln_status = 'unknown'
                            recommendation = f"Manual version verification needed for {lib_name}"
                        
                        # ALWAYS create the library check record - NEVER SKIP
                        library_check, created = FrontendLibraryCheck.objects.get_or_create(
                            scan=scan,
                            asset=asset,
                            library_name=lib_name,
                            detected_version=lib_version,
                            defaults={
                                'latest_version': latest_version,
                                'vulnerability_status': vuln_status,
                                'risk_level': risk_level,
                                'source_urls': [lib.get('source_url', asset_url)],
                                'recommendation': recommendation
                            }
                        )

                        if not created:
                            updates = []
                            if library_check.latest_version != latest_version:
                                library_check.latest_version = latest_version
                                updates.append('latest_version')
                            if library_check.vulnerability_status != vuln_status:
                                library_check.vulnerability_status = vuln_status
                                updates.append('vulnerability_status')
                            if library_check.risk_level != risk_level:
                                library_check.risk_level = risk_level
                                updates.append('risk_level')
                            if library_check.recommendation != recommendation:
                                library_check.recommendation = recommendation
                                updates.append('recommendation')
                            if updates:
                                library_check.save(update_fields=updates)
                        
                        if created:
                            logger.info(f"✓✓✓ SAVED to DB: FrontendLibraryCheck id={library_check.id}")
                            logger.info(f"✓ Created library check: {lib_name} v{lib_version} - {vuln_status} ({risk_level})")
                        else:
                            logger.info(f"ℹ️ Already exists: {lib_name} v{lib_version}")

                        # ============================================================
                        # VULNERABILITY PROCESSING - IMPROVED CVE ID EXTRACTION
                        # ============================================================
                        logger.info(f"\n{'='*80}")
                        logger.info(f"VULNERABILITY PROCESSING FOR {lib_name} v{lib_version}")
                        logger.info(f"{'='*80}")
                        logger.info(f"Total vulnerabilities found: {len(vulnerabilities)}")
                        logger.info(f"Max CVSS score: {max_cvss}")
                        logger.info(f"Risk level: {risk_level}")
                        logger.info(f"Asset: {asset.value} (ID={asset.id})")
                        logger.info(f"Scan: {scan.id}")

                        if not vulnerabilities:
                            logger.warning(f"⚠️ No vulnerabilities to process for {lib_name} v{lib_version}")
                        else:
                            logger.info(f"Processing {len(vulnerabilities)} vulnerabilities...")
                            # Print summary of vulnerabilities with IMPROVED extraction
                            for i, v in enumerate(vulnerabilities, 1):
                                # ✅ FIXED: Try multiple fields for CVE ID extraction
                                v_id = (
                                    v.get('cve_id') or 
                                    v.get('primary_cve') or 
                                    (v.get('cve_ids', [None])[0] if v.get('cve_ids') else None) or
                                    v.get('id', 'UNKNOWN')
                                )
                                v_cvss = v.get('cvss_score', 0.0)
                                v_source = v.get('source', 'UNKNOWN')
                                logger.info(f"  {i}. {v_id} (CVSS: {v_cvss}) [Source: {v_source}]")
                        
                        
                        # Create findings for CVEs (only if vulnerabilities found)
                        for idx, vuln in enumerate(vulnerabilities, 1):
                            logger.info(f"\n{'='*60}")
                            logger.info(f"Processing vulnerability {idx}/{len(vulnerabilities)} for {lib_name} v{lib_version}")
                            
                            # ✅ FIXED: Extract CVE ID with IMPROVED handling for Retire.js
                            # cve_id = (
                            #     vuln.get('cve_id') or                                          # OSV format / Fixed Retire.js
                            #     vuln.get('primary_cve') or                                     # Retire.js (original)
                            #     (vuln.get('cve_ids', [None])[0] if vuln.get('cve_ids') else None) or  # ← NEW: Handle list
                            #     vuln.get('id', 'UNKNOWN')                                      # Fallback
                            # )
                            
                            # logger.info(f"CVE ID: {cve_id}")
                            # logger.info(f"Source: {vuln.get('source', 'UNKNOWN')}")

                            # Extract vulnerability ID
                            vuln_id = (
                                vuln.get('cve_id') or 
                                vuln.get('primary_cve') or 
                                (vuln.get('cve_ids', [None])[0] if vuln.get('cve_ids') else None) or
                                vuln.get('id', 'UNKNOWN')
                            )

                            # Determine if this is a CVE-based vulnerability
                            is_cve_vulnerability = vuln_id and vuln_id.startswith('CVE-')

                            logger.info(f"Vulnerability ID: {vuln_id} (CVE: {is_cve_vulnerability})")
                            logger.info(f"Source: {vuln.get('source', 'UNKNOWN')}")
                            
                            # Debug: Print vulnerability structure if needed
                            if logger.isEnabledFor(logging.DEBUG):
                                logger.debug(f"Vulnerability data: {json.dumps(vuln, indent=2, default=str)}")
                            
                            # Extract CVSS score with fallback to severity mapping
                            cvss_score = vuln.get('cvss_score', 0.0)
                            if cvss_score is None or cvss_score == 0.0:
                                cvss_score = max_cvss

                            try:
                                cvss_score = float(cvss_score)
                            except (ValueError, TypeError):
                                cvss_score = 0.0

                            # If still 0, use severity mapping
                            if cvss_score == 0.0:
                                severity = vuln.get('severity', 'UNKNOWN').upper()
                                severity_map = {
                                    'CRITICAL': 9.5,
                                    'HIGH': 7.5,
                                    'MEDIUM': 5.0,
                                    'MODERATE': 5.0,
                                    'LOW': 3.0,
                                    'UNKNOWN': 0.0
                                }
                                cvss_score = severity_map.get(severity, 5.0)
                                logger.info(f"Using severity-based CVSS: {cvss_score} (from {severity})")

                            cvss_vector = vuln.get('cvss_vector', '')

                            logger.info(f"📊 CVSS Score: {cvss_score}")
                            
                            # Step 1: Create CVE record ONLY if this is a CVE vulnerability
                            cve_record = None
                            if is_cve_vulnerability:
                                logger.info(f"📝 Creating/updating CVE record for {vuln_id}")
                                cve_record, cve_created = CVE.objects.get_or_create(
                                    cve_id=vuln_id,
                                    defaults={
                                        'cvss_score': cvss_score,
                                        'cvss_vector': cvss_vector or '',
                                        'description': (
                                            vuln.get('summary') or 
                                            vuln.get('description') or 
                                            vuln.get('details', '')
                                        )[:500],
                                        'published_date': vuln.get('published'),
                                        'last_modified': vuln.get('modified')
                                    }
                                )
                                
                                if cve_created:
                                    logger.info(f"✅ Created new CVE record: {cve_id}")
                                else:
                                    # Update existing CVE if we have better data
                                    updated = False
                                    if cvss_score > 0.0 and (not cve_record.cvss_score or cve_record.cvss_score == 0.0):
                                        cve_record.cvss_score = cvss_score
                                        updated = True
                                    if cvss_vector and not cve_record.cvss_vector:
                                        cve_record.cvss_vector = cvss_vector
                                        updated = True
                                    if updated:
                                        cve_record.save()
                                        logger.info(f"Updated CVE record: {cve_id}")
                                    else:
                                        logger.info(f"CVE record already exists: {cve_id}")
                            else:
                                logger.info(f"Non-CVE vulnerability: {vuln_id} - will create Finding without CVE link")
                            
                            # Step 2: Create Finding - IMPROVED title for non-CVE vulnerabilities
                            if is_cve_vulnerability:
                                finding_title = f"{lib_name} {lib_version} - {vuln_id}"
                            else:
                                # For non-CVE vulnerabilities, use source and summary
                                source = vuln.get('source', 'Security')
                                summary = vuln.get('summary', 'Vulnerability')
                                
                                # Clean up summary for title (take first line or first 60 chars)
                                if summary:
                                    summary_clean = summary.split('\n')[0].split('http')[0].strip()
                                    if len(summary_clean) > 60:
                                        summary_clean = summary_clean[:57] + '...'
                                else:
                                    summary_clean = 'Vulnerability'
                                
                                finding_title = f"{lib_name} {lib_version} - {source} Vulnerability: {summary_clean}"

                            logger.info(f"Finding title: {finding_title}")
                            
                            # FIXED: More specific deduplication - match exact title + asset
                            existing_finding = Finding.objects.filter(
                                asset=asset,
                                title=finding_title,  # ← FIXED: Exact match, not __icontains
                                category='CVE'
                            ).first()
                            
                            if existing_finding:
                                logger.info(f"📎 Finding already exists (ID={existing_finding.id}): {finding_title}")
                                
                                # Link existing finding to current scan
                                scan_finding, sf_created = ScanFinding.objects.get_or_create(
                                    scan=scan,
                                    finding=existing_finding
                                )
                                if sf_created:
                                    logger.info(f"✅ Linked existing finding to scan {scan.id}")
                                else:
                                    logger.info(f"ℹ️ Finding already linked to scan {scan.id}")
                            else:
                                # CREATE NEW FINDING
                                logger.info(f"💾 Creating new finding: {finding_title}")
                                logger.info(f"   Asset: {asset.value}")
                                logger.info(f"   CVSS Score: {cvss_score}")
                                logger.info(f"   Risk Level: {risk_level}")
                                
                                try:
                                    # Get recommendation
                                    recommendation_text = (
                                        vuln.get('recommendation') or
                                        vuln.get('remediation') or
                                        vuln.get('details') or 
                                        vuln.get('summary') or 
                                        f'Update {lib_name} to a patched version'
                                    )
                                    
                                    # Get evidence/description
                                    evidence_text = (
                                        vuln.get('summary') or
                                        vuln.get('description') or
                                        vuln.get('details') or
                                        f"Vulnerable library: {lib_name} v{lib_version}"
                                    )

                                    # Add vulnerability source and references to evidence for non-CVE vulns
                                    if not is_cve_vulnerability:
                                        refs = vuln.get('references', [])
                                        if refs:
                                            evidence_text += f"\n\nReferences:\n" + "\n".join(refs[:3])

                                    finding = Finding.objects.create(
                                        asset=asset,
                                        title=finding_title,
                                        category='CVE',
                                        nuclei_template_id='library-vuln-check',
                                        nuclei_severity=vuln.get('severity', 'medium').lower(),
                                        cvss_score=cvss_score,
                                        cvss_vector=cvss_vector or '',
                                        risk_rating=risk_level,
                                        scoring_confidence='High',
                                        evidence=evidence_text[:1000],  # Limit to 1000 chars
                                        recommendation=recommendation_text[:500],  # Limit to 500 chars
                                        status='open'
                                    )
                                    
                                    logger.info(f"✅✅✅ CREATED Finding ID={finding.id}: {finding_title}")
                                    logger.info(f"   Database ID: {finding.id}")
                                    logger.info(f"   CVSS Score in DB: {finding.cvss_score}")
                                    
                                    # Link finding to scan
                                    scan_finding = ScanFinding.objects.create(
                                        scan=scan, 
                                        finding=finding
                                    )
                                    logger.info(f"✅ Linked finding to scan via ScanFinding ID={scan_finding.id}")
                                    
                                    # Link finding to CVE (ONLY if we have a CVE record)
                                    if cve_record:
                                        FindingCVE.objects.create(
                                            finding=finding,
                                            cve=cve_record,
                                            relevance='direct'
                                        )
                                        logger.info(f"✅ Linked finding to CVE {vuln_id}")
                                    else:
                                        logger.info(f"ℹ️ Non-CVE vulnerability - no CVE link created")
                                    
                                    # Verify the finding was created
                                    verify_finding = Finding.objects.filter(id=finding.id).first()
                                    if verify_finding:
                                        logger.info(f"✅ VERIFICATION: Finding exists in database")
                                        logger.info(f"   Title: {verify_finding.title}")
                                        logger.info(f"   CVSS: {verify_finding.cvss_score}")
                                        logger.info(f"   Category: {verify_finding.category}")
                                    else:
                                        logger.error(f"❌ VERIFICATION FAILED: Finding not found in database!")
                                        
                                except Exception as e:
                                    logger.error(f"❌ FAILED to create finding: {e}")
                                    import traceback
                                    logger.error(f"Traceback:\n{traceback.format_exc()}")
                                    # Don't raise - continue processing other vulnerabilities
                        
                        logger.info(f"\n{'='*80}")
                        logger.info(f"FINISHED processing {len(vulnerabilities)} vulnerabilities for {lib_name}")
                        logger.info(f"{'='*80}\n")
                    
                    # Store CMS info if detected
                    if detected_cms:
                        logger.info(f"🔍 Detected CMS: {detected_cms['name']} v{detected_cms['version']}")
                        
                        # Create informational finding for CMS
                        if detected_cms['version'] not in ['Unknown', 'SaaS']:
                            existing_cms_finding = Finding.objects.filter(
                                asset=asset,
                                title__icontains=f"{detected_cms['name']} CMS"
                            ).first()
                            
                            if not existing_cms_finding:
                                finding = Finding.objects.create(
                                    asset=asset,
                                    title=f"{detected_cms['name']} CMS Detected - {detected_cms['version']}",
                                    category='Misconfiguration',
                                    nuclei_template_id='cms-detection',
                                    nuclei_severity='info',
                                    risk_rating='Low',
                                    scoring_confidence='High',
                                    evidence=f"CMS: {detected_cms['name']} v{detected_cms['version']}",
                                    recommendation=f"Ensure {detected_cms['name']} is up to date and properly configured",
                                    status='open'
                                )
                                ScanFinding.objects.create(scan=scan, finding=finding)
                                logger.info(f"✓ Created CMS finding")
                            else:
                                # Link existing CMS finding to current scan
                                ScanFinding.objects.get_or_create(
                                    scan=scan,
                                    finding=existing_cms_finding
                                )
                                logger.info(f"📎 Linked existing CMS finding")
                
                except Exception as e:
                    logger.error(f"❌ Library/CVE detection failed for {asset.value}: {e}")
                    import traceback
                    logger.error(f"Full traceback:\n{traceback.format_exc()}")
        
        # Mark scan as completed
        scan.status = 'completed'
        scan.finished_at = timezone.now()
        scan.duration_seconds = (scan.finished_at - scan.started_at).total_seconds()
        scan.save()
        
        logger.info(f"✅ Scan {scan.id} completed in {scan.duration_seconds}s")
        
        return Response({
            'message': 'Scan completed successfully',
            'scan_id': scan.id,
            'status': 'completed',
            'duration_seconds': scan.duration_seconds,
            'scanned_assets_count': len(assets),
            'skipped_assets_count': len(inactive_assets),
            'skipped_targets': inactive_assets,
        })
        
    except Exception as e:
        logger.error(f"❌ Scan {scan.id} failed: {e}")
        import traceback
        logger.error(f"Full traceback:\n{traceback.format_exc()}")
        
        scan.status = 'failed'
        scan.error_message = str(e)
        scan.finished_at = timezone.now()
        scan.save()
        
        return Response(
            {'error': f'Scan failed: {str(e)}'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
def execute_domain_scan(request, domain_id):
    """
    Execute scan for all assets under a domain in one request.

    POST /api/domains/{domain_id}/scan/
    Body (optional):
    {
        "templates": ["ssl", "dns", "email"],
        "scan_type": "on-demand",
        "check_libraries": true,
        "check_cves": true,
        "extract_org": true
    }
    """
    domain = _scoped_domain(request, domain_id)
    asset_ids = list(domain.assets.values_list('id', flat=True))

    if not asset_ids:
        return Response(
            {
                'error': f'No assets found for domain {domain.root_domain}',
                'domain_id': domain.id,
                'domain': domain.root_domain,
            },
            status=status.HTTP_400_BAD_REQUEST,
        )

    payload = request.data.copy() if hasattr(request.data, 'copy') else dict(request.data)
    payload['asset_ids'] = asset_ids

    # execute_nuclei_scan is itself an @api_view, so it re-authenticates the
    # request we hand it. Forward the caller's auth (and any org-targeting
    # header) so that inner authentication succeeds instead of 401-ing.
    forwarded = {}
    auth_header = request.META.get('HTTP_AUTHORIZATION')
    if auth_header:
        forwarded['HTTP_AUTHORIZATION'] = auth_header
    org_header = request.META.get('HTTP_X_ORGANISATION_ID')
    if org_header:
        forwarded['HTTP_X_ORGANISATION_ID'] = org_header

    factory = APIRequestFactory()
    delegated_request = factory.post('/api/nuclei/scan/', payload, format='json', **forwarded)

    return execute_nuclei_scan(delegated_request)

def _process_nuclei_results(scan, asset, results, templates):
    """
    Process Nuclei scan results and store in appropriate tables
    WITH DEDUPLICATION
    """
    processed_findings = set()

    logger.info(f"Processing {len(results)} Nuclei scan results for {asset.value}")
    logger.info(f"Full Nuclei scan results for {asset.value}: {json.dumps(results, indent=2, default=str)}")
    
    for result in results:
        template_id = result.get('template-id', '')
        template_name = result.get('info', {}).get('name', '')
        severity = result.get('info', {}).get('severity', 'info')
        
        # Create unique key for deduplication
        finding_key = f"{template_id}_{asset.id}_{template_name}"
        
        # Skip if already processed
        if finding_key in processed_findings:
            logger.debug(f"Skipping duplicate finding: {template_name}")
            continue
        
        processed_findings.add(finding_key)
        
        # Process by category (non-exclusive - a result can belong to multiple categories)
        template_lower = template_id.lower()
        
        if 'ssl' in template_lower or 'tls' in template_lower:
            _create_ssl_check(scan, asset, result)
        
        if 'dns' in template_lower:
            _create_dns_check(scan, asset, result)
        
        if 'email' in template_lower or 'spf' in template_lower or 'dmarc' in template_lower or 'dkim' in template_lower:
            _create_email_check(scan, asset, result)
        
        if 'header' in template_lower:
            _create_header_check(scan, asset, result)
        
        if 'javascript' in template_lower or 'library' in template_lower:
            _create_library_check(scan, asset, result)
        
        # Always create a Finding record
        _create_finding(scan, asset, result)

def _process_technology_results(scan, asset, tech_result):
    """
    Save detected technologies and CMS into DB
    """

    def _save_technology(tech_name, tech_version='Unknown', tech_category='general'):
        if not tech_name:
            return

        latest_version = 'Unknown'

        if supports_latest_version_lookup(tech_name):
            try:
                resolved_latest = fetch_latest_library_version(tech_name)
                if resolved_latest:
                    latest_version = resolved_latest
            except Exception as e:
                logger.warning(f"⚠️ Latest version lookup failed for technology {tech_name}: {e}")

        technology_check, created = TechnologyCheck.objects.get_or_create(
            scan=scan,
            asset=asset,
            technology_name=tech_name,
            version=tech_version,
            category=tech_category,
            defaults={
                'latest_version': latest_version,
            }
        )

        if not created and technology_check.latest_version != latest_version:
            technology_check.latest_version = latest_version
            technology_check.save(update_fields=['latest_version'])

    technologies = list(tech_result.get('technologies', []))
    cms = tech_result.get('cms')
    web_server = tech_result.get('web_server')
    programming_languages = tech_result.get('programming_languages', [])
    cdn = tech_result.get('cdn', [])
    analytics = tech_result.get('analytics', [])
    security = tech_result.get('security', [])
    cms_extensions = tech_result.get('cms_extensions', {})

    if web_server:
        technologies.append({
            'name': web_server.get('name'),
            'version': web_server.get('version', 'Unknown'),
            'category': web_server.get('type', 'Web Server')
        })

    for item in programming_languages + cdn + analytics + security:
        technologies.append({
            'name': item.get('name'),
            'version': item.get('version', 'Unknown'),
            'category': item.get('type', 'general')
        })

    # ---------------- GENERAL TECHNOLOGIES ----------------
    for tech in technologies:
        _save_technology(
            tech.get('name'),
            tech.get('version', 'Unknown'),
            tech.get('category') or tech.get('type', 'general')
        )

    # ---------------- CMS ----------------
    if cms:
        _save_technology(
            cms.get('name'),
            cms.get('version', 'Unknown'),
            'CMS'
        )

    # ---------------- CMS EXTENSIONS ----------------
    for plugin in cms_extensions.get('plugins', []):
        _save_technology(plugin.get('name'), plugin.get('version', 'Unknown'), plugin.get('type', 'CMS Plugin'))

    theme = cms_extensions.get('themes')
    if theme:
        _save_technology(theme.get('name'), theme.get('version', 'Unknown'), theme.get('type', 'CMS Theme'))

    for module in cms_extensions.get('modules', []):
        _save_technology(module.get('name'), module.get('version', 'Unknown'), module.get('type', 'CMS Module'))

    for extension in cms_extensions.get('extensions', []):
        _save_technology(extension.get('name'), extension.get('version', 'Unknown'), extension.get('type', 'CMS Extension'))

# def _process_library_detection(scan, asset, tech_result, check_cves):

#     detected_libraries = tech_result.get('libraries', [])

#     for lib in detected_libraries:

#         lib_name = lib.get('name')
#         lib_version = lib.get('version', 'Unknown')

#         vulnerabilities = []
#         vuln_status = 'unknown'
#         risk_level = 'Low'
#         recommendation = f"Manual version verification needed for {lib_name}"

#         # CVE checking
#         if check_cves and lib_version not in ['Unknown', 'unknown', 'latest']:

#             try:
#                 ecosystem = get_ecosystem_for_library(lib_name.lower())

#                 cve_result = check_library_vulnerabilities(
#                     lib_name,
#                     lib_version,
#                     ecosystem
#                 )

#                 vulnerabilities = cve_result.get('vulnerabilities', [])
#                 max_cvss = cve_result.get('max_cvss_score', 0.0)

#                 if vulnerabilities:
#                     vuln_status = 'vulnerable'

#                     if max_cvss >= 9:
#                         risk_level = 'Critical'
#                     elif max_cvss >= 7:
#                         risk_level = 'High'
#                     elif max_cvss >= 4:
#                         risk_level = 'Medium'

#                     recommendation = f"Update {lib_name} immediately"

#                 else:
#                     vuln_status = 'up-to-date'

#             except Exception as e:
#                 logger.error(f"CVE check failed: {e}")
#                 vuln_status = 'check-failed'

#         FrontendLibraryCheck.objects.get_or_create(
#             scan=scan,
#             asset=asset,
#             library_name=lib_name,
#             detected_version=lib_version,
#             defaults={
#                 'latest_version': 'Unknown',
#                 'vulnerability_status': vuln_status,
#                 'risk_level': risk_level,
#                 'recommendation': recommendation
#             }
#         )


def _create_ssl_check(scan, asset, result):
    """Create SSL/TLS check record with deduplication"""
    template_id = result.get('template-id', '')
    info = result.get('info', {})
    
    # Determine check type
    if 'certificate' in template_id or 'cert' in template_id:
        check_type = 'certificate'
    elif 'cipher' in template_id:
        check_type = 'cipher'
    elif 'protocol' in template_id or 'tls-version' in template_id:
        check_type = 'protocol'
    else:
        check_type = 'hsts'
    
    finding_name = info.get('name', '')
    
    # Check for duplicate
    existing = SSLTLSCheck.objects.filter(
        scan=scan,
        asset=asset,
        check_type=check_type,
        finding=finding_name
    ).first()
    
    if existing:
        logger.debug(f"Skipping duplicate SSL check: {finding_name}")
        return
    
    classification = info.get('classification', {})
    cvss_score = classification.get('cvss-score', 0.0)
    try:
        cvss_score = float(cvss_score)
    except (ValueError, TypeError):
        cvss_score = 0.0
    
    SSLTLSCheck.objects.create(
        scan=scan,
        asset=asset,
        check_type=check_type,
        finding=finding_name,
        example=result.get('matched-at', ''),
        cvss_score=cvss_score,
        risk_rating=_map_severity_to_risk(info.get('severity', 'info')),
        recommendation=info.get('remediation', '')
    )

    


def _create_dns_check(scan, asset, result):
    """Create DNS check record with deduplication"""
    template_id = result.get('template-id', '')
    info = result.get('info', {})
    
    # Determine check type
    if 'dnssec' in template_id:
        check_type = 'dnssec'
    elif 'zone-transfer' in template_id or 'axfr' in template_id:
        check_type = 'zone_transfer'
    elif 'hijack' in template_id:
        check_type = 'hijacking'
    else:
        check_type = 'subdomain_takeover'
    
    finding_name = info.get('name', '')
    
    # Check for duplicate
    existing = DNSSecurityCheck.objects.filter(
        scan=scan,
        asset=asset,
        check_type=check_type,
        finding=finding_name
    ).first()
    
    if existing:
        return
    
    DNSSecurityCheck.objects.create(
        scan=scan,
        asset=asset,
        check_type=check_type,
        finding=finding_name,
        example=result.get('matched-at', ''),
        cvss_score=info.get('classification', {}).get('cvss-score', 0.0),
        risk_rating=_map_severity_to_risk(info.get('severity', 'info')),
        recommendation=info.get('remediation', '')
    )


def _create_email_check(scan, asset, result):
    """Create email security check record with deduplication"""
    template_id = result.get('template-id', '').lower()
    info = result.get('info', {})
    
    # Determine check type
    if 'spf' in template_id:
        check_type = 'SPF'
    elif 'dkim' in template_id:
        check_type = 'DKIM'
    else:
        check_type = 'DMARC'
    
    # Check for duplicate
    existing = EmailSecurityCheck.objects.filter(
        scan=scan,
        asset=asset,
        check_type=check_type
    ).first()
    
    if existing:
        return
    
    # Determine status
    severity = info.get('severity', 'info')
    status_value = 'FAIL' if severity in ['high', 'critical', 'medium'] else 'PASS'
    
    EmailSecurityCheck.objects.create(
        scan=scan,
        asset=asset,
        check_type=check_type,
        status=status_value,
        details=info.get('description', ''),
        cvss_score=info.get('classification', {}).get('cvss-score', 0.0),
        risk_rating=_map_severity_to_risk(severity),
        recommendation=info.get('remediation', ''),
        record_value=result.get('extracted-results', [''])[0] if result.get('extracted-results') else ''
    )


def check_security_headers(target_url: str) -> Dict[str, Dict[str, str]]:
    """
    Directly check security headers on a target URL without Nuclei.
    
    Args:
        target_url (str): Target URL to check
    
    Returns:
        Dict: Security headers findings with status and value
    """
    # Security headers to check with risk ratings and recommendations
    SECURITY_HEADERS = {
        'Strict-Transport-Security': {
            'risk_rating': 'High',
            'recommendation': 'Enable HSTS to enforce HTTPS. Set "Strict-Transport-Security: max-age=31536000; includeSubDomains; preload"',
            'cvss_score': 7.5
        },
        'Content-Security-Policy': {
            'risk_rating': 'High',
            'recommendation': 'Implement Content-Security-Policy to prevent XSS and injection attacks.',
            'cvss_score': 6.1
        },
        'X-Frame-Options': {
            'risk_rating': 'Medium',
            'recommendation': 'Set "X-Frame-Options: DENY" to prevent clickjacking attacks.',
            'cvss_score': 5.3
        },
        'X-Content-Type-Options': {
            'risk_rating': 'Medium',
            'recommendation': 'Set "X-Content-Type-Options: nosniff" to prevent MIME type sniffing.',
            'cvss_score': 5.3
        },
        'Referrer-Policy': {
            'risk_rating': 'Low',
            'recommendation': 'Set "Referrer-Policy: strict-origin-when-cross-origin" to control referrer information.',
            'cvss_score': 3.7
        },
        'Permissions-Policy': {
            'risk_rating': 'Low',
            'recommendation': 'Implement Permissions-Policy to restrict browser features and APIs.',
            'cvss_score': 3.7
        },
        'X-XSS-Protection': {
            'risk_rating': 'Low',
            'recommendation': 'Set "X-XSS-Protection: 1; mode=block" for legacy browser protection.',
            'cvss_score': 3.7
        }
    }
    
    findings = {}
    
    try:
        # Ensure URL has protocol
        if not target_url.startswith(('http://', 'https://')):
            target_url = f'https://{target_url}'
        
        logger.info(f"[HEADER_CHECK] Checking security headers for {target_url}")
        
        # Make HEAD request first (faster), fall back to GET if needed
        try:
            response = requests.head(target_url, timeout=10, verify=False, allow_redirects=True)
        except:
            response = requests.get(target_url, timeout=10, verify=False, allow_redirects=True)
        
        response_headers = response.headers
        
        # Check each security header
        for header_name, header_info in SECURITY_HEADERS.items():
            header_value = response_headers.get(header_name, '')
            
            if header_value:
                status_val = 'present'
                risk_rating = 'Low'
                cvss_score = None
                logger.info(f"  ✓ {header_name}: PRESENT")
            else:
                status_val = 'missing'
                risk_rating = header_info['risk_rating']
                cvss_score = header_info['cvss_score']
                logger.info(f"  ✗ {header_name}: MISSING")
            
            findings[header_name] = {
                'status': status_val,
                'value': header_value,
                'risk_rating': risk_rating,
                'cvss_score': cvss_score,
                'recommendation': header_info['recommendation']
            }
        
        logger.info(f"[HEADER_CHECK] Completed for {target_url}")
        
    except requests.exceptions.RequestException as e:
        logger.error(f"[HEADER_CHECK] Failed to connect to {target_url}: {e}")
        # Return all headers as missing if we can't reach the target
        for header_name, header_info in SECURITY_HEADERS.items():
            findings[header_name] = {
                'status': 'missing',
                'value': '',
                'risk_rating': header_info['risk_rating'],
                'cvss_score': header_info['cvss_score'],
                'recommendation': header_info['recommendation']
            }
    except Exception as e:
        logger.error(f"[HEADER_CHECK] Unexpected error for {target_url}: {e}")
    
    return findings


def _process_header_checks(scan, asset, headers_findings: Dict[str, Dict[str, str]]):
    """
    Process security header findings and create SecurityHeaderCheck records.
    
    Args:
        scan: Scan object
        asset: Asset object
        headers_findings: Dict of header findings from check_security_headers()
    """
    for header_name, finding in headers_findings.items():
        # Check for duplicate
        existing = SecurityHeaderCheck.objects.filter(
            scan=scan,
            asset=asset,
            header=header_name
        ).first()
        
        if existing:
            logger.debug(f"Skipping duplicate header check: {header_name}")
            continue
        
        SecurityHeaderCheck.objects.create(
            scan=scan,
            asset=asset,
            header=header_name,
            status=finding['status'],
            cvss_score=finding['cvss_score'],
            risk_rating=finding['risk_rating'],
            recommendation=finding['recommendation'],
            header_value=finding['value']
        )
        logger.info(f"Created header check: {header_name} - {finding['status']}")


def _create_header_check(scan, asset, result):
    """Create security header check record with deduplication (legacy Nuclei-based)"""
    info = result.get('info', {})
    header_name = result.get('matcher-name', '') or info.get('name', '')
    
    # Check for duplicate
    existing = SecurityHeaderCheck.objects.filter(
        scan=scan,
        asset=asset,
        header=header_name
    ).first()
    
    if existing:
        return
    
    SecurityHeaderCheck.objects.create(
        scan=scan,
        asset=asset,
        header=header_name,
        status='missing',
        cvss_score=info.get('classification', {}).get('cvss-score', 0.0),
        risk_rating=_map_severity_to_risk(info.get('severity', 'info')),
        recommendation=info.get('remediation', ''),
        header_value=''
    )


def _create_library_check(scan, asset, result):
    """Create frontend library check record with deduplication"""
    info = result.get('info', {})
    
    library_name = result.get('matcher-name', 'unknown')
    detected_version = result.get('extracted-results', ['unknown'])[0] if result.get('extracted-results') else 'unknown'
    
    # Check for duplicate
    existing = FrontendLibraryCheck.objects.filter(
        scan=scan,
        asset=asset,
        library_name=library_name,
        detected_version=detected_version
    ).first()
    
    if existing:
        return
    
    FrontendLibraryCheck.objects.create(
        scan=scan,
        asset=asset,
        library_name=library_name,
        detected_version=detected_version,
        latest_version='unknown',
        vulnerability_status='vulnerable' if info.get('severity') in ['high', 'critical'] else 'outdated',
        risk_level=_map_severity_to_risk(info.get('severity', 'info')),
        source_urls=[result.get('matched-at', '')],
        recommendation=info.get('remediation', '')
    )


def _create_finding(scan, asset, result):
    """Create general finding record with deduplication - FIXED CVSS EXTRACTION"""
    info = result.get('info', {})
    title = info.get('name', '')
    template_id = result.get('template-id', '')
    
    # Check for duplicate
    existing = Finding.objects.filter(
        asset=asset,
        title=title,
        nuclei_template_id=template_id
    ).first()
    
    if existing:
        ScanFinding.objects.get_or_create(scan=scan, finding=existing)
        logger.debug(f"📎 Linked existing finding: {title}")
        return
    
    # ============ FIXED CVSS EXTRACTION ============
    classification = info.get('classification', {})
    
    cvss_score = None
    cvss_vector = None
    
    # Method 1: Try cvss-score field
    if 'cvss-score' in classification:
        try:
            cvss_score = float(classification['cvss-score'])
        except (ValueError, TypeError):
            pass
    
    # Method 2: Try parsing cvss-metrics vector
    if 'cvss-metrics' in classification and classification['cvss-metrics']:
        cvss_vector = classification['cvss-metrics']
        if cvss_score is None:
            score_match = re.search(r'/BS:(\d+\.\d+)', cvss_vector)
            if score_match:
                try:
                    cvss_score = float(score_match.group(1))
                except (ValueError, TypeError):
                    pass
    
    # Method 3: Map from severity if still None
    if cvss_score is None:
        severity = info.get('severity', 'info').lower()
        severity_to_cvss = {
            'critical': 9.5,
            'high': 7.5,
            'medium': 5.0,
            'low': 3.0,
            'info': 0.0
        }
        cvss_score = severity_to_cvss.get(severity, 0.0)
        logger.warning(f"⚠️ No CVSS score found for '{title}', using severity mapping: {cvss_score}")
    
    # Ensure cvss_score is a float
    try:
        cvss_score = float(cvss_score)
    except (ValueError, TypeError):
        cvss_score = 0.0
    
    logger.info(f"💾 Creating finding: {title} (CVSS: {cvss_score})")
    
    finding = Finding.objects.create(
        asset=asset,
        title=title,
        category=_determine_category(template_id),
        nuclei_template_id=template_id,
        nuclei_severity=info.get('severity', 'info'),
        cvss_score=cvss_score,  # ✅ FIXED: Now guaranteed to have valid float
        cvss_vector=cvss_vector or '',
        risk_rating=_map_severity_to_risk(info.get('severity', 'info')),
        scoring_confidence='High',
        evidence=result.get('matched-at', ''),
        recommendation=info.get('remediation', 'No remediation provided'),
        status='open'
    )
    
    ScanFinding.objects.create(scan=scan, finding=finding)
    logger.info(f"✅ Created finding ID={finding.id} with CVSS={cvss_score}")

def _map_severity_to_risk(severity):
    """Map Nuclei severity to risk rating"""
    severity_map = {
        'critical': 'Critical',
        'high': 'High',
        'medium': 'Medium',
        'low': 'Low',
        'info': 'Low'
    }
    return severity_map.get(severity.lower(), 'Low')


def _determine_category(template_id):
    """Determine finding category from template ID"""
    template_id = template_id.lower()
    
    if 'cve' in template_id:
        return 'CVE'
    elif 'ssl' in template_id or 'tls' in template_id:
        return 'SSL'
    elif 'dns' in template_id:
        return 'DNS'
    elif 'email' in template_id or 'spf' in template_id or 'dmarc' in template_id:
        return 'Email'
    else:
        return 'Misconfiguration'


# ============================================
# DASHBOARD & STATISTICS
# ============================================

@api_view(['GET'])
def dashboard_metrics(request):
    """
    Get dashboard metrics
    GET /api/dashboard/metrics/
    """
    org = get_user_organisation(request)
    scans_qs = Scan.objects.filter(organisation=org)
    total_scans = scans_qs.count()
    completed_scans = scans_qs.filter(status='completed').count()

    # Get findings by risk level
    findings_by_risk = Finding.objects.filter(organisation=org).values('risk_rating').annotate(
        count=Count('id')
    )

    # Recent scans
    recent_scans = scans_qs.filter(
        created_at__gte=timezone.now() - timedelta(days=7)
    ).count()

    return Response({
        'total_scans': total_scans,
        'completed_scans': completed_scans,
        'failed_scans': scans_qs.filter(status='failed').count(),
        'recent_scans_week': recent_scans,
        'findings_by_risk': {item['risk_rating']: item['count'] for item in findings_by_risk},
        'total_assets': Asset.objects.filter(organisation=org).count(),
        'total_domains': Domain.objects.filter(organisation=org).count(),
    })

@api_view(['GET'])
def debug_scan_data(request, scan_id):
    """
    Debug endpoint to check scan data integrity
    GET /api/debug/scan/{scan_id}/
    """
    scan = _scoped_scan(request, scan_id)
    
    # Count findings
    scan_findings_count = ScanFinding.objects.filter(scan=scan).count()
    unique_findings = ScanFinding.objects.filter(scan=scan).values('finding_id').distinct().count()
    
    # Count by category
    findings_by_category = {}
    scan_findings_qs = ScanFinding.objects.filter(scan=scan).select_related('finding')
    for sf in scan_findings_qs:
        cat = sf.finding.category
        findings_by_category[cat] = findings_by_category.get(cat, 0) + 1
    
    # Check CVSS scores
    findings_with_cvss = Finding.objects.filter(
        id__in=ScanFinding.objects.filter(scan=scan).values('finding_id')
    ).exclude(cvss_score__isnull=True).exclude(cvss_score=0.0).count()
    
    findings_without_cvss = Finding.objects.filter(
        id__in=ScanFinding.objects.filter(scan=scan).values('finding_id')
    ).filter(Q(cvss_score__isnull=True) | Q(cvss_score=0.0)).count()
    
    # Sample findings
    sample_findings = Finding.objects.filter(
        id__in=ScanFinding.objects.filter(scan=scan).values('finding_id')
    ).exclude(cvss_score__isnull=True).exclude(cvss_score=0.0).order_by('-cvss_score')[:5]
    
    sample_findings_data = [
        {
            'id': f.id,
            'title': f.title,
            'cvss_score': f.cvss_score,
            'category': f.category
        }
        for f in sample_findings
    ]
    
    # Check assets
    assets_scanned = ScanAsset.objects.filter(scan=scan).count()
    assets_with_org = ScanAsset.objects.filter(
        scan=scan,
        asset__domain__owner__isnull=False
    ).exclude(asset__domain__owner='').count()
    
    # Get organization names
    org_names = []
    for sa in ScanAsset.objects.filter(scan=scan).select_related('asset__domain'):
        if sa.asset.domain and sa.asset.domain.owner:
            org_names.append({
                'asset': sa.asset.value,
                'domain': sa.asset.domain.root_domain,
                'organization': sa.asset.domain.owner
            })
    
    return Response({
        'scan_id': scan_id,
        'scan_status': scan.status,
        'scan_duration_seconds': scan.duration_seconds,
        'total_scan_findings_records': scan_findings_count,
        'unique_findings': unique_findings,
        'findings_by_category': findings_by_category,
        'cvss_statistics': {
            'findings_with_cvss_score': findings_with_cvss,
            'findings_without_cvss_score': findings_without_cvss,
            'sample_findings_with_cvss': sample_findings_data
        },
        'assets': {
            'total_scanned': assets_scanned,
            'with_organization_name': assets_with_org,
            'organization_names': org_names
        },
        'check_counts': {
            'library_checks': scan.library_checks.count(),
            'ssl_checks': scan.ssl_checks.count(),
            'email_checks': scan.email_checks.count(),
            'header_checks': scan.header_checks.count(),
            'dns_checks': scan.dns_checks.count(),
            'technology_checks': scan.technology_checks.count()
        }
    })


@api_view(['GET'])
def executive_dashboard(request):
    """
    Executive Summary Dashboard - Comprehensive metrics
    GET /api/dashboard/executive/
    
    Returns all key metrics organized by priority tiers
    """
    logger.info("Generating executive dashboard...")

    # Platform admins may view the executive summary across every organisation
    # (narrow to one org by passing organisation_id). Everyone else is scoped to
    # their own organisation. org_filter is spread into each queryset below; an
    # empty dict means "all organisations".
    if is_platform_admin(request.user):
        org_id = requested_org_id(request)
        org_filter = {'organisation_id': org_id} if org_id else {}
    else:
        org_filter = {'organisation': get_user_organisation(request)}

    # ============================================
    # TIER 1 - CORE OVERVIEW
    # ============================================

    # Overall Security Posture
    total_findings = Finding.objects.filter(**org_filter).count()
    critical_findings = Finding.objects.filter(**org_filter, risk_rating='Critical', status='open').count()
    high_findings = Finding.objects.filter(**org_filter, risk_rating='High', status='open').count()
    medium_findings = Finding.objects.filter(**org_filter, risk_rating='Medium', status='open').count()
    low_findings = Finding.objects.filter(**org_filter, risk_rating='Low', status='open').count()
    
    # Calculate overall security score (0-100)
    # total_checks = (
    #     FrontendLibraryCheck.objects.count() +
    #     SSLTLSCheck.objects.count() +
    #     EmailSecurityCheck.objects.count() +
    #     SecurityHeaderCheck.objects.count() +
    #     DNSSecurityCheck.objects.count()
    # )
    
    # failed_checks = (
    #     FrontendLibraryCheck.objects.filter(
    #         Q(vulnerability_status='vulnerable') | Q(vulnerability_status='outdated')
    #     ).count() +
    #     SSLTLSCheck.objects.exclude(risk_rating='Low').count() +
    #     EmailSecurityCheck.objects.filter(status='FAIL').count() +
    #     SecurityHeaderCheck.objects.filter(status='missing').count() +
    #     DNSSecurityCheck.objects.exclude(risk_rating='Low').count()
    # )
    
    # overall_security_score = 100
    # if total_checks > 0:
    #     overall_security_score = round(((total_checks - failed_checks) / total_checks) * 100)

    # Calculate overall security score (0-100)
    total_checks = (
        FrontendLibraryCheck.objects.filter(**org_filter).count() +
        SSLTLSCheck.objects.filter(**org_filter).count() +
        EmailSecurityCheck.objects.filter(**org_filter).count() +
        SecurityHeaderCheck.objects.filter(**org_filter).count() +
        DNSSecurityCheck.objects.filter(**org_filter).count()
    )

    failed_checks = (
        FrontendLibraryCheck.objects.filter(**org_filter).filter(
            Q(vulnerability_status='vulnerable') | Q(vulnerability_status='outdated')
        ).count() +
        SSLTLSCheck.objects.filter(**org_filter).exclude(risk_rating='Low').count() +
        EmailSecurityCheck.objects.filter(**org_filter, status='FAIL').count() +
        SecurityHeaderCheck.objects.filter(**org_filter, status='missing').count() +
        DNSSecurityCheck.objects.filter(**org_filter).exclude(risk_rating='Low').count()
    )

    check_score = 100
    if total_checks > 0:
        check_score = round(((total_checks - failed_checks) / total_checks) * 100)

    vuln_penalty = (3 * critical_findings) + (2 * high_findings) + (1 * medium_findings) + (0.5 * low_findings)
    overall_security_score = max(0, check_score - min(20, vuln_penalty))
    
    # Key Metrics
    total_scans = Scan.objects.filter(**org_filter).count()
    completed_scans = Scan.objects.filter(**org_filter, status='completed').count()
    failed_scans = Scan.objects.filter(**org_filter, status='failed').count()
    scans_this_week = Scan.objects.filter(
        **org_filter,
        created_at__gte=timezone.now() - timedelta(days=7)
    ).count()

    # Asset Overview
    total_domains = Domain.objects.filter(**org_filter).count()
    total_assets = Asset.objects.filter(**org_filter).count()

    # Scan Status Summary
    pending_scans = Scan.objects.filter(**org_filter, status__in=['queued', 'running']).count()
    
    # ============================================
    # TIER 2 - CRITICAL INSIGHTS
    # ============================================
    
    # Worst Performing Owners (Top 5)
    worst_owners = []
    domains_with_owner = Domain.objects.filter(**org_filter).exclude(owner__isnull=True).exclude(owner='')
    
    for domain in domains_with_owner:
        # Get all assets for this domain
        asset_ids = domain.assets.values_list('id', flat=True)
        
        # Count findings by risk for this domain's assets
        owner_critical = Finding.objects.filter(
            asset_id__in=asset_ids,
            risk_rating='Critical'
        ).count()
        
        owner_high = Finding.objects.filter(
            asset_id__in=asset_ids,
            risk_rating='High'
        ).count()
        
        # Calculate security score for this owner
        owner_total_checks = (
            FrontendLibraryCheck.objects.filter(asset_id__in=asset_ids).count() +
            SSLTLSCheck.objects.filter(asset_id__in=asset_ids).count() +
            EmailSecurityCheck.objects.filter(asset_id__in=asset_ids).count() +
            SecurityHeaderCheck.objects.filter(asset_id__in=asset_ids).count() +
            DNSSecurityCheck.objects.filter(asset_id__in=asset_ids).count()
        )
        
        owner_failed_checks = (
            FrontendLibraryCheck.objects.filter(
                asset_id__in=asset_ids
            ).filter(
                Q(vulnerability_status='vulnerable') | Q(vulnerability_status='outdated')
            ).count() +
            SSLTLSCheck.objects.filter(asset_id__in=asset_ids).exclude(risk_rating='Low').count() +
            EmailSecurityCheck.objects.filter(asset_id__in=asset_ids, status='FAIL').count() +
            SecurityHeaderCheck.objects.filter(asset_id__in=asset_ids, status='missing').count() +
            DNSSecurityCheck.objects.filter(asset_id__in=asset_ids).exclude(risk_rating='Low').count()
        )
        
        owner_score = 100
        if owner_total_checks > 0:
            owner_score = round(((owner_total_checks - owner_failed_checks) / owner_total_checks) * 100)
        
        if owner_critical > 0 or owner_high > 0:
            worst_owners.append({
                'owner': domain.owner,
                'critical_count': owner_critical,
                'high_count': owner_high,
                'security_score': owner_score,
                'total_assets': len(asset_ids)
            })
    
    # Sort by critical first, then high, then score
    worst_owners.sort(key=lambda x: (-x['critical_count'], -x['high_count'], x['security_score']))
    worst_owners = worst_owners[:5]
    
    # Worst Performing Assets (Top 10)
    worst_assets = []
    for asset in Asset.objects.filter(**org_filter):
        finding_count = Finding.objects.filter(asset=asset).count()
        critical_count = Finding.objects.filter(asset=asset, risk_rating='Critical').count()
        high_count = Finding.objects.filter(asset=asset, risk_rating='High').count()
        
        # Determine risk level
        risk_level = 'Low'
        if critical_count > 0:
            risk_level = 'Critical'
        elif high_count > 0:
            risk_level = 'High'
        elif Finding.objects.filter(asset=asset, risk_rating='Medium').exists():
            risk_level = 'Medium'
        
        if finding_count > 0:
            worst_assets.append({
                'asset_id': asset.id,
                'asset_value': asset.value,
                'asset_type': asset.asset_type,
                'finding_count': finding_count,
                'critical_count': critical_count,
                'high_count': high_count,
                'risk_level': risk_level
            })
    
    worst_assets.sort(key=lambda x: (-x['critical_count'], -x['high_count'], -x['finding_count']))
    worst_assets = worst_assets[:10]
    
    # Most Common Vulnerabilities (Top 10)
    vulnerability_counts = Finding.objects.filter(
        **org_filter, category='CVE'
    ).values('title', 'risk_rating').annotate(
        frequency=Count('id')
    ).order_by('-frequency')[:10]
    
    most_common_vulnerabilities = [
        {
            'vulnerability_name': item['title'],
            'frequency': item['frequency'],
            'risk_level': item['risk_rating']
        }
        for item in vulnerability_counts
    ]
    
    # ============================================
    # TIER 3 - DETAILED ANALYSIS
    # ============================================
    
    # Most Vulnerable Libraries (Top 5)
    library_stats = FrontendLibraryCheck.objects.filter(**org_filter).values('library_name').annotate(
        outdated_count=Count('id', filter=Q(vulnerability_status='outdated')),
        vulnerable_count=Count('id', filter=Q(vulnerability_status='vulnerable')),
        total_count=Count('id')
    ).filter(
        Q(outdated_count__gt=0) | Q(vulnerable_count__gt=0)
    ).order_by('-vulnerable_count', '-outdated_count')[:5]
    
    most_vulnerable_libraries = [
        {
            'library_name': item['library_name'],
            'outdated_count': item['outdated_count'],
            'vulnerable_count': item['vulnerable_count'],
            'total_instances': item['total_count']
        }
        for item in library_stats
    ]
    
    # Most Commonly Missing Headers (Top 5)
    header_stats = SecurityHeaderCheck.objects.filter(
        **org_filter, status='missing'
    ).values('header').annotate(
        frequency=Count('id'),
        affected_assets=Count('asset', distinct=True)
    ).order_by('-frequency')[:5]
    
    most_missing_headers = [
        {
            'header_name': item['header'],
            'frequency': item['frequency'],
            'affected_asset_count': item['affected_assets']
        }
        for item in header_stats
    ]
    
    # ============================================
    # TIER 4 - PERFORMANCE INSIGHTS
    # ============================================
    
    # Best Performing Owners (Top 5)
    best_owners = []
    for domain in domains_with_owner:
        asset_ids = domain.assets.values_list('id', flat=True)
        
        owner_critical = Finding.objects.filter(
            asset_id__in=asset_ids,
            risk_rating='Critical'
        ).count()
        
        owner_high = Finding.objects.filter(
            asset_id__in=asset_ids,
            risk_rating='High'
        ).count()
        
        owner_total_checks = (
            FrontendLibraryCheck.objects.filter(asset_id__in=asset_ids).count() +
            SSLTLSCheck.objects.filter(asset_id__in=asset_ids).count() +
            EmailSecurityCheck.objects.filter(asset_id__in=asset_ids).count() +
            SecurityHeaderCheck.objects.filter(asset_id__in=asset_ids).count() +
            DNSSecurityCheck.objects.filter(asset_id__in=asset_ids).count()
        )
        
        owner_failed_checks = (
            FrontendLibraryCheck.objects.filter(
                asset_id__in=asset_ids
            ).filter(
                Q(vulnerability_status='vulnerable') | Q(vulnerability_status='outdated')
            ).count() +
            SSLTLSCheck.objects.filter(asset_id__in=asset_ids).exclude(risk_rating='Low').count() +
            EmailSecurityCheck.objects.filter(asset_id__in=asset_ids, status='FAIL').count() +
            SecurityHeaderCheck.objects.filter(asset_id__in=asset_ids, status='missing').count() +
            DNSSecurityCheck.objects.filter(asset_id__in=asset_ids).exclude(risk_rating='Low').count()
        )
        
        owner_score = 100
        if owner_total_checks > 0:
            owner_score = round(((owner_total_checks - owner_failed_checks) / owner_total_checks) * 100)
        
        best_owners.append({
            'owner': domain.owner,
            'security_score': owner_score,
            'total_assets': len(asset_ids),
            'critical_count': owner_critical,
            'high_count': owner_high
        })
    
    # Sort by score (highest first), then by least critical/high issues
    best_owners.sort(key=lambda x: (-x['security_score'], x['critical_count'], x['high_count']))
    best_owners = best_owners[:5]
    
    # Domain Security Score Distribution
    score_distribution = {
        'excellent': 0,    # 81-100
        'good': 0,         # 61-80
        'fair': 0,         # 41-60
        'at_risk': 0,      # 21-40
        'critical': 0      # 0-20
    }
    
    for domain in Domain.objects.filter(**org_filter):
        asset_ids = domain.assets.values_list('id', flat=True)
        if not asset_ids:
            continue

        domain_total_checks = (
            FrontendLibraryCheck.objects.filter(asset_id__in=asset_ids).count() +
            SSLTLSCheck.objects.filter(asset_id__in=asset_ids).count() +
            EmailSecurityCheck.objects.filter(asset_id__in=asset_ids).count() +
            SecurityHeaderCheck.objects.filter(asset_id__in=asset_ids).count() +
            DNSSecurityCheck.objects.filter(asset_id__in=asset_ids).count()
        )
        
        if domain_total_checks == 0:
            continue
        
        domain_failed_checks = (
            FrontendLibraryCheck.objects.filter(
                asset_id__in=asset_ids
            ).filter(
                Q(vulnerability_status='vulnerable') | Q(vulnerability_status='outdated')
            ).count() +
            SSLTLSCheck.objects.filter(asset_id__in=asset_ids).exclude(risk_rating='Low').count() +
            EmailSecurityCheck.objects.filter(asset_id__in=asset_ids, status='FAIL').count() +
            SecurityHeaderCheck.objects.filter(asset_id__in=asset_ids, status='missing').count() +
            DNSSecurityCheck.objects.filter(asset_id__in=asset_ids).exclude(risk_rating='Low').count()
        )
        
        domain_score = round(((domain_total_checks - domain_failed_checks) / domain_total_checks) * 100)
        
        if domain_score >= 81:
            score_distribution['excellent'] += 1
        elif domain_score >= 61:
            score_distribution['good'] += 1
        elif domain_score >= 41:
            score_distribution['fair'] += 1
        elif domain_score >= 21:
            score_distribution['at_risk'] += 1
        else:
            score_distribution['critical'] += 1
    
    # Assets by Owner
    assets_by_owner = []
    for domain in domains_with_owner:
        asset_ids = list(domain.assets.values_list('id', flat=True))
        
        owner_critical = Finding.objects.filter(
            asset_id__in=asset_ids,
            risk_rating='Critical'
        ).count()
        
        owner_high = Finding.objects.filter(
            asset_id__in=asset_ids,
            risk_rating='High'
        ).count()
        
        owner_total_checks = (
            FrontendLibraryCheck.objects.filter(asset_id__in=asset_ids).count() +
            SSLTLSCheck.objects.filter(asset_id__in=asset_ids).count() +
            EmailSecurityCheck.objects.filter(asset_id__in=asset_ids).count() +
            SecurityHeaderCheck.objects.filter(asset_id__in=asset_ids).count() +
            DNSSecurityCheck.objects.filter(asset_id__in=asset_ids).count()
        )
        
        owner_failed_checks = (
            FrontendLibraryCheck.objects.filter(
                asset_id__in=asset_ids
            ).filter(
                Q(vulnerability_status='vulnerable') | Q(vulnerability_status='outdated')
            ).count() +
            SSLTLSCheck.objects.filter(asset_id__in=asset_ids).exclude(risk_rating='Low').count() +
            EmailSecurityCheck.objects.filter(asset_id__in=asset_ids, status='FAIL').count() +
            SecurityHeaderCheck.objects.filter(asset_id__in=asset_ids, status='missing').count() +
            DNSSecurityCheck.objects.filter(asset_id__in=asset_ids).exclude(risk_rating='Low').count()
        )
        
        avg_score = 100
        if owner_total_checks > 0:
            avg_score = round(((owner_total_checks - owner_failed_checks) / owner_total_checks) * 100)
        
        assets_by_owner.append({
            'owner': domain.owner,
            'total_assets': len(asset_ids),
            'critical_issues': owner_critical,
            'high_issues': owner_high,
            'avg_security_score': avg_score
        })
    
    assets_by_owner.sort(key=lambda x: -x['total_assets'])
    
    # ============================================
    # FINAL RESPONSE
    # ============================================
    
    return Response({
        'tier_1_core_overview': {
            'overall_security_posture': {
                'status': 'At Risk' if critical_findings > 0 else 'Good' if high_findings == 0 else 'Fair',
                'security_score': overall_security_score,
                'critical_issues': critical_findings,
                'high_issues': high_findings,
                'medium_issues': medium_findings,
                'low_issues': low_findings
            },
            'key_metrics': {
                'total_scans': total_scans,
                'completed_scans': completed_scans,
                'failed_scans': failed_scans,
                'scans_this_week': scans_this_week
            },
            'asset_overview': {
                'total_domains': total_domains,
                'total_assets': total_assets
            },
            'findings_by_risk_level': {
                'critical': critical_findings,
                'high': high_findings,
                'medium': medium_findings,
                'low': low_findings,
                'total_findings': total_findings
            },
            'scan_status_summary': {
                'completed': completed_scans,
                'failed': failed_scans,
                'pending_in_progress': pending_scans
            }
        },
        'tier_2_critical_insights': {
            'worst_performing_owners': worst_owners,
            'worst_performing_assets': worst_assets,
            'most_common_vulnerabilities': most_common_vulnerabilities
        },
        'tier_3_detailed_analysis': {
            'most_vulnerable_libraries': most_vulnerable_libraries,
            'most_commonly_missing_headers': most_missing_headers
        },
        'tier_4_performance_insights': {
            'best_performing_owners': best_owners,
            'domain_security_score_distribution': {
                'excellent_81_100': score_distribution['excellent'],
                'good_61_80': score_distribution['good'],
                'fair_41_60': score_distribution['fair'],
                'at_risk_21_40': score_distribution['at_risk'],
                'critical_0_20': score_distribution['critical']
            },
            'assets_by_owner': assets_by_owner
        }
    })