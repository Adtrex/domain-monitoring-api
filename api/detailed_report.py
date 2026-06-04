"""
Detailed, step-by-step security report generator for NITDA-style actionable reports.
Generates a DOCX file with grouped findings, plain-language risks, asset lists, and full remediation steps.
"""

from io import BytesIO
import docx
from docx.shared import Pt, RGBColor, Inches
from datetime import datetime


def generate_detailed_report_docx(report_title, executive_summary, grouped_findings, organization_name=None, output_path=None):
    document = docx.Document()

    # Title
    title_p = document.add_paragraph()
    title_p.alignment = 1
    if organization_name:
        title_text = f"{report_title} for {organization_name}"
    else:
        title_text = report_title
    title_run = title_p.add_run(title_text)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = "Verdana"

    # Date
    date_p = document.add_paragraph()
    date_p.alignment = 1
    date_run = date_p.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.name = "Verdana"
    document.add_paragraph()

    # Executive Summary
    exec_heading = document.add_paragraph()
    exec_heading.add_run("Executive Summary").bold = True
    exec_sum = document.add_paragraph(executive_summary)
    exec_sum.alignment = 3
    for run in exec_sum.runs:
        run.font.name = "Verdana"
        run.font.size = Pt(11)
    document.add_paragraph()

    # Findings by Severity
    for severity, findings in grouped_findings.items():
        if not findings:
            continue
        sev_heading = document.add_paragraph()
        sev_heading.add_run(f"{severity} Priority Recommendations").bold = True
        for idx, finding in enumerate(findings, 1):
            # Title
            f_title = document.add_paragraph()
            f_title.add_run(f"{idx}. {finding['title']}").bold = True
            # Risk
            risk_p = document.add_paragraph(f"Risk: {finding['risk']}")
            # Affected Assets
            assets_p = document.add_paragraph("Affected Assets:")
            for asset in finding['affected_assets']:
                assets_p.add_run(f"\n● {asset}")
            # Recommendations
            recs_p = document.add_paragraph("Recommendations:")
            for rec in finding['recommendations']:
                recs_p.add_run(f"\n● {rec}")
            # Reference Links
            if finding.get('references'):
                refs_p = document.add_paragraph("References:")
                for ref in finding['references']:
                    refs_p.add_run(f"\n{ref}")
            document.add_paragraph()

    # Save or return
    output = BytesIO()
    document.save(output)
    output.seek(0)
    if output_path:
        with open(output_path, "wb") as f:
            f.write(output.getvalue())
    return output

# Example usage (to be replaced with real data wiring):
if __name__ == "__main__":
    report_title = "NITDA Digital Asset Security Assessment Report"
    executive_summary = "During our regular monitoring ... (see user example above)"
    grouped_findings = {
        "Critical": [
            {
                "title": "SSL/TLS Weak Protocol Support",
                "risk": "Attackers can intercept, decrypt, or manipulate encrypted communications through man-in-the-middle attacks.",
                "affected_assets": ["nitda.gov.ng (ports 443, 465, 993, 995)", "iicp.nitda.gov.ng (ports 443, 3389, 21)"],
                "recommendations": [
                    "Disable all SSL 2.0, SSL 3.0, TLS 1.0, and TLS 1.1 protocols immediately",
                    "Enforce TLS 1.2 as minimum, with TLS 1.3 preferred",
                    "Disable weak cipher suites across all affected services",
                    "If any of the above services are no longer in active use, decommission them properly. Remove DNS records, shut down the service, revoke associated certificates, and close open ports."
                ],
                "references": ["https://owasp.org/www-project-top-ten/2017/A6_2017-Security_Misconfiguration"]
            }
        ],
        "High": [],
        "Medium": [],
        "Low": []
    }
    generate_detailed_report_docx(report_title, executive_summary, grouped_findings, output_path="detailed_report.docx")
